import streamlit as st
import anthropic
import json
import plotly.graph_objects as go
import time
import requests
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime
import gspread
from google.oauth2.service_account import Credentials
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment as XlAlign

# Load API key — Streamlit secrets take priority, fallback to hardcoded for local dev

def _mval(field):
    """Get display string from market field — handles both new dict and legacy string format."""
    if isinstance(field, dict):
        return field.get("value", "N/A")
    return str(field) if field else "N/A"

TAVILY_KEY = ""  # optional

st.set_page_config(page_title="Schaeffler Innovation Assistant", page_icon="🟢", layout="wide", initial_sidebar_state="expanded")

# ═══════════════════════════════════════════════════════════════
#  INTRO PAGE  —  runs before everything, including API key check
# ═══════════════════════════════════════════════════════════════
if not st.session_state.get("_intro_done", False):
    import streamlit.components.v1 as _components

    # Single combined call: style + layout in one block, no split calls
    st.markdown("""
<style>
html,body,.stApp,[data-testid="stAppViewContainer"]{background:#050f07!important;}
.block-container{padding:0!important;max-width:100%!important;}
section[data-testid="stSidebar"],header[data-testid="stHeader"],
#MainMenu,footer,[data-testid="stToolbar"]{display:none!important;}
div[data-testid="stButton"]>button{
    background:#007A3D!important;color:#fff!important;border:none!important;
    border-radius:3px!important;font-size:16px!important;font-weight:700!important;
    letter-spacing:2.5px!important;padding:14px 40px!important;
    font-family:Arial,sans-serif!important;
    width:auto!important;display:inline-block!important;
}
div[data-testid="stButton"]>button:hover{background:#005A2B!important;}
div[data-testid="stButton"]{display:flex!important;justify-content:center!important;margin-top:4px!important;}
</style>""", unsafe_allow_html=True)

    # Hero rendered via components.html — guaranteed rendering, not affected by markdown parser
    _components.html("""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
*{margin:0;padding:0;box-sizing:border-box;}
body{
  background:#050f07;
  min-height:85vh;
  display:flex;align-items:center;justify-content:center;
  font-family:Arial,'Helvetica Neue',Helvetica,sans-serif;
  overflow:hidden;position:relative;
  padding:40px 24px;
}
.glow{
  position:absolute;top:-80px;left:50%;transform:translateX(-50%);
  width:800px;height:460px;
  background:radial-gradient(ellipse at center,rgba(0,122,61,0.22) 0%,transparent 68%);
  pointer-events:none;
}
.grid{
  position:absolute;inset:0;
  background-image:
    linear-gradient(rgba(0,122,61,0.055) 1px,transparent 1px),
    linear-gradient(90deg,rgba(0,122,61,0.055) 1px,transparent 1px);
  background-size:44px 44px;pointer-events:none;
}
.content{position:relative;z-index:2;text-align:center;max-width:680px;width:100%;}
.badge-row{display:flex;align-items:center;justify-content:center;gap:18px;margin-bottom:28px;}
.s-badge{
  width:56px;height:56px;background:#007A3D;border-radius:9px;
  display:flex;align-items:center;justify-content:center;flex-shrink:0;
  box-shadow:0 0 32px rgba(0,122,61,0.5);
}
.s-badge span{font-size:43px;font-weight:900;color:#fff;line-height:1;}
.wordmark{text-align:left;}
.wordmark .name{font-size:35px;font-weight:800;letter-spacing:5px;color:#fff;line-height:1;}
.wordmark .tagline{font-size:10px;letter-spacing:4.5px;font-weight:400;color:rgba(255,255,255,0.38);margin-top:5px;}
.divider{width:44px;height:2px;background:#007A3D;margin:0 auto 26px auto;box-shadow:0 0 10px rgba(0,122,61,0.6);}
h1{font-size:28px;font-weight:700;letter-spacing:1.5px;color:#fff;margin:0 0 14px 0;line-height:1.35;}
.subtitle{font-size:16px;color:rgba(255,255,255,0.42);margin:0 0 38px 0;line-height:1.75;letter-spacing:0.2px;}
.pills{display:flex;flex-wrap:wrap;justify-content:center;gap:6px;margin-bottom:8px;}
.pill{
  background:rgba(0,122,61,0.10);border:1px solid rgba(0,122,61,0.28);
  border-radius:3px;padding:7px 14px;
  font-size:12px;font-weight:700;letter-spacing:1.3px;
  color:rgba(255,255,255,0.55);white-space:nowrap;
}
</style>
</head>
<body>
<div class="glow"></div>
<div class="grid"></div>
<div class="content">
  <div class="badge-row">
    <div class="s-badge"><span>S</span></div>
    <div class="wordmark">
      <div class="name">SCHAEFFLER</div>
      <div class="tagline">WE PIONEER MOTION</div>
    </div>
  </div>
  <div class="divider"></div>
  <h1>AI Innovation Research Assistant</h1>
  <p class="subtitle">
    Structured evaluation for radical &amp; disruptive innovation ideas<br>
    at the fuzzy front end of Schaeffler&apos;s innovation process.
  </p>
  <div class="pills">
    <div class="pill">01 &middot; QUADRANT CLASSIFIER</div>
    <div class="pill">02 &middot; MARKET INTELLIGENCE</div>
    <div class="pill">03 &middot; PATENT INTELLIGENCE</div>
    <div class="pill">04 &middot; TECHNICAL FEASIBILITY</div>
    <div class="pill">05 &middot; P&sup3; PERSPECTIVE</div>
    <div class="pill">06 &middot; SCORING &amp; SYNTHESIS</div>
  </div>
</div>
</body>
</html>
""", height=480, scrolling=False)

    if st.button("BEGIN INNOVATION RESEARCH  \u2192", key="_intro_btn"):
        st.session_state["_intro_done"] = True
        st.rerun()

    st.markdown(
        "<p style='text-align:center;font-family:Arial,sans-serif;font-size:13px;"
        "letter-spacing:1px;color:rgba(255,255,255,0.15);margin-top:16px;'>"
        "MBA Capstone &middot; EBS Universit&auml;t f&uuml;r Wirtschaft und Recht &middot; Arpan Chowdhury</p>",
        unsafe_allow_html=True
    )
    st.stop()
# ═══════════════════════════════════════════════════════════════
#  END INTRO
# ═══════════════════════════════════════════════════════════════

# ── API key — loaded from Streamlit secrets ───────────────────────────────────
try:
    ANTHROPIC_KEY = st.secrets["ANTHROPIC_API_KEY"]
    if not ANTHROPIC_KEY or not ANTHROPIC_KEY.startswith("sk-"):
        raise ValueError("Key looks invalid")
except Exception:
    st.error(
        "⚠️ **Anthropic API key not found or invalid.**\n\n"
        "Add your key to Streamlit secrets:\n"
        "1. Open your app on Streamlit Cloud → **Settings → Secrets**\n"
        "2. Add: `ANTHROPIC_API_KEY = \"sk-ant-...your-key...\"`\n"
        "3. Save and reboot the app.\n\n"
        "For local dev, create `.streamlit/secrets.toml` with the same line."
    )
    st.stop()

# ── Google Sheets — Idea Log ──────────────────────────────────────────────────
SHEET_ID = "1Ya-z55BtzRS7NYiKiM8U8E0-NChueTVprJovvUrvZ6s"

_SA_INFO = {
    "type": "service_account",
    "project_id": "absolute-dahlia-450007-r8",
    "private_key_id": "292e1cb7c9d249e34ad3f89ab2f180e865ea1576",
    "private_key": "-----BEGIN PRIVATE KEY-----\nMIIEvAIBADANBgkqhkiG9w0BAQEFAASCBKYwggSiAgEAAoIBAQCbuv2KGVKo3RN+\nY2vvVAsjGRBFxu8j8p8tGYxrvaxc0g0wefRZq0hOJbh9zBZu+dL4JwWf5SAJQ+Yp\nLdqctTIl+M6vBJNL1J8iNZxtmOmKrQeGvI+O+CnkU/4ZVMLk+lhOUrSVyJ8rPXOl\nT4E5ar8j+ZpWtueKtODgdKVgYHpTX4kjYj+/ochlFG6DQfhSWmJDf52Gfd/nldZt\n13+n3/uHYmLI1wKNFEpVaq3Qd2r191bcVAY8/NrOtRaqtHsoINAP2Sn6rEuGthaa\niZECrIkYiMKkjTWk8E8XIv9Zv2JnBOWPErIHRbGGHV6WBz5p0Z0L1GLGDiuhvha8\nvq6AiVsXAgMBAAECggEAHCGClDQv1NYeo9mU0UY2vs/Tuy8M2ssEivKPBZVdMeU1\nwbh99ca1iHxS39KCiOhy/iWaZABRMatEw9KHJ4Cpvuc7eq0SaIPPfS//AmM5aLYJ\n4oJkUlisxJSRlYTUseUxF3DkMxxq+DYhEk8S0krgnUCE6z4eBFXZO2KGzyqOXknf\nCS0ThZ2ky0/ytuIoj8OeCkanuBQZXj219zH0Kiamqt7AWs7qSUoafnejkkQGJ5xR\nKrYr+Y0D4L9LFmZEbidlfoHZOcz7vrHvkx3RrL+RA2QJU7XolzR6sl7ecwzLGifs\nvsCxzD52+7pOvwNUp/hHpol/2bD8W0xkDcRpqCFrKQKBgQDKJ6V12ElUixKz9djx\n3X8SLpQTJihAqIts5gesaNp1mAG7/3nAq1cThprwKaYPukkPYbwGjN/IPXwZKfZr\nfyHxFBA/a5ooz2bJoEm4w4w8dIrqxbDd9K2HuqKSnW0ouU/Hustv7gE+zkyGFRIx\nJ08hJwQaMuAfulzt/7vhbOvjDwKBgQDFNcvjltsm8gw80jjCxh3+6r+UDFPncVDL\niL5P/q6pfCzulsdxFVZUfx9yh9mCrxGopbdZnmzQW86GuhfnPC+AQpYdltGTOUdx\n+calCeSAmtfoztI3P+CmR4DDiZ4c2rfPlqOwwHueJXo62u8984y7NsGBdCRCrBSG\njVzFWFBneQKBgHTF/BUbsBhPEam0rPHh0cJN96ksFHptIcTxB6O3GeJtwSq4w7rg\n/ra/vYZXeJ6DLCrfeP6Lp8UCh0n97GNiF9grj8siu/UxAR4dIhjBlKNjas99DNLZ\nwNeznq90kpbAnO4x38wzPrLp9lhJma2dGF99Kyh7FO4e+Ale/UeVZJlPAoGAB8N6\nZ1dFAV9+A9byzRgnjiWHrThfBTl8yMZ1V4jbL2joC+x7pYQFhgYLIuMeOPrTYyRC\n95A5EGrM0pj43+2KoS394uRRE86pdV8z5sNg738pCM07kVk+as1d0FTWmKQzoER5\n5TdupmcrTK3ZxUKVQ7mAHKyJ0OYdWL6v7ETxxWECgYA2RXMn67Pb0CtfX1Gs7ztB\n8N3PMOtYdz/TQbOK0u4waWrzbnFUCvsH39yoqiDofYdUDN/E/31IXoTzCi3N4F2U\ncvJ3oid6ZubsBA2gJrHTYvd/yljV+o8kuzUV4d2/fxTuphhh2TyoYa3BUD+W7alr\nDs4l7RfxIUIgMviRmVpAdw==\n-----END PRIVATE KEY-----\n",
    "client_email": "schaeffler-capstone@absolute-dahlia-450007-r8.iam.gserviceaccount.com",
    "client_id": "105367783720778854419",
    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
    "token_uri": "https://oauth2.googleapis.com/token",
    "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
    "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/schaeffler-capstone%40absolute-dahlia-450007-r8.iam.gserviceaccount.com",
    "universe_domain": "googleapis.com"
}

SHEET_COLUMNS = [
    "Date", "Submitter Name", "Position", "Department",
    "Full Idea Description", "Clarifying Q&A",
    "Quadrant", "Innovation Cluster", "Product Family",
    "Market Score", "Patent Score", "Feasibility Score", "P³ Score", "IPI Score",
    "Recommendation", "Key Concerns", "Next Steps",
    "Market Name", "Market Size 2024", "CAGR",
    "TRL Level", "Build Strategy"
]

def _fix_sa_info(raw: dict) -> dict:
    """Normalise a service-account dict so the private key always has real newlines."""
    sa = dict(raw)
    key = sa.get("private_key", "")
    # Step 1: collapse double-escaped \n → single backslash+n
    key = key.replace("\\n", "\n")
    # Step 2: replace backslash+n with a real newline character.
    # chr(10) is used instead of "\n" to avoid any Python source escaping ambiguity.
    key = key.replace("\n", chr(10))
    sa["private_key"] = key
    return sa

def _sheets_client():
    """
    Build an authorised gspread client.
    Authentication priority:
      1. st.secrets["GOOGLE_CREDENTIALS"]  — JSON string pasted into Streamlit secrets
                                             (most reliable; preserves newlines exactly)
      2. st.secrets["gcp_service_account"] — TOML section in Streamlit secrets
      3. Hardcoded _SA_INFO fallback        — will fail if GCP key has been revoked

    To set up option 1 in Streamlit Cloud:
      Settings → Secrets → paste:
        GOOGLE_CREDENTIALS = '''<paste full JSON from GCP service account key file>'''
    """
    # Option 1: JSON string in secrets (most reliable — newlines preserved verbatim)
    try:
        raw_json = st.secrets["GOOGLE_CREDENTIALS"]
        sa_info = _fix_sa_info(json.loads(raw_json))
        return gspread.service_account_from_dict(sa_info)
    except KeyError:
        pass
    except Exception:
        pass

    # Option 2: TOML [gcp_service_account] section
    try:
        sa_info = _fix_sa_info(dict(st.secrets["gcp_service_account"]))
        return gspread.service_account_from_dict(sa_info)
    except KeyError:
        pass
    except Exception:
        pass

    # Option 3: Hardcoded fallback (may fail if GCP key was revoked)
    sa_info = _fix_sa_info(dict(_SA_INFO))
    return gspread.service_account_from_dict(sa_info)

def load_past_ideas():
    try:
        ws = _sheets_client().open_by_key(SHEET_ID).sheet1
        return ws.get_all_records()
    except Exception:
        return []

def save_idea_to_sheets(row_data: dict):
    try:
        sh = _sheets_client().open_by_key(SHEET_ID)
        ws = sh.sheet1
        existing = ws.get_all_values()
        # Write header if sheet is completely empty
        if not existing:
            ws.append_row(SHEET_COLUMNS, value_input_option="RAW")
        row = [str(row_data.get(col, "")) for col in SHEET_COLUMNS]
        # append_row always adds to the next empty row — no index arithmetic needed
        ws.append_row(row, value_input_option="USER_ENTERED")
        return True, None
    except Exception as e:
        return False, str(e)


def check_similar_ideas(new_idea, past_ideas):
    if not past_ideas:
        return []
    past_summaries = "\n".join([
        (
            f"- [{r.get('Date','')}] {r.get('Submitter Name','Unknown')} ({r.get('Department','')}) | "
            f"Quadrant: {r.get('Quadrant','') or 'Unknown'} | "
            f"IPI: {r.get('IPI Score','') or 'Unknown'} | "
            f"Recommendation: {r.get('Recommendation','') or 'Unknown'} | "
            f"Idea: {str(r.get('Full Idea Description',''))[:200]}"
        )
        for r in past_ideas[:30]
    ])
    result = call_claude(
        'You compare innovation ideas. The past ideas list includes their Quadrant, IPI Score, and Recommendation — use these exact values in your output, do not write "Unknown" if the value is present. Return ONLY valid JSON: {"similar": [{"date": "...", "submitter": "...", "department": "...", "idea_snippet": "...", "quadrant": "...", "ipi": "...", "recommendation": "...", "similarity": "High/Medium", "reason": "one sentence"}]}. Return empty similar array if nothing is genuinely similar.',
        f"New idea: {new_idea}\n\nPast researched ideas:\n{past_summaries}",
        max_tokens=600
    )
    try:
        raw = result.strip().replace("```json","").replace("```","").strip()
        fb = raw.find("{"); lb = raw.rfind("}") + 1
        if fb >= 0: raw = raw[fb:lb]
        return json.loads(raw).get("similar", [])
    except Exception:
        return []

def generate_ideas_excel():
    records = load_past_ideas()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Innovation Ideas"
    hdr_fill  = PatternFill("solid", fgColor="1F3864")
    hdr_font  = Font(bold=True, color="FFFFFF", size=10)
    alt_fill  = PatternFill("solid", fgColor="EAF1FB")
    if not records:
        ws.append(["No ideas recorded yet — complete a full pipeline assessment to populate this log."])
        ws["A1"].font = Font(italic=True, color="555555")
    else:
        headers = list(records[0].keys())
        ws.append(headers)
        for cell in ws[1]:
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = XlAlign(horizontal="center", wrap_text=True)
        for i, record in enumerate(records, start=2):
            ws.append([record.get(h, "") for h in headers])
            if i % 2 == 0:
                for cell in ws[i]:
                    cell.fill = alt_fill
        for col in ws.columns:
            max_len = max((len(str(cell.value or "")) for cell in col), default=10)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf

# ── Styling — sidebar identity + sidebar-toggle arrow fix ────
st.markdown("""
<style>
/* ── Sidebar: Schaeffler green identity ── */
section[data-testid="stSidebar"] {
    background-color: #007A3D !important;
    border-right: none !important;
    overflow-x: hidden !important;
}
section[data-testid="stSidebar"] > div {
    overflow-x: hidden !important;
}
section[data-testid="stSidebar"] * {
    color: #FFFFFF !important;
}


section[data-testid="stSidebar"] .stButton > button {
    background: transparent !important;
    color: rgba(255,255,255,0.65) !important;
    border: none !important;
    border-radius: 4px !important;
    font-size: 14px !important;
    text-align: left !important;
    justify-content: flex-start !important;
    display: flex !important;
    width: 100% !important;
}
section[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(255,255,255,0.10) !important;
    color: #FFFFFF !important;
}





/* ── Sidebar selectbox — match sidebar style ── */
section[data-testid="stSidebar"] .stSelectbox label {
    color: rgba(255,255,255,0.55) !important;
    font-size: 13px !important;
    letter-spacing: 1.5px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
}
section[data-testid="stSidebar"] .stSelectbox > div > div {
    background: transparent !important;
    border: 1px solid rgba(255,255,255,0.2) !important;
    color: rgba(255,255,255,0.8) !important;
    font-size: 14px !important;
    font-family: 'Arial','Helvetica Neue',Helvetica,sans-serif !important;
    border-radius: 4px !important;
}
section[data-testid="stSidebar"] .stSelectbox > div > div:hover {
    border-color: rgba(255,255,255,0.45) !important;
}
section[data-testid="stSidebar"] .stSelectbox svg {
    fill: rgba(255,255,255,0.6) !important;
    color: rgba(255,255,255,0.6) !important;
}

/* ── Language + Ideas Log fixed at bottom of sidebar ── */
.ideas-log-fixed {
    position: fixed !important;
    bottom: 0 !important;
    left: 0 !important;
    width: var(--sidebar-width, 246px) !important;
    max-width: 260px !important;
    background: #007A3D !important;
    border-top: 1px solid rgba(255,255,255,0.18) !important;
    padding: 10px 14px 16px 14px !important;
    z-index: 9999 !important;
    box-sizing: border-box !important;
}
.ideas-log-fixed .stButton > button,
.ideas-log-fixed .stDownloadButton > button {
    background: transparent !important;
    border: none !important;
    color: rgba(255,255,255,0.5) !important;
    font-size: 13px !important;
    font-family: 'Arial','Helvetica Neue',Helvetica,sans-serif !important;
    font-weight: 400 !important;
    letter-spacing: 0.3px !important;
    padding: 4px 0 !important;
    text-align: left !important;
    width: 100% !important;
}
.ideas-log-fixed .stButton > button:hover,
.ideas-log-fixed .stDownloadButton > button:hover {
    color: rgba(255,255,255,0.85) !important;
    background: transparent !important;
}
.ideas-log-fixed .stSelectbox label {
    color: rgba(255,255,255,0.45) !important;
    font-size: 11px !important;
    letter-spacing: 1.5px !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    margin-bottom: 2px !important;
}
.ideas-log-fixed .stSelectbox > div > div {
    background: transparent !important;
    border: 1px solid rgba(255,255,255,0.18) !important;
    color: rgba(255,255,255,0.65) !important;
    font-size: 13px !important;
    font-family: 'Arial','Helvetica Neue',Helvetica,sans-serif !important;
    border-radius: 3px !important;
    padding: 2px 6px !important;
}
.ideas-log-fixed .stSelectbox svg {
    fill: rgba(255,255,255,0.45) !important;
}

/* ── Source tag pill ── */
.source-tag {
    background:#1e3a5f; color:#93c5fd;
    font-size:11px; padding:2px 8px;
    border-radius:3px; margin-left:6px;
    font-family: monospace;
}

/* ── Global font consistency — Arial/Helvetica Neue throughout ── */
html, body, [class*="css"],
.stApp, .stMarkdown, .stText,
.stButton > button,
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stSelectbox > div > div,
.stMetric, .stChatMessage,
.stCaption, .stAlert,
h1, h2, h3, h4, h5, h6, p, div, span, label, li {
    font-family: 'Arial', 'Helvetica Neue', Helvetica, sans-serif !important;
}
/* Keep monospace only for code blocks */
code, pre, .stCode {
    font-family: 'Courier New', Courier, monospace !important;
}
</style>

<script>


// Inject favicon dynamically as an SVG data URI with Schaeffler S logo
(function() {
    var svgFavicon = "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 32 32'%3E%3Crect width='32' height='32' rx='5' fill='%23007A3D'/%3E%3Ctext x='16' y='24' font-family='Arial,Helvetica,sans-serif' font-size='22' font-weight='700' fill='white' text-anchor='middle'%3ES%3C/text%3E%3C/svg%3E";
    var existing = document.querySelector("link[rel*='icon']");
    if (existing) {
        existing.href = svgFavicon;
    } else {
        var link = document.createElement('link');
        link.rel = 'icon';
        link.type = 'image/svg+xml';
        link.href = svgFavicon;
        document.head.appendChild(link);
    }
    var shortcut = document.querySelector("link[rel='shortcut icon']");
    if (shortcut) shortcut.href = svgFavicon;
})();


</script>
""", unsafe_allow_html=True)

# ── Constants ─────────────────────────────────────────────────
BG    = "#0f1e35"
BLUE  = "#60a5fa"
DIM   = "#4a6fa5"
WHITE = "#e2e8f0"
NAVY  = "#1F3864"

# ── Language system ────────────────────────────────────────────
_LANG = {
    "en": {
        "pipeline":  "INNOVATION PIPELINE",
        "idea_cap":  "Idea",
        "quad_cap":  "Quadrant",
        "lang_label":"Language / Sprache",
        "stages": [
            (1,"01 · Quadrant Classifier"),
            (2,"02 · Market Intelligence"),
            (3,"03 · Patent Intelligence"),
            (4,"04 · Technical Feasibility"),
            (5,"05 · P³ Perspective"),
            (6,"06 · Scoring & Synthesis"),
        ],
        "dl_market":  "⬇️  Download Market Intelligence Report",
        "dl_patent":  "⬇️  Download Patent Intelligence Report",
        "dl_feasib":  "⬇️  Download Technical Feasibility Report",
        "dl_org":     "⬇️  Download P³ Perspective Report",
        "dl_master":  "⬇️  Download Full Innovation Assessment Report",
        "dl_ideas":   "↓  Download Ideas Log",
        "dl_spinner": "Generating report — please wait…",
        "dl_caption": "Covers all 6 stages: Quadrant Classification · Market · Patent · Feasibility · P³ Score · IPI",
        "claude_suffix": "",
        # Shared stage header label (used in all 6 stage info boxes)
        "stage_what_label": "WHAT THIS STAGE DOES",
        # Stage 01
        "s1_title": "Stage 01 · Quadrant Classifier",
        "s1_what": "Maps your idea onto Schaeffler's modified Ansoff matrix — Exploit, Extend, Radical, or Disrupt. Ideas in Radical and Disrupt proceed through the full pipeline. Others are redirected to the right Schaeffler product division.",
        "s1_what_label": "WHAT THIS STAGE DOES",
        "s1_you_get": "<b style='color:#e2e8f0;'>You get:</b> Quadrant classification · Schaeffler Motion product family fit · Strategic trend alignment · Innovation pathway (Start-Up Mode vs Innovation Factory)",
        "s1_step1": "Step 1 — Who are you & describe your idea",
        "s1_name": "Your name *", "s1_name_ph": "e.g. Anna Müller",
        "s1_role": "Position / Role *", "s1_role_ph": "e.g. Senior Engineer",
        "s1_dept": "Department *", "s1_dept_ph": "e.g. E-Mobility R&D",
        "s1_idea_label": "What is your innovation idea?",
        "s1_idea_ph": "e.g. A self-lubricating bearing system that uses micro-reservoirs embedded within the bearing material to release lubricant automatically based on temperature and load sensing...",
        "s1_submit": "Submit idea",
        "s1_warn_identity": "Please fill in your name, position, and department before continuing.",
        "s1_warn_empty": "Please enter your idea first.",
        "s1_warn_brief": "A bit brief — can you add more detail? What does it do, and for whom?",
        "s1_step2": "Step 2 — Three quick questions",
        "s1_step2_caption": "These answers help refine the classification. The idea description itself drives the result — use these to add context, not to override what the idea clearly is.",
        "s1_q1": "**1. Has the core technology behind this idea been demonstrated anywhere — in a lab, a startup, a research paper, or a competitor product?**",
        "s1_q1a": "Yes — it has been demonstrated somewhere (even if not commercialised)",
        "s1_q1b": "No — the underlying technology is genuinely novel or theoretical",
        "s1_q1_detail": "Optional — where has it been demonstrated, or what makes it genuinely new?",
        "s1_q1_ph": "e.g. MIT lab prototype, or 'no known demonstration of this mechanism'",
        "s1_q2": "**2. Does this idea target markets or applications that Schaeffler currently operates in?**",
        "s1_q2_caption": "Schaeffler's current markets: automotive (ICE & EV), industrial machinery, rail, aerospace, energy, two-wheelers. If the idea fits any of these, select yes.",
        "s1_q2a": "Yes — it targets automotive, industrial, rail, aerospace, energy or adjacent sectors Schaeffler already serves",
        "s1_q2b": "No — it targets a market genuinely outside Schaeffler's current scope (e.g. consumer electronics, healthcare devices, retail)",
        "s1_q2_detail": "Optional — which specific market or application area?",
        "s1_q2_ph": "e.g. EV drivetrain OEMs, or 'medical implant manufacturers — entirely new for Schaeffler'",
        "s1_q3": "**3. Is the problem this idea solves already recognised and being worked on by the industry?**",
        "s1_q3a": "Yes — the problem is well known and others are actively trying to solve it",
        "s1_q3b": "No — the problem itself is new, underappreciated, or not yet widely recognised",
        "s1_q3_detail": "Optional — describe the problem in one sentence",
        "s1_q3_ph": "e.g. Bearing failure in EV drivetrains due to high-frequency current leakage",
        "s1_classify_btn": "Classify my idea →",
        "s1_result": "Result",
        "s1_qualifies": "✓ This idea qualifies for the full Innovation pipeline.",
        "s1_continue": "Continue to Stage 02: Market Intelligence →",
        "s1_full_run": "⚡ Run Full Analysis — all 5 stages",
        "s1_chat_header": "💬 Questions about this classification?",
        "s1_chat_ph": "Ask about the classification...",
        "s1_startover": "← Start over",
        "s1_spinner_check": "Checking your idea...",
        "s1_spinner_similar": "Checking against past ideas...",
        "s1_spinner_classify": "Classifying your idea...",
        "s1_similar_warn": "⚠️ **{n} similar idea(s) previously researched** — review before proceeding.",
        # Stage 02
        "s2_title": "Stage 02 · Market Intelligence",
        "s2_what": "Analyses the commercial opportunity behind your idea — how big the market is, how fast it is growing, who the competitors are, and how well the idea fits across Schaeffler's 10 customer sector clusters.",
        "s2_you_get": "<b style='color:#e2e8f0;'>You get:</b> Market size & CAGR with sources · Sector cluster fit chart · Competitor landscape · Market Intelligence Score (0–10)",
        "s2_run_btn": "Run Market Intelligence →",
        "s2_continue": "Continue to Stage 03: Patent Intelligence →",
        "s2_rerun": "← Re-run analysis",
        "s2_success": "✓ Market Intelligence complete. Final score: **{score}/10**",
        "s2_chat_header": "💬 Questions about the market analysis?",
        "s2_chat_ph": "Ask about the market, sectors, competitors...",
        # Stage 03
        "s3_title": "Stage 03 · Patent Intelligence",
        "s3_what": "Maps the patent landscape for your idea's core technology — who is filing, whether they are competitors or potential customers, where the IP white spaces are, and how Schaeffler's existing patent portfolio relates to the idea.",
        "s3_you_get": "<b style='color:#e2e8f0;'>You get:</b> Patent Ansoff map with all key filers plotted · IP white spaces · Schaeffler IP gap analysis · Patent Intelligence Score (0–10)",
        "s3_run_btn": "Run Patent Intelligence →",
        "s3_continue": "Continue to Stage 04: Technical Feasibility →",
        "s3_rerun": "← Re-run analysis",
        "s3_success": "✓ Patent Intelligence complete. Final score: **{score}/10**",
        "s3_chat_header": "💬 Questions about the patent analysis?",
        "s3_chat_ph": "Ask about patents, IP risks, white spaces...",
        # Stage 04
        "s4_title": "Stage 04 · Technical Feasibility",
        "s4_what": "Assesses whether the core technology of your idea actually exists and how mature it is — using the Technology Readiness Level (TRL) framework used by NASA, the EU, and industrial R&D organisations.",
        "s4_you_get": "<b style='color:#e2e8f0;'>You get:</b> Technology existence verdict · TRL level (1–9) with rationale · Schaeffler entry readiness · Technical Feasibility Score (0–10)",
        "s4_run_btn": "Run Technical Feasibility →",
        "s4_continue": "Continue to Stage 05: P³ Perspective →",
        "s4_rerun": "← Re-run analysis",
        "s4_success": "✓ Technical Feasibility complete. Final score: **{score}/10**",
        "s4_chat_header": "💬 Questions about the feasibility analysis?",
        "s4_chat_ph": "Ask about TRL, technology gaps, development timeline...",
        # Stage 05
        "s5_title": "Stage 05 · P³ Perspective",
        "s5_what": "Assesses whether Schaeffler has the organisational capability to develop and commercialise this idea — using the P³ formula: Performance = Portfolio × People × Process.",
        "s5_you_get": "<b style='color:#e2e8f0;'>You get:</b> P³ readiness scores · Build vs Partner recommendation · Key capability gaps · P³ Perspective Score (0–10)",
        "s5_run_btn": "Run P³ Perspective →",
        "s5_continue": "Continue to Stage 06: Scoring & Synthesis →",
        "s5_rerun": "← Re-run analysis",
        "s5_success": "✓ P³ Perspective complete. Final score: **{score}/10**",
        "s5_chat_header": "💬 Questions about the org readiness analysis?",
        "s5_chat_ph": "Ask about P³ scores, capability gaps, partnerships...",
        # Stage 06
        "s6_title": "Stage 06 · Scoring & Synthesis",
        "s6_what": "Synthesises all five pipeline dimensions into a single Innovation Potential Index (IPI) score and strategic recommendation — with full narrative, risk analysis, and concrete next steps.",
        "s6_you_get": "<b style='color:#e2e8f0;'>You get:</b> IPI score (0–10) · Strategic recommendation · Radar chart · Key concerns & next steps · Full narrative synthesis",
        "s6_run_btn": "Run Scoring & Synthesis →",
        "s6_rerun": "← Adjust weights and re-run",
        "s6_chat_header": "💬 Questions about the overall assessment?",
        "s6_chat_ph": "Ask about the IPI score, recommendation, or next steps...",
        "s6_image_btn": "🎨 Generate Solution Blueprint",
        "s6_image_redo": "🔄 Regenerate Blueprint",
    },
    "de": {
        "pipeline":  "INNOVATIONS-PIPELINE",
        "idea_cap":  "Idee",
        "quad_cap":  "Quadrant",
        "lang_label":"Sprache",
        "stages": [
            (1,"01 · Quadrant-Klassifikator"),
            (2,"02 · Marktintelligenz"),
            (3,"03 · Patentintelligenz"),
            (4,"04 · Technische Machbarkeit"),
            (5,"05 · P³ Perspektive"),
            (6,"06 · Bewertung & Synthese"),
        ],
        "dl_market":  "⬇️  Marktintelligenz-Bericht herunterladen",
        "dl_patent":  "⬇️  Patentintelligenz-Bericht herunterladen",
        "dl_feasib":  "⬇️  Technischen Machbarkeitsbericht herunterladen",
        "dl_org":     "⬇️  Org. Bereitschaftsbericht herunterladen",
        "dl_master":  "⬇️  Vollständigen Innovationsbericht herunterladen",
        "dl_ideas":   "↓  Ideen-Log herunterladen",
        "dl_spinner": "Bericht wird erstellt — bitte warten…",
        "dl_caption": "Umfasst alle 6 Stufen: Quadrant-Klassifikation · Markt · Patente · Machbarkeit · Org. Bereitschaft · IPI",
        # Shared stage header label (used in all 6 stage info boxes)
        "stage_what_label": "WAS DIESE STUFE TUT",
        # Stage 01
        "s1_title": "Stufe 01 · Quadrant-Klassifikator",
        "s1_what": "Ordnet Ihre Idee in die modifizierte Ansoff-Matrix von Schaeffler ein — Exploit, Extend, Radical oder Disrupt. Ideen in Radical und Disrupt durchlaufen die gesamte Pipeline. Andere werden an die passende Produktsparte weitergeleitet.",
        "s1_what_label": "WAS DIESE STUFE TUT",
        "s1_you_get": "<b style='color:#e2e8f0;'>Sie erhalten:</b> Quadrant-Klassifikation · Passung zur Schaeffler Motion-Produktfamilie · Strategische Trendausrichtung · Innovationspfad (Start-Up-Modus vs. Innovationsfabrik)",
        "s1_step1": "Schritt 1 — Wer sind Sie & beschreiben Sie Ihre Idee",
        "s1_name": "Ihr Name *", "s1_name_ph": "z.B. Anna Müller",
        "s1_role": "Position / Rolle *", "s1_role_ph": "z.B. Senior Engineer",
        "s1_dept": "Abteilung *", "s1_dept_ph": "z.B. E-Mobilität F&E",
        "s1_idea_label": "Was ist Ihre Innovationsidee?",
        "s1_idea_ph": "z.B. Ein selbstschmierendes Lagersystem, das Mikro-Reservoirs im Lagermaterial nutzt, um Schmiermittel automatisch basierend auf Temperatur- und Lastsensoren freizusetzen...",
        "s1_submit": "Idee einreichen",
        "s1_warn_identity": "Bitte füllen Sie Name, Position und Abteilung aus, bevor Sie fortfahren.",
        "s1_warn_empty": "Bitte geben Sie zuerst Ihre Idee ein.",
        "s1_warn_brief": "Etwas kurz — können Sie mehr Details hinzufügen? Was macht es, und für wen?",
        "s1_step2": "Schritt 2 — Drei kurze Fragen",
        "s1_step2_caption": "Diese Antworten helfen bei der Verfeinerung der Klassifikation. Die Ideenbeschreibung selbst bestimmt das Ergebnis — nutzen Sie diese nur zur Ergänzung.",
        "s1_q1": "**1. Wurde die Kerntechnologie dieser Idee irgendwo demonstriert — in einem Labor, einem Startup, einer Forschungsarbeit oder einem Konkurrenzprodukt?**",
        "s1_q1a": "Ja — sie wurde irgendwo demonstriert (auch wenn nicht kommerzialisiert)",
        "s1_q1b": "Nein — die zugrunde liegende Technologie ist wirklich neu oder theoretisch",
        "s1_q1_detail": "Optional — wo wurde sie demonstriert, oder was macht sie wirklich neu?",
        "s1_q1_ph": "z.B. MIT-Labor-Prototyp, oder 'kein bekannter Nachweis dieses Mechanismus'",
        "s1_q2": "**2. Zielt diese Idee auf Märkte oder Anwendungen, in denen Schaeffler derzeit tätig ist?**",
        "s1_q2_caption": "Aktuelle Märkte von Schaeffler: Automotive (Verbrenner & EV), Industriemaschinen, Bahn, Luft- und Raumfahrt, Energie, Zweiräder.",
        "s1_q2a": "Ja — sie zielt auf Automotive, Industrie, Bahn, Luft- und Raumfahrt, Energie oder angrenzende Sektoren",
        "s1_q2b": "Nein — sie zielt auf einen Markt außerhalb des aktuellen Schaeffler-Portfolios (z.B. Unterhaltungselektronik, Medizintechnik)",
        "s1_q2_detail": "Optional — welcher spezifische Markt oder Anwendungsbereich?",
        "s1_q2_ph": "z.B. EV-Antriebsstrang-OEMs, oder 'Medizinimplantat-Hersteller — völlig neu für Schaeffler'",
        "s1_q3": "**3. Wird das Problem, das diese Idee löst, bereits von der Branche erkannt und bearbeitet?**",
        "s1_q3a": "Ja — das Problem ist bekannt und andere versuchen es aktiv zu lösen",
        "s1_q3b": "Nein — das Problem selbst ist neu, unterschätzt oder noch nicht weit verbreitet",
        "s1_q3_detail": "Optional — beschreiben Sie das Problem in einem Satz",
        "s1_q3_ph": "z.B. Lagerversagen in EV-Antriebssträngen durch hochfrequente Stromableitung",
        "s1_classify_btn": "Meine Idee klassifizieren →",
        "s1_result": "Ergebnis",
        "s1_qualifies": "✓ Diese Idee qualifiziert sich für die vollständige Innovationspipeline.",
        "s1_continue": "Weiter zu Stufe 02: Marktintelligenz →",
        "s1_full_run": "⚡ Vollanalyse — alle 5 Stufen",
        "s1_chat_header": "💬 Fragen zur Klassifikation?",
        "s1_chat_ph": "Fragen zur Klassifikation stellen...",
        "s1_startover": "← Von vorne beginnen",
        "s1_spinner_check": "Idee wird geprüft...",
        "s1_spinner_similar": "Vergleich mit früheren Ideen...",
        "s1_spinner_classify": "Idee wird klassifiziert...",
        "s1_similar_warn": "⚠️ **{n} ähnliche Idee(n) wurden bereits untersucht** — bitte vor dem Fortfahren prüfen.",
        # Stage 02
        "s2_title": "Stufe 02 · Marktintelligenz",
        "s2_what": "Analysiert die kommerzielle Chance hinter Ihrer Idee — wie groß der Markt ist, wie schnell er wächst, wer die Wettbewerber sind, und wie gut die Idee zu Schaefflers 10 Kundensektorclustern passt.",
        "s2_you_get": "<b style='color:#e2e8f0;'>Sie erhalten:</b> Marktgröße & CAGR mit Quellen · Sektorcluster-Passungsdiagramm · Wettbewerbslandschaft · Marktintelligenz-Score (0–10)",
        "s2_run_btn": "Marktintelligenz starten →",
        "s2_continue": "Weiter zu Stufe 03: Patentintelligenz →",
        "s2_rerun": "← Analyse wiederholen",
        "s2_success": "✓ Marktintelligenz abgeschlossen. Endergebnis: **{score}/10**",
        "s2_chat_header": "💬 Fragen zur Marktanalyse?",
        "s2_chat_ph": "Fragen zu Markt, Sektoren, Wettbewerbern...",
        # Stage 03
        "s3_title": "Stufe 03 · Patentintelligenz",
        "s3_what": "Kartiert die Patentlandschaft für die Kerntechnologie Ihrer Idee — wer anmeldet, ob es Wettbewerber oder potenzielle Kunden sind, wo die IP-Weißräume liegen, und wie Schaefflers bestehendes Patentportfolio mit der Idee zusammenhängt.",
        "s3_you_get": "<b style='color:#e2e8f0;'>Sie erhalten:</b> Patent-Ansoff-Karte · IP-Weißräume · Schaeffler-IP-Lückenanalyse · Patentintelligenz-Score (0–10)",
        "s3_run_btn": "Patentintelligenz starten →",
        "s3_continue": "Weiter zu Stufe 04: Technische Machbarkeit →",
        "s3_rerun": "← Analyse wiederholen",
        "s3_success": "✓ Patentintelligenz abgeschlossen. Endergebnis: **{score}/10**",
        "s3_chat_header": "💬 Fragen zur Patentanalyse?",
        "s3_chat_ph": "Fragen zu Patenten, IP-Risiken, Weißräumen...",
        # Stage 04
        "s4_title": "Stufe 04 · Technische Machbarkeit",
        "s4_what": "Bewertet, ob die Kerntechnologie Ihrer Idee tatsächlich existiert und wie ausgereift sie ist — unter Verwendung des TRL-Rahmens (Technology Readiness Level) von NASA, EU und industriellen F&E-Organisationen.",
        "s4_you_get": "<b style='color:#e2e8f0;'>Sie erhalten:</b> Technologie-Existenzbewertung · TRL-Stufe (1–9) mit Begründung · Schaeffler-Eintrittsvorbereitung · Technische Machbarkeits-Score (0–10)",
        "s4_run_btn": "Technische Machbarkeit starten →",
        "s4_continue": "Weiter zu Stufe 05: P³ Perspektive →",
        "s4_rerun": "← Analyse wiederholen",
        "s4_success": "✓ Technische Machbarkeit abgeschlossen. Endergebnis: **{score}/10**",
        "s4_chat_header": "💬 Fragen zur Machbarkeitsanalyse?",
        "s4_chat_ph": "Fragen zu TRL, Technologielücken, Entwicklungszeitplan...",
        # Stage 05
        "s5_title": "Stufe 05 · P³ Perspektive",
        "s5_what": "Bewertet, ob Schaeffler die P³-Perspektive hat, diese Idee zu entwickeln und zu kommerzialisieren — mit der P³-Formel: Leistung = Portfolio × People × Process.",
        "s5_you_get": "<b style='color:#e2e8f0;'>Sie erhalten:</b> P³-Bereitschaftsscores · Build-vs-Partner-Empfehlung · Wichtige Kompetenzlücken · P³ Score (0–10)",
        "s5_run_btn": "P³ Perspektive starten →",
        "s5_continue": "Weiter zu Stufe 06: Bewertung & Synthese →",
        "s5_rerun": "← Analyse wiederholen",
        "s5_success": "✓ P³ Perspektive abgeschlossen. Endergebnis: **{score}/10**",
        "s5_chat_header": "💬 Fragen zur Org.-Bereitschaftsanalyse?",
        "s5_chat_ph": "Fragen zu P³-Scores, Kompetenzlücken, Partnerschaften...",
        # Stage 06
        "s6_title": "Stufe 06 · Bewertung & Synthese",
        "s6_what": "Synthetisiert alle fünf Pipeline-Dimensionen zu einem einzigen Innovationspotenzialindex (IPI) und einer strategischen Empfehlung — mit vollständiger Erzählung, Risikoanalyse und konkreten nächsten Schritten.",
        "s6_you_get": "<b style='color:#e2e8f0;'>Sie erhalten:</b> IPI-Score (0–10) · Strategische Empfehlung · Radardiagramm · Wichtige Bedenken & nächste Schritte · Vollständige narrative Synthese",
        "s6_run_btn": "Bewertung & Synthese starten →",
        "s6_rerun": "← Gewichtungen anpassen und neu starten",
        "s6_chat_header": "💬 Fragen zur Gesamtbewertung?",
        "s6_chat_ph": "Fragen zu IPI-Score, Empfehlung oder nächsten Schritten...",
        "s6_image_btn": "🎨 Lösungs-Blueprint generieren",
        "s6_image_redo": "🔄 Blueprint neu generieren",
        "claude_suffix": (
            "\n\nWICHTIG: Antworte AUSSCHLIESSLICH auf Deutsch. "
            "Alle Analysen, Beschreibungen, Begründungen, Zusammenfassungen und sonstigen Texte "
            "müssen vollständig auf Deutsch verfasst sein — auch alle String-Werte in JSON-Antworten "
            "(z.B. market_name, rationale, focus, summary, growth_drivers usw.). "
            "JSON-Schlüssel bleiben auf Englisch. Eigennamen, Marken und internationale "
            "Fachbegriffe dürfen auf Englisch bleiben."
        ),
    },
}

def T(key):
    lang = st.session_state.get("ui_lang", "en")
    return _LANG.get(lang, _LANG["en"]).get(key, _LANG["en"].get(key, key))

def _lang_suffix():
    lang = st.session_state.get("ui_lang", "en")
    return _LANG.get(lang, _LANG["en"]).get("claude_suffix", "")

def _one_click_dl(label, gen_fn, filename):
    """Single-click: show progress bar, generate report, auto-download via JS."""
    import base64
    import streamlit.components.v1 as _stc
    MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    col_btn, _ = st.columns([3,1])
    with col_btn:
        if st.button(label, type="primary", use_container_width=True):
            prog = st.progress(0, text=T("dl_spinner"))
            try:
                prog.progress(20, text=T("dl_spinner"))
                buf = gen_fn()
                prog.progress(80, text=T("dl_spinner"))
                b64 = base64.b64encode(buf.getvalue()).decode()
                prog.progress(100, text="✓ Ready — downloading…")
                time.sleep(0.3)
                prog.empty()
                _stc.html(f"""
<a id="_sdl" href="data:{MIME};base64,{b64}" download="{filename}"
   style="display:inline-block;margin:6px 0 2px;padding:10px 22px;
          background:#007A3D;color:#fff;font-weight:700;border-radius:5px;
          text-decoration:none;font-family:Arial,sans-serif;font-size:14px;">
  ⬇️ {filename}
</a>
<script>(function(){{var a=document.getElementById('_sdl');if(a)a.click();}})();</script>
""", height=55)
            except Exception as exc:
                prog.empty()
                st.error(f"Report error: {exc}")

# ── Helpers ───────────────────────────────────────────────────
def call_claude(system, user, max_tokens=2000):
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
    full_sys = system + _lang_suffix()
    for attempt in range(3):
        try:
            msg = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=max_tokens,
                system=full_sys,
                messages=[{"role": "user", "content": user}]
            )
            return msg.content[0].text
        except Exception as e:
            err = str(e)
            if "overloaded" in err.lower() and attempt < 2:
                time.sleep(3)
            elif "auth" in err.lower() or "api_key" in err.lower() or "401" in err:
                st.error(f"API key error: {err}. Check your Streamlit secrets.")
                st.stop()
            else:
                raise e

def call_claude_chat(system, history, max_tokens=500):
    client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)
    full_sys = system + _lang_suffix()
    for attempt in range(3):
        try:
            msg = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=max_tokens,
                system=full_sys,
                messages=history
            )
            return msg.content[0].text
        except Exception as e:
            if "overloaded" in str(e).lower() and attempt < 2:
                time.sleep(3)
            else:
                return "The API is briefly overloaded — please try again in a moment."

def tavily_search(query):
    if not TAVILY_KEY:
        return []
    try:
        resp = requests.post("https://api.tavily.com/search", json={
            "api_key": TAVILY_KEY, "query": query,
            "search_depth": "basic", "max_results": 5
        }, timeout=10)
        return [{"title": r.get("title",""), "url": r.get("url",""),
                 "content": r.get("content","")} for r in resp.json().get("results",[])]
    except:
        return []


import re as _re_global
import re as _re_xml_san
_XML_CTRL_RE = _re_xml_san.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')
def _xs(s):
    """Strip XML-illegal control chars from any value before add_run()."""
    if s is None: return ''
    return _XML_CTRL_RE.sub('', str(s))


def _parse_json(raw: str) -> dict:
    """Robustly parse JSON from a Claude response.
    Handles: markdown fences, preamble/postamble, rubric annotations,
    trailing commas, and all bad characters inside string values:
    bare newlines, carriage returns, tabs, control chars, lone backslashes.
    Bare internal double-quotes are handled by the lookahead heuristic.
    """
    if not raw:
        raise ValueError("Empty response")
    text = raw.strip()
    # 1. Strip markdown fences
    text = _re_global.sub(r"```json\s*", "", text)
    text = _re_global.sub(r"```\s*", "", text).strip()
    # 2. Extract outermost { … }
    fb = text.find("{"); lb = text.rfind("}")
    if fb >= 0 and lb > fb:
        text = text[fb:lb+1]
    # 3. Strip inline rubric annotations: "field": 7 (some text) → "field": 7
    text = _re_global.sub(
        r'("[\w_]+"\s*:\s*)(\d+(?:\.\d+)?)\s*\([^)]*\)', r'\1\2', text
    )
    # 4. Remove trailing commas before } or ]
    text = _re_global.sub(r',\s*([}\]])', r'\1', text)

    # 5. Character-by-character sanitiser — fixes everything inside string values
    def _sanitise(s):
        VALID_ESC = set('"\\/ bfnrtu')  # valid JSON escape chars after backslash
        out = []
        in_str = False
        i = 0
        n = len(s)
        while i < n:
            c = s[i]
            if not in_str:
                if c == '"':
                    in_str = True
                out.append(c)
                i += 1
            else:
                # Inside a string value
                if c == '\\':
                    nxt = s[i+1] if i+1 < n else ''
                    if nxt in VALID_ESC:
                        out.append(c); out.append(nxt); i += 2  # valid escape — pass through
                    else:
                        out.append('\\\\'); i += 1              # lone backslash — escape it
                elif c == '"':
                    # Heuristic: real string terminator if next structural char
                    # (skipping whitespace) is , } ] or :
                    j = i + 1
                    while j < n and s[j] in ' \t\r\n':
                        j += 1
                    nxt_struct = s[j] if j < n else ''
                    if nxt_struct in (',', '}', ']', ':'):
                        in_str = False          # close the string
                        out.append(c)
                    else:
                        out.append('\\"')       # bare quote inside string — escape it
                    i += 1
                elif c in ('\n', '\r', '\t'):
                    out.append(' '); i += 1     # whitespace control chars → space
                elif ord(c) < 32:
                    i += 1                      # other control chars → drop
                else:
                    out.append(c); i += 1
        return ''.join(out)

    text = _sanitise(text)
    # 6. Final trailing-comma pass (sanitise may have shifted content)
    text = _re_global.sub(r',\s*([}\]])', r'\1', text)
    return json.loads(text)


# ── Session state ─────────────────────────────────────────────
defaults = {
    "active_stage": 1,
    "ui_lang": "en",
    # User identity
    "user_name": "",
    "user_position": "",
    "user_dept": "",
    # Stage 01
    "s1_step": 1,
    "s1_idea": "",
    "s1_questions": [],
    "s1_answers": [],
    "s1_classification": {},
    "s1_chat": [],
    "s1_similar_ideas": [],
    # Stage 02
    "s2_step": "intro",
    "s2_data": {},
    "s2_chat": [],
    # Stage 06 visual brief
    "s6_action_brief": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Sidebar ───────────────────────────────────────────────────
with st.sidebar:
    # ── Logo ──────────────────────────────────────────────────
    st.markdown(f"""
<div style="padding:20px 12px 12px 12px;">
  <div style="font-family:'Arial','Helvetica Neue',Helvetica,sans-serif;font-size:30px;font-weight:700;letter-spacing:3px;color:#FFFFFF;line-height:1;">SCHAEFFLER</div>
  <div style="font-family:'Arial','Helvetica Neue',Helvetica,sans-serif;font-size:12px;letter-spacing:3.5px;color:rgba(255,255,255,0.7);margin-top:3px;font-weight:400;">WE PIONEER MOTION</div>
  <div style="background:rgba(255,255,255,0.2);height:1px;margin:16px 0 12px 0;"></div>
  <div style="font-family:'Arial','Helvetica Neue',Helvetica,sans-serif;font-size:14px;letter-spacing:2px;color:rgba(255,255,255,0.65);font-weight:600;">{T("pipeline")}</div>
</div>
""", unsafe_allow_html=True)

    st.markdown('<div style="background:rgba(255,255,255,0.15);height:1px;margin:4px 0 10px;"></div>', unsafe_allow_html=True)

    # ── Stage navigation ───────────────────────────────────────
    stages = T("stages")
    completed = set()
    if st.session_state.get("s1_classification"): completed.add(1)
    if st.session_state.get("s2_data"):           completed.add(2)
    if st.session_state.get("s3_data"):           completed.add(3)
    if st.session_state.get("s4_data"):           completed.add(4)
    if st.session_state.get("s5_data"):           completed.add(5)
    if st.session_state.get("s6_data"):           completed.add(6)

    for num, label in stages:
        active = st.session_state.active_stage
        if num == active:
            st.markdown(f"""<div style="font-family:'Arial','Helvetica Neue',Helvetica,sans-serif;background:rgba(255,255,255,0.15);border-radius:4px;padding:7px 12px;margin:2px 0;font-size:14px;font-weight:700;color:#FFFFFF;border-left:3px solid #FFFFFF;letter-spacing:0.3px;text-align:left;">&#9658; {label}</div>""", unsafe_allow_html=True)
        elif num in completed:
            if st.button(f"✓  {label}", key=f"nav_{num}", use_container_width=True):
                st.session_state.active_stage = num
                st.rerun()
        else:
            st.markdown(f"""<div style="font-family:'Arial','Helvetica Neue',Helvetica,sans-serif;padding:7px 12px;margin:2px 0;font-size:14px;color:rgba(255,255,255,0.35);text-align:left;">&#9675; {label}</div>""", unsafe_allow_html=True)

    if st.session_state.s1_idea:
        st.markdown("---")
        st.caption(f"**{T('idea_cap')}:** {st.session_state.s1_idea[:60]}...")
        if st.session_state.s1_classification:
            st.caption(f"**{T('quad_cap')}:** {st.session_state.s1_classification.get('quadrant','')}")

    # ── Language + Ideas Log — fixed at bottom of sidebar ──
    st.markdown('<div class="ideas-log-fixed">', unsafe_allow_html=True)
    # Language selector
    _lang_opts   = ["English", "Deutsch"]
    _lang_to_key = {"English":"en","Deutsch":"de"}
    _key_to_lang = {"en":"English","de":"Deutsch"}
    _sel = st.selectbox(
        T("lang_label"), _lang_opts,
        index=_lang_opts.index(_key_to_lang.get(st.session_state.ui_lang,"English")),
        key="_lang_sel"
    )
    if _lang_to_key[_sel] != st.session_state.ui_lang:
        st.session_state.ui_lang = _lang_to_key[_sel]
        # Clear all cached stage outputs so every stage re-runs in the new language
        for _k in ["s2_data","s3_data","s4_data","s5_data","s6_data",
                   "s2_step","s3_step","s4_step","s5_step","s6_step",
                   "s2_chat","s3_chat","s4_chat","s5_chat","s6_chat",
                   "s1_classification","s1_chat","s1_similar_ideas"]:
            if _k in st.session_state:
                del st.session_state[_k]
        st.session_state.active_stage = 1
        st.rerun()
    st.markdown('<div style="background:rgba(255,255,255,0.12);height:1px;margin:6px 0 4px;"></div>', unsafe_allow_html=True)
    st.markdown("""
<a href="https://docs.google.com/spreadsheets/d/1Ya-z55BtzRS7NYiKiM8U8E0-NChueTVprJovvUrvZ6s/edit?usp=sharing"
   target="_blank"
   style="display:block;color:rgba(255,255,255,0.5);font-size:13px;font-weight:400;
          font-family:'Arial','Helvetica Neue',Helvetica,sans-serif;letter-spacing:0.3px;
          text-decoration:none;padding:4px 0;width:100%;line-height:1.8;">
  ↗ View Ideas Log
</a>""", unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# ── Ansoff chart helper ───────────────────────────────────────
def ansoff_chart(quadrant, tech_score, market_score):
    q_cols = {"EXPLOIT":"#1a2d45","EXTEND":"#1e3a5f","RADICAL":"#1F3864","DISRUPT":"#0d2137","DISRUPTIVE":"#0d2137"}
    text_col = "#e2e8f0"; dim_col = "#4a6fa5"; grid_col = "#2a4a70"

    fig = go.Figure()
    # Schaeffler convention: X=Technology (left=Established, right=New)
    #                        Y=Market (bottom=Established, top=New)
    # EXPLOIT=bottom-left, EXTEND=top-left, DISRUPTIVE=bottom-right, RADICAL=top-right
    # Schaeffler convention (from paper):
    # X axis = Market (left=Established → right=New to the World)
    # Y axis = Technology (bottom=Established → top=New to the World)
    # EXPLOIT = bottom-left (existing tech, existing market)
    # EXTEND  = top-left   (new tech, existing market) — Schaeffler calls this differently
    # RADICAL = top-right  (new tech, new market) — appears top-right
    # DISRUPTIVE = bottom-right (existing market novelty, new tech) 
    # Per Lau 2023: EXPLOIT(low/low) EXTEND(low/high) RADICAL(high/high) DISRUPTIVE(high/low)
    # X=Technology, Y=Market in paper. Radical+Disruptive on right side of X.
    for q in [
        dict(x=[0,5,5,0],   y=[0,0,5,5],   name="EXPLOIT",    lx=2.5, ly=2.5),
        dict(x=[0,5,5,0],   y=[5,5,10,10], name="EXTEND",     lx=2.5, ly=7.5),
        dict(x=[5,10,10,5], y=[5,5,10,10], name="RADICAL",    lx=7.5, ly=7.5),
        dict(x=[5,10,10,5], y=[0,0,5,5],   name="DISRUPTIVE", lx=7.5, ly=2.5),
    ]:
        fig.add_trace(go.Scatter(
            x=q["x"]+[q["x"][0]], y=q["y"]+[q["y"][0]],
            fill="toself", fillcolor=q_cols[q["name"]],
            line=dict(color=grid_col, width=1),
            mode="lines", showlegend=False, hoverinfo="skip"
        ))
        fig.add_annotation(x=q["lx"], y=q["ly"], text=f"<b>{q['name']}</b>",
            showarrow=False,
            font=dict(size=14, color=BLUE if q["name"]==quadrant else dim_col))

    fig.add_shape(type="line",x0=5,x1=5,y0=0,y1=10,line=dict(color=grid_col,width=1.5,dash="dot"))
    fig.add_shape(type="line",x0=0,x1=10,y0=5,y1=5,line=dict(color=grid_col,width=1.5,dash="dot"))

    # 4-level axis tick labels (X=Technology, Y=Market — Schaeffler convention)
    for x_pos, x_label in [(1.25,"Established"),(3.75,"Adjacent"),(6.25,"New to Schaeffler"),(8.75,"New to the World")]:
        fig.add_annotation(x=x_pos, y=-1.0, text=x_label, showarrow=False,
            font=dict(size=9, color=dim_col), textangle=0)
    for y_pos, y_label in [(1.25,"Established"),(3.75,"Adjacent"),(6.25,"New to\nSchaeffler"),(8.75,"New to\nthe World")]:
        fig.add_annotation(x=-1.3, y=y_pos, text=y_label, showarrow=False,
            font=dict(size=9, color=dim_col), textangle=-90)
    # Axis spine labels
    fig.add_annotation(x=5, y=-1.8, text="← Technology Dimension (Newness) →", showarrow=False,
        font=dict(size=10, color=text_col))
    fig.add_annotation(x=-2.4, y=5, text="← Market Dimension (Newness) →", showarrow=False,
        font=dict(size=10, color=text_col), textangle=-90)
    # Dividing lines at each level boundary (2.5, 5, 7.5)
    for v in [2.5, 5.0, 7.5]:
        fig.add_shape(type="line", x0=v, x1=v, y0=0, y1=10,
            line=dict(color=grid_col, width=0.8, dash="dot"))
        fig.add_shape(type="line", x0=0, x1=10, y0=v, y1=v,
            line=dict(color=grid_col, width=0.8, dash="dot"))

    fig.add_trace(go.Scatter(
        x=[tech_score], y=[market_score], mode="markers+text",
        marker=dict(size=20, color=BLUE, line=dict(color="white",width=2)),
        text=["  Your idea"], textposition="middle right",
        textfont=dict(size=12,color="white",family="Arial Bold"),
        showlegend=False,
        hovertemplate=f"<b>{quadrant}</b><br>Tech: {tech_score}/10<br>Market: {market_score}/10<extra></extra>"
    ))
    fig.update_layout(
        title=dict(text="Schaeffler Innovation Framework — Modified Ansoff Matrix",
                   font=dict(size=13,color=text_col),x=0.5),
        xaxis=dict(range=[-1.6,11],showticklabels=False,showgrid=False,zeroline=False,
                   title="Technology Dimension →",title_font=dict(size=11,color=dim_col)),
        yaxis=dict(range=[-1.6,11],showticklabels=False,showgrid=False,zeroline=False,
                   title="Market Dimension →",title_font=dict(size=11,color=dim_col)),
        plot_bgcolor=BG, paper_bgcolor=BG, height=420,
        margin=dict(l=110,r=30,t=50,b=70), font=dict(color=text_col)
    )
    return fig

def get_dot_position(quadrant, confidence):
    # X=Technology, Y=Market (Schaeffler convention)
    # base = (tech_x, market_y)
    base = {
        "EXPLOIT":    (2.5, 2.5),   # low tech, low market  — bottom-left
        "EXTEND":     (2.5, 7.5),   # low tech, high market — top-left
        "RADICAL":    (7.5, 7.5),   # high tech, high market — top-right
        "DISRUPTIVE": (7.5, 2.5),   # high tech, low market — bottom-right
        "DISRUPT":    (7.5, 2.5),   # alias
    }
    nudge = {"High":1.5,"Medium":1.0,"Low":0.5}.get(confidence,1.0)
    t, m = base.get(quadrant,(5,5))
    if quadrant in ("EXPLOIT",):      t-=nudge; m-=nudge
    elif quadrant in ("EXTEND",):     t-=nudge; m+=nudge
    elif quadrant in ("RADICAL",):    t+=nudge; m+=nudge
    elif quadrant in ("DISRUPTIVE","DISRUPT"): t+=nudge; m-=nudge
    return round(max(0.5,min(9.5,t)),1), round(max(0.5,min(9.5,m)),1)

# ════════════════════════════════════════════════════════════
# STAGE 01 — QUADRANT CLASSIFIER
# ════════════════════════════════════════════════════════════




def generate_org_report(idea, quadrant, s1c, s5d):
    """Generate a Stage 05 P³ Perspective Word report."""
    org_data   = s5d.get("org_data", {})
    portfolio  = org_data.get("p3_portfolio", {})
    people     = org_data.get("p3_people", {})
    process    = org_data.get("p3_process", {})
    bop        = org_data.get("build_or_partner", {})
    gaps       = org_data.get("org_gaps", [])
    partners   = org_data.get("partnership_candidates", [])

    org_ctx = (
        f"Idea: {idea}\nQuadrant: {quadrant}\n"
        f"P3 Portfolio score: {s5d.get('p_portfolio',5)}/10  Rationale: {portfolio.get('rationale','')}\n"
        f"Cluster fit: {portfolio.get('cluster_fit','')}  Strengths: {chr(44).join(portfolio.get('strengths',[])[:3])}\n"
        f"P3 People score: {s5d.get('p_people',5)}/10  Rationale: {people.get('rationale','')}\n"
        f"Matched competencies: {chr(44).join(people.get('matched_competencies',[])[:4])}\n"
        f"Critical gap: {people.get('competency_gap','')}  Closure route: {people.get('sourcing_route','')}\n"
        f"P3 Process score: {s5d.get('p_process',5)}/10  Rationale: {process.get('rationale','')}\n"
        f"Applicable assets: {chr(44).join(process.get('applicable_assets',[])[:3])}\n"
        f"Investment required: {process.get('investment_required','')}  Time to close: {process.get('time_to_close','')}\n"
        f"Build strategy: {bop.get('recommendation','')}  Rationale: {bop.get('rationale','')}\n"
        f"Time internal: {bop.get('time_to_trl6_internal','')}  Time with partner: {bop.get('time_to_trl6_partner','')}\n"
        f"Overall org readiness score: {s5d.get('final_score',5)}/10"
    )
    extended = call_claude(
        '''You are a senior Schaeffler innovation strategist writing a detailed P³ Perspective report.
Write specific, substantive content referencing Schaeffler P3 formula (Performance = Portfolio x People x Process).
Return ONLY valid JSON, no markdown backticks:
{
  "executive_summary": "3-4 full paragraphs: overall P³ Perspective verdict, key P3 strengths and gaps, build-or-partner recommendation, and strategic rationale.",
  "portfolio_analysis": "2-3 full paragraphs on strategic portfolio fit — how this idea aligns with Schaeffler innovation clusters, trends, product families, and the broader electrification strategy post-Vitesco.",
  "people_analysis": "2-3 full paragraphs on human capital readiness — which Schaeffler competencies match, what is critically missing, and the most realistic route to close the gap (hire, upskill, acquire, partner).",
  "process_analysis": "2-3 full paragraphs on process and infrastructure readiness — which Schaeffler assets and processes apply directly, what needs to be built, and estimated investment and timeline.",
  "partnership_strategy": "2 full paragraphs on the recommended partnership strategy — who Schaeffler should partner with, what type of arrangement, and how to structure it.",
  "risks": ["Org risk 1 with mitigation", "Org risk 2 with mitigation", "Org risk 3", "Org risk 4"],
  "recommendations": ["Concrete org action 1 with timeline", "Concrete action 2", "Concrete action 3", "Concrete action 4"]
}''',
        org_ctx,
        max_tokens=3500
    )
    raw_e = extended.strip().replace("```json","").replace("```","").strip()
    fb = raw_e.find("{"); lb = raw_e.rfind("}") + 1
    if fb >= 0: raw_e = raw_e[fb:lb]
    try:
        ext = json.loads(raw_e)
    except:
        ext = {
            "executive_summary": f"Schaeffler\'s P³ Perspective for this idea scores {s5d.get('final_score',5)}/10. The P3 assessment shows Portfolio fit at {s5d.get('p_portfolio',5)}/10, People readiness at {s5d.get('p_people',5)}/10, and Process readiness at {s5d.get('p_process',5)}/10. Recommended build strategy: {bop.get('recommendation','Co-develop')}. {bop.get('rationale','')}",
            "portfolio_analysis": f"Portfolio fit: {portfolio.get('rationale','')} Cluster fit: {portfolio.get('cluster_fit','')}",
            "people_analysis": f"People readiness: {people.get('rationale','')} Critical gap: {people.get('competency_gap','')} Closure route: {people.get('sourcing_route','')}",
            "process_analysis": f"Process readiness: {process.get('rationale','')} Investment required: {process.get('investment_required','')} Time to close: {process.get('time_to_close','')}",
            "partnership_strategy": f"Partnership strategy: {bop.get('recommendation','Co-develop')}. Time to TRL6 internally: {bop.get('time_to_trl6_internal','')}. With partner: {bop.get('time_to_trl6_partner','')}.",
            "risks": ["Competency gap requires immediate sourcing action — delay increases time to TRL6", "Partnership negotiations can be slow — initiate early", "Internal process gaps may create bottlenecks in development", "Build vs buy decision requires board-level sign-off"],
            "recommendations": [f"Initiate {bop.get('recommendation','co-development')} process within 30 days", f"Address critical gap: {people.get('competency_gap','key competency')} via {people.get('sourcing_route','targeted hiring')}", "Map applicable Schaeffler assets to innovation project plan", "Establish steering committee for cross-divisional coordination"]
        }

    NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6)
    WHITE=RGBColor(0xFF,0xFF,0xFF); GREY=RGBColor(0x55,0x55,0x55)
    BLACK=RGBColor(0x00,0x00,0x00); LBLUE=RGBColor(0x60,0xA5,0xFA)

    def set_bg(cell, hx):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); tcPr.append(shd)
    def h1(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8"); bot.set(qn("w:space"),"3"); bot.set(qn("w:color"),"2E75B6")
        pBdr.append(bot); pPr.append(pBdr); r=p.add_run(_xs(text))
        r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY
    def h2(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(_xs(text)); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=NAVY
    def body(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; r=p.add_run(_xs(text)); r.font.size=Pt(10.5)
    def kv(doc, label, value):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r1=p.add_run(_xs(f"{label}: ")); r1.bold=True; r1.font.size=Pt(10.5); r1.font.color.rgb=NAVY
        r2=p.add_run(_xs(value)); r2.font.size=Pt(10.5)
    def bul(doc, text):
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(_xs(text)); r.font.size=Pt(10.5)

    doc=DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; c=t.cell(0,0); set_bg(c,"1F3864")
    p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("ORGANISATIONAL READINESS REPORT"); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=WHITE
    p2=c.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(10)
    r2=p2.add_run("Schaeffler AI Innovation Research Assistant  ·  Stage 05"); r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    doc.add_paragraph()
    p=doc.add_paragraph(); r=p.add_run("P³ Perspective Assessment")
    r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
    p2=doc.add_paragraph(); r2=p2.add_run(_xs(f"Score: {s5d.get('final_score',5)}/10  ·  Strategy: {bop.get('recommendation','')}  ·  Quadrant: {quadrant}  ·  {datetime.now().strftime('%d %B %Y')}")); r2.font.size=Pt(10); r2.italic=True; r2.font.color.rgb=GREY
    doc.add_paragraph()

    h1(doc,"P³ Score Summary")
    p3_tbl=doc.add_table(rows=1,cols=4); p3_tbl.style="Table Grid"
    for i,h in enumerate(["Dimension","Score","Weight","Description"]):
        c=p3_tbl.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    for i,(dim,score,wt,desc) in enumerate([
        ("Portfolio",f"{s5d.get('p_portfolio',5):.1f}/10","35%",portfolio.get('cluster_fit','')),
        ("People",f"{s5d.get('p_people',5):.1f}/10","40%",people.get('competency_gap','')),
        ("Process",f"{s5d.get('p_process',5):.1f}/10","25%",process.get('investment_required','')),
    ]):
        row=p3_tbl.add_row(); fill="EAF1FB" if i%2==0 else "FFFFFF"
        for c in row.cells: set_bg(c,fill)
        for j,val in enumerate([dim,score,wt,desc[:80] if desc else ""]):
            r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)
    fr=p3_tbl.add_row()
    for c in fr.cells: set_bg(c,"1F3864")
    r=fr.cells[0].paragraphs[0].add_run("OVERALL READINESS"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    r2=fr.cells[1].paragraphs[0].add_run(_xs(f"{s5d.get('final_score',5)}/10")); r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=LBLUE
    doc.add_paragraph()

    h1(doc,"Executive Summary"); body(doc, ext.get("executive_summary",""))
    h1(doc,"P³ Portfolio — Strategic Fit"); body(doc, ext.get("portfolio_analysis",""))
    if portfolio.get("strengths"):
        h2(doc,"Portfolio Strengths")
        for s in portfolio["strengths"]: bul(doc, f"✓ {s}")
    if portfolio.get("gaps"):
        h2(doc,"Portfolio Gaps")
        for g in portfolio["gaps"]: bul(doc, f"✗ {g}")
    h1(doc,"P³ People — Competency Readiness"); body(doc, ext.get("people_analysis",""))
    if people.get("matched_competencies"):
        h2(doc,"Matched Competencies")
        for c in people["matched_competencies"]: bul(doc, f"✓ {c}")
    kv(doc,"Critical competency gap", people.get("competency_gap",""))
    kv(doc,"Closure route", people.get("sourcing_route",""))
    h1(doc,"P³ Process — Infrastructure & Assets"); body(doc, ext.get("process_analysis",""))
    if process.get("applicable_assets"):
        h2(doc,"Applicable Assets")
        for a in process["applicable_assets"]: bul(doc, a)
    kv(doc,"Investment required", process.get("investment_required",""))
    kv(doc,"Estimated time to close", process.get("time_to_close",""))

    if gaps:
        h1(doc,"Organisational Gaps Register")
        gt=doc.add_table(rows=1,cols=4); gt.style="Table Grid"
        for i,h in enumerate(["Gap","Severity","Closure Route","Timeline"]):
            c=gt.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,g in enumerate(gaps):
            row=gt.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            sev_fill={"High":"FFE4E4","Medium":"FFF8E4","Low":"E4FFE9"}.get(g.get("severity",""),"FFFFFF")
            for c in row.cells: set_bg(c,sev_fill if idx==0 else fill)
            for j,val in enumerate([g.get("gap",""),g.get("severity",""),g.get("closure_route",""),g.get("timeline","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)

    if partners:
        h1(doc,"Partnership Candidates"); body(doc, ext.get("partnership_strategy",""))
        pt=doc.add_table(rows=1,cols=4); pt.style="Table Grid"
        for i,h in enumerate(["Organisation","Type","Rationale","Route"]):
            c=pt.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,p in enumerate(partners):
            row=pt.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([p.get("name",""),p.get("type",""),p.get("rationale",""),p.get("route","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)

    h1(doc,"Build-or-Partner Recommendation")
    kv(doc,"Recommendation", bop.get("recommendation",""))
    kv(doc,"Time to TRL 6 — Internal", bop.get("time_to_trl6_internal",""))
    kv(doc,"Time to TRL 6 — With Partner", bop.get("time_to_trl6_partner",""))
    body(doc, bop.get("rationale",""))
    h1(doc,"Risks"); 
    for risk in ext.get("risks",[]): bul(doc, risk)
    h1(doc,"Recommendations")
    for rec in ext.get("recommendations",[]): bul(doc, rec)

    doc.add_paragraph()
    ft=doc.add_table(rows=1,cols=1); ft.style="Table Grid"; fc=ft.cell(0,0); set_bg(fc,"1F3864")
    fp=fc.paragraphs[0]; fp.paragraph_format.space_before=Pt(6); fp.paragraph_format.space_after=Pt(6)
    fr=fp.add_run(_xs(f"Schaeffler AI Innovation Research Assistant  ·  Stage 05: P³ Perspective  ·  {datetime.now().strftime('%d %B %Y')}  ·  Capstone Project — Arpan Chowdhury, EBS Universität"))
    fr.font.size=Pt(8); fr.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    _sanitize_doc(doc)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


def generate_master_report(idea, quadrant, s1c, s2d, s3d, s4d, s5d_org, s6d):
    """Generate the full comprehensive Innovation Assessment Word report covering all stages."""
    ipi       = s6d.get("ipi", 0)
    weights   = s6d.get("weights", {"market":35,"patent":25,"feasibility":25,"org":15})
    synthesis = s6d.get("synthesis", {})
    scores    = s6d.get("scores", {})
    market    = s2d.get("market", {})
    comp      = s2d.get("comp", {})
    sectors   = s2d.get("sectors", {})
    landscape = s3d.get("landscape", {})
    ansoff_d  = s3d.get("ansoff_data", {})
    existence = s4d.get("existence", {})
    trl       = s4d.get("trl", {})
    org_data  = s5d_org.get("org_data", {})
    portfolio = org_data.get("p3_portfolio", {})
    people    = org_data.get("p3_people", {})
    process   = org_data.get("p3_process", {})
    bop       = org_data.get("build_or_partner", {})

    # Generate enriched narrative for the master report
    master_ctx = (
        f"Idea: {idea}\nQuadrant: {quadrant}\nIPI: {ipi}/10\nRecommendation: {synthesis.get('recommendation','')}\n\n"
        f"STAGE 02 — Market ({weights.get('market',35)}%): {scores.get('market',5)}/10\n"
        f"Market: {market.get('market_name','')}  Size 2024: {_mval(market.get('market_size_current') or market.get('market_size_2024',''))}  CAGR: {_mval(market.get('cagr',''))}\n"
        f"Primary sectors: {', '.join(s2d.get('sectors',{}).get('primary_sectors',[]))}\n"
        f"Competition: {comp.get('competitive_intensity','')}  White space: {comp.get('white_space','')}\n\n"
        f"STAGE 03 — Patent ({weights.get('patent',25)}%): {scores.get('patent',5)}/10\n"
        f"Activity: {landscape.get('activity_level','')}  Trend: {landscape.get('filing_trend','')}\n"
        f"Novelty: {ansoff_d.get('novelty_signal','')}  IP risk: {ansoff_d.get('ip_risk','')}\n"
        f"White spaces: {'; '.join(landscape.get('white_spaces',[])[:3])}\n\n"
        f"STAGE 04 — Feasibility ({weights.get('feasibility',25)}%): {scores.get('feasibility',5)}/10\n"
        f"TRL: {trl.get('trl_level',3)} — {trl.get('trl_label','')}\n"
        f"Existence: {existence.get('existence_verdict','')}  Entry readiness: {trl.get('schaeffler_entry_readiness','')}\n"
        f"Time to production: {existence.get('time_to_readiness','')}\n\n"
        f"STAGE 05 — P³ Score ({weights.get('org',15)}%): {scores.get('p3',5)}/10\n"
        f"Portfolio: {s5d_org.get('p_portfolio',5)}/10  People: {s5d_org.get('p_people',5)}/10  Process: {s5d_org.get('p_process',5)}/10\n"
        f"Strategy: {bop.get('recommendation','')}  Time with partner: {bop.get('time_to_trl6_partner','')}\n"
        f"Critical gap: {people.get('competency_gap','')}\n\n"
        f"Strongest signals: {'; '.join(synthesis.get('strongest_signals',[])[:3])}\n"
        f"Key concerns: {'; '.join(synthesis.get('key_concerns',[])[:3])}\n"
        f"Strategic fit: {synthesis.get('strategic_fit','')}"
    )
    enriched = call_claude(
        '''You are a senior Schaeffler innovation strategist writing a comprehensive multi-stage Innovation Assessment Report.
Write rich, specific, analytical content. Return ONLY valid JSON, no markdown backticks:
{
  "executive_summary": "4-5 full paragraphs: the headline verdict, what the data shows across all 4 dimensions (market, IP, feasibility, org), the single most compelling reason to proceed, the single most important risk, and the concrete next step.",
  "strategic_narrative": "4-5 full paragraphs synthesising the opportunity holistically — how the market signals, IP landscape, technology readiness, and org capability interact. Reference Schaeffler electrification strategy, Vitesco merger, E-Mobility growth, OEM relationships and P3 formula.",
  "market_highlights": "2 full paragraphs on the top 3 market signals — most important size/growth data point, the most relevant sector fit, and the most significant competitive dynamic.",
  "ip_highlights": "2 full paragraphs on the top IP insights — novelty position, the most threatening filer, and the most valuable white space to capture.",
  "feasibility_highlights": "2 full paragraphs on the technology readiness picture — TRL rationale, most convincing evidence found, and the one critical gap that needs solving before TRL6.",
  "org_highlights": "2 full paragraphs on P³ Perspective — the strongest existing capability, the most critical gap, and the recommended build-partner path.",
  "risk_synthesis": ["Top cross-cutting risk 1 with specific mitigation", "Risk 2 with mitigation", "Risk 3 with mitigation", "Risk 4 with mitigation", "Risk 5 with mitigation"],
  "action_plan": ["Immediate action (0-30 days): specific step 1", "Short-term (1-3 months): specific step 2", "Medium-term (3-6 months): specific step 3", "Longer-term (6-12 months): specific step 4", "Strategic (12+ months): specific step 5"]
}''',
        master_ctx, max_tokens=4000
    )
    raw_e = enriched.strip().replace("```json","").replace("```","").strip()
    fb = raw_e.find("{"); lb = raw_e.rfind("}") + 1
    if fb >= 0: raw_e = raw_e[fb:lb]
    try:
        enr = json.loads(raw_e)
    except:
        enr = {
            "executive_summary": f"This innovation idea ({idea[:80]}) has been assessed across four dimensions of Schaeffler\'s Innovation Pipeline, yielding an Innovation Potential Index (IPI) of {ipi}/10. The recommendation is: {synthesis.get('recommendation','PROCEED WITH CONDITIONS')}. {synthesis.get('recommendation_rationale','')} The strongest signal is: {synthesis.get('strongest_signals',['market opportunity'])[0] if synthesis.get('strongest_signals') else 'market opportunity'}. The primary concern is: {synthesis.get('key_concerns',['technical maturity'])[0] if synthesis.get('key_concerns') else 'technical maturity'}.",
            "strategic_narrative": synthesis.get("narrative", f"The idea targets {market.get('market_name','a growing market')} and sits in Schaeffler\'s {quadrant} innovation quadrant. {synthesis.get('strategic_fit','')}"),
            "market_highlights": f"Market size: {_mval(market.get('market_size_current') or market.get('market_size_2024',''))} (2024), growing to {_mval(market.get('market_size_forecast') or market.get('market_size_2030',''))} by 2030 at {_mval(market.get('cagr',''))} CAGR. Primary sector fit: {', '.join(s2d.get('sectors',{}).get('primary_sectors',[]))}. Competitive intensity: {comp.get('competitive_intensity','')}.",
            "ip_highlights": f"Novelty signal: {ansoff_d.get('novelty_signal','Moderate')}. IP risk: {ansoff_d.get('ip_risk','Medium')}. White spaces: {'; '.join(landscape.get('white_spaces',[])[:2])}.",
            "feasibility_highlights": f"TRL {trl.get('trl_level',3)}: {trl.get('trl_label','')}. Existence: {existence.get('existence_verdict','')}. Time to production: {existence.get('time_to_readiness','')}.",
            "org_highlights": f"Org readiness: {scores.get('p3',5)}/10. Strategy: {bop.get('recommendation','Co-develop')}. Critical gap: {people.get('competency_gap','')}.",
            "risk_synthesis": synthesis.get("risks", ["IP risk requires FTO analysis", "Technical maturity requires R&D investment", "Competitive dynamics require monitoring", "Org readiness gaps require targeted hiring/partnering"]),
            "action_plan": synthesis.get("next_steps", ["Commission engineering feasibility review", "Conduct FTO IP analysis", "Identify pilot customer", "Present to innovation steering committee"])
        }

    NAVY=RGBColor(0x1F,0x38,0x64); BLUE=RGBColor(0x2E,0x75,0xB6)
    WHITE=RGBColor(0xFF,0xFF,0xFF); GREY=RGBColor(0x55,0x55,0x55)
    BLACK=RGBColor(0x00,0x00,0x00); LBLUE=RGBColor(0x60,0xA5,0xFA)
    GREEN=RGBColor(0x22,0xC5,0x5E); AMBER=RGBColor(0xF5,0x9E,0x0B)

    def set_bg(cell, hx):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); tcPr.append(shd)
    def h1(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8"); bot.set(qn("w:space"),"3"); bot.set(qn("w:color"),"2E75B6")
        pBdr.append(bot); pPr.append(pBdr); r=p.add_run(_xs(text))
        r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY
    def h2(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(_xs(text)); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=NAVY
    def body(doc, text):
        if not text: return
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; r=p.add_run(_xs(text)); r.font.size=Pt(10.5)
    def kv(doc, label, value):
        if not value: return
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r1=p.add_run(_xs(f"{label}: ")); r1.bold=True; r1.font.size=Pt(10.5); r1.font.color.rgb=NAVY
        r2=p.add_run(_xs(value)); r2.font.size=Pt(10.5)
    def bul(doc, text):
        if not text: return
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(_xs(text)); r.font.size=Pt(10.5)

    doc=DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # ── Cover header ──────────────────────────────────────────────
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; c=t.cell(0,0); set_bg(c,"1F3864")
    p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("SCHAEFFLER INNOVATION ASSESSMENT REPORT"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=WHITE
    p2=c.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(14)
    r2=p2.add_run("AI Innovation Research Assistant  ·  Full Pipeline Assessment  ·  Stages 01–06")
    r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    doc.add_paragraph()

    # ── Title ──────────────────────────────────────────────────────
    p=doc.add_paragraph(); r=p.add_run(_xs(market.get("market_name", idea[:80]) or idea[:80]))
    r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
    p2=doc.add_paragraph()
    rec_text = synthesis.get("recommendation","")
    r2=p2.add_run(_xs(f"IPI: {ipi}/10  ·  {rec_text}  ·  Quadrant: {quadrant}  ·  {datetime.now().strftime('%d %B %Y')}"))
    r2.font.size=Pt(10); r2.italic=True; r2.font.color.rgb=GREY

    # ── Idea box ──────────────────────────────────────────────────
    doc.add_paragraph()
    tb=doc.add_table(rows=1,cols=2); tb.style="Table Grid"
    c1=tb.cell(0,0); c2=tb.cell(0,1); set_bg(c1,"1F3864"); set_bg(c2,"EAF1FB"); c1.width=Inches(0.12)
    c1.paragraphs[0].add_run("")
    rp=c2.paragraphs[0]; rp.paragraph_format.space_before=Pt(8); rp.paragraph_format.space_after=Pt(2)
    rb=rp.add_run("Innovation Idea"); rb.bold=True; rb.font.size=Pt(9); rb.font.color.rgb=NAVY
    rp2=c2.add_paragraph(); rp2.paragraph_format.space_before=Pt(0); rp2.paragraph_format.space_after=Pt(8)
    ri=rp2.add_run(_xs(idea)); ri.font.size=Pt(10); ri.italic=True
    doc.add_paragraph()

    # ── IPI Score table ───────────────────────────────────────────
    h1(doc,"Innovation Potential Index (IPI)")
    ipi_tbl=doc.add_table(rows=1,cols=4); ipi_tbl.style="Table Grid"
    for i,h in enumerate(["Stage","Score","Weight","Contribution"]):
        c=ipi_tbl.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    for i,(stage,score,wt,contrib) in enumerate([
        ("02 · Market Intelligence",    f"{scores.get('market',5):.1f}/10",      f"{weights.get('market',35)}%",      f"{scores.get('market',5)*weights.get('market',35)/100:.2f}"),
        ("03 · Patent Intelligence",    f"{scores.get('patent',5):.1f}/10",      f"{weights.get('patent',25)}%",      f"{scores.get('patent',5)*weights.get('patent',25)/100:.2f}"),
        ("04 · Technical Feasibility",  f"{scores.get('feasibility',5):.1f}/10", f"{weights.get('feasibility',25)}%", f"{scores.get('feasibility',5)*weights.get('feasibility',25)/100:.2f}"),
        ("05 · P³ Perspective",f"{scores.get('p3',5):.1f}/10",        f"{weights.get('org',15)}%",         f"{scores.get('p3',5)*weights.get('org',15)/100:.2f}"),
    ]):
        row=ipi_tbl.add_row(); fill="EAF1FB" if i%2==0 else "FFFFFF"
        for c in row.cells: set_bg(c,fill)
        for j,val in enumerate([stage,score,wt,contrib]):
            r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(10); r.bold=(j==0)
    fr=ipi_tbl.add_row()
    for c in fr.cells: set_bg(c,"1F3864")
    r=fr.cells[0].paragraphs[0].add_run("INNOVATION POTENTIAL INDEX"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    r2=fr.cells[3].paragraphs[0].add_run(_xs(f"{ipi:.1f}/10")); r2.bold=True; r2.font.size=Pt(13); r2.font.color.rgb=LBLUE
    doc.add_paragraph()

    # ── Recommendation panel ──────────────────────────────────────
    h1(doc,"Recommendation")
    rec_col_map = {"PROCEED":"EAF9F0","PROCEED WITH CONDITIONS":"FFF8E4","DEFER":"FFF0E4","REJECT":"FFE9E9"}
    rec_fill = rec_col_map.get(rec_text, "EAF1FB").lstrip("#")
    rec_tbl=doc.add_table(rows=1,cols=1); rec_tbl.style="Table Grid"
    rc=rec_tbl.cell(0,0); set_bg(rc, rec_fill)
    p=rc.paragraphs[0]; p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(_xs(rec_text)); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY
    p2=rc.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(8)
    r2=p2.add_run(_xs(synthesis.get("recommendation_rationale",""))); r2.font.size=Pt(10)
    doc.add_paragraph()

    if synthesis.get("strongest_signals"):
        h2(doc,"Strongest Signals")
        for s in synthesis["strongest_signals"]: bul(doc, f"✓  {s}")
    if synthesis.get("key_concerns"):
        h2(doc,"Key Concerns")
        for c in synthesis["key_concerns"]: bul(doc, f"⚠  {c}")
    if synthesis.get("conditions"):
        h2(doc,"Conditions to Proceed")
        for c in synthesis["conditions"]: bul(doc, f"→  {c}")

    # ── Executive Summary ─────────────────────────────────────────
    h1(doc,"Executive Summary")
    body(doc, enr.get("executive_summary",""))

    # ── Strategic Narrative ───────────────────────────────────────
    h1(doc,"Strategic Narrative & Full Synthesis")
    body(doc, enr.get("strategic_narrative",""))
    if synthesis.get("strategic_fit"):
        doc.add_paragraph()
        body(doc, synthesis.get("strategic_fit",""))

    # ── Stage 02 — Market Intelligence ───────────────────────────
    h1(doc,"Stage 02 · Market Intelligence  ·  Score: " + str(scores.get("market",5)) + "/10")
    body(doc, enr.get("market_highlights",""))
    doc.add_paragraph()
    kv(doc,"Market", market.get("market_name",""))
    kv(doc,"Size (2024)", _mval(market.get("market_size_current") or market.get("market_size_2024","")))
    kv(doc,"Projected (2030)", _mval(market.get("market_size_forecast") or market.get("market_size_2030","")))
    kv(doc,"CAGR", market.get("cagr",""))
    kv(doc,"Maturity", market.get("market_maturity",""))
    kv(doc,"Geography", market.get("geographic_focus",""))
    kv(doc,"Competitive intensity", comp.get("competitive_intensity",""))
    kv(doc,"White space", comp.get("white_space",""))
    kv(doc,"Schaeffler advantage", comp.get("schaeffler_advantage",""))
    if market.get("growth_drivers"):
        h2(doc,"Growth Drivers")
        for d in market["growth_drivers"][:4]: bul(doc, d)
    if comp.get("competitors"):
        h2(doc,"Key Competitors")
        ct=doc.add_table(rows=1,cols=3); ct.style="Table Grid"
        for i,h in enumerate(["Company","Type","Relevance"]):
            c=ct.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,ci in enumerate(comp["competitors"]):
            row=ct.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([ci.get("name",""),ci.get("type",""),ci.get("relevance","")+" "+ci.get("source","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)
    if sectors.get("sector_scores"):
        h2(doc,"Schaeffler Sector Cluster Scores")
        primary=sectors.get("primary_sectors",[])
        st2=doc.add_table(rows=1,cols=3); st2.style="Table Grid"
        for i,h in enumerate(["Sector","Score","Rationale"]):
            c=st2.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,(sec,data) in enumerate(sectors["sector_scores"].items()):
            row=st2.add_row(); fill="EAF5EA" if sec in primary else ("EAF1FB" if idx%2==0 else "FFFFFF")
            for c in row.cells: set_bg(c,fill)
            r0=row.cells[0].paragraphs[0].add_run(_xs(sec)); r0.font.size=Pt(10); r0.bold=(sec in primary)
            r1=row.cells[1].paragraphs[0].add_run(_xs(f"{data.get('score',0)}/10")); r1.font.size=Pt(10); r1.bold=True
            r2=row.cells[2].paragraphs[0].add_run(_xs(data.get("rationale",""))); r2.font.size=Pt(9.5)

    # ── Stage 03 — Patent Intelligence ───────────────────────────
    h1(doc,"Stage 03 · Patent Intelligence  ·  Score: " + str(scores.get("patent",5)) + "/10")
    body(doc, enr.get("ip_highlights",""))
    doc.add_paragraph()
    kv(doc,"Filing activity", landscape.get("activity_level",""))
    kv(doc,"Filing trend", landscape.get("filing_trend",""))
    kv(doc,"Trend rationale", landscape.get("filing_trend_rationale",""))
    kv(doc,"Novelty signal", ansoff_d.get("novelty_signal",""))
    kv(doc,"Novelty rationale", ansoff_d.get("novelty_rationale",""))
    kv(doc,"IP risk", ansoff_d.get("ip_risk",""))
    kv(doc,"IP risk rationale", ansoff_d.get("ip_risk_rationale",""))
    sp=ansoff_d.get("schaeffler_position",{})
    kv(doc,"Schaeffler existing IP", sp.get("existing_ip",""))
    kv(doc,"IP gap addressed", sp.get("gap",""))
    if landscape.get("white_spaces"):
        h2(doc,"IP White Spaces")
        for ws in landscape["white_spaces"]: bul(doc, ws)
    if landscape.get("technology_keywords"):
        h2(doc,"Key Technology Search Terms")
        body(doc, "  ·  ".join(landscape["technology_keywords"]))
    if landscape.get("key_filers"):
        h2(doc,"Key Patent Filers")
        ft=doc.add_table(rows=1,cols=4); ft.style="Table Grid"
        for i,h in enumerate(["Company","Type","Threat","Filing Focus"]):
            c=ft.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,fi in enumerate(landscape["key_filers"]):
            row=ft.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([fi.get("company",""),fi.get("type",""),fi.get("threat_level",""),fi.get("focus","")]):
                r=row.cells[j].paragraphs[0].add_run(val); r.font.size=Pt(9.5); r.bold=(j==0)
        if landscape.get("landscape_summary"):
            doc.add_paragraph(); body(doc, landscape["landscape_summary"])

    # ── Stage 04 — Technical Feasibility ─────────────────────────
    h1(doc,"Stage 04 · Technical Feasibility  ·  Score: " + str(scores.get("feasibility",5)) + "/10")
    body(doc, enr.get("feasibility_highlights",""))
    doc.add_paragraph()
    kv(doc,"TRL Level", trl.get("trl_label",""))
    kv(doc,"Existence verdict", existence.get("existence_verdict",""))
    kv(doc,"Schaeffler entry readiness", trl.get("schaeffler_entry_readiness",""))
    kv(doc,"Time to production readiness", existence.get("time_to_readiness",""))
    kv(doc,"Technology core", existence.get("technology_core",""))
    if trl.get("trl_rationale"):
        doc.add_paragraph(); body(doc, trl["trl_rationale"])
    if existence.get("existence_summary"):
        body(doc, existence["existence_summary"])
    if existence.get("evidence"):
        h2(doc,"Evidence (Academic, Startups, Pilots, Programmes)")
        ev_tbl=doc.add_table(rows=1,cols=4); ev_tbl.style="Table Grid"
        for i,h in enumerate(["Type","Title & Description","Relevance","Source"]):
            c=ev_tbl.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,ev in enumerate(existence["evidence"]):
            row=ev_tbl.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([ev.get("type",""), ev.get("title","")+" — "+ev.get("description",""), ev.get("relevance","")+" / "+ev.get("confidence",""), ev.get("source","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==1)
    if existence.get("technology_gaps"):
        h2(doc,"Technology Gaps to Bridge")
        for gap in existence["technology_gaps"]: bul(doc, gap)
    if existence.get("keywords"):
        h2(doc,"Technology Keyword Map")
        body(doc, "  ·  ".join(existence["keywords"]))
    if trl.get("key_technical_risks"):
        h2(doc,"Key Technical Risks")
        rt=doc.add_table(rows=1,cols=3); rt.style="Table Grid"
        for i,h in enumerate(["Risk","Severity","Mitigation"]):
            c=rt.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,risk in enumerate(trl["key_technical_risks"]):
            row=rt.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([risk.get("risk",""),risk.get("severity",""),risk.get("mitigation","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)

    # ── Stage 05 — P³ Perspective ──────────────────────
    h1(doc,"Stage 05 · P³ Perspective  ·  Score: " + str(scores.get("p3",5)) + "/10")
    body(doc, enr.get("org_highlights",""))
    doc.add_paragraph()
    kv(doc,"P³ Portfolio", f"{s5d_org.get('p_portfolio',5)}/10")
    kv(doc,"P³ People", f"{s5d_org.get('p_people',5)}/10")
    kv(doc,"P³ Process", f"{s5d_org.get('p_process',5)}/10")
    kv(doc,"Build strategy", bop.get("recommendation",""))
    kv(doc,"Time to TRL6 — Internal", bop.get("time_to_trl6_internal",""))
    kv(doc,"Time to TRL6 — With Partner", bop.get("time_to_trl6_partner",""))
    kv(doc,"Critical competency gap", people.get("competency_gap",""))
    kv(doc,"Closure route", people.get("sourcing_route",""))
    if portfolio.get("rationale"): body(doc, portfolio["rationale"])
    if people.get("rationale"): body(doc, people["rationale"])
    if process.get("rationale"): body(doc, process["rationale"])
    if process.get("applicable_assets"):
        h2(doc,"Applicable Schaeffler Assets")
        for a in process["applicable_assets"]: bul(doc, a)
    if people.get("matched_competencies"):
        h2(doc,"Matched Competencies")
        for c in people["matched_competencies"]: bul(doc, f"✓  {c}")
    if org_data.get("org_gaps"):
        h2(doc,"Organisational Gaps")
        gt=doc.add_table(rows=1,cols=4); gt.style="Table Grid"
        for i,h in enumerate(["Gap","Severity","Closure Route","Timeline"]):
            c=gt.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,g in enumerate(org_data["org_gaps"]):
            row=gt.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([g.get("gap",""),g.get("severity",""),g.get("closure_route",""),g.get("timeline","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)
    if org_data.get("partnership_candidates"):
        h2(doc,"Partnership Candidates")
        pt=doc.add_table(rows=1,cols=4); pt.style="Table Grid"
        for i,h in enumerate(["Organisation","Type","Rationale","Route"]):
            c=pt.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,p in enumerate(org_data["partnership_candidates"]):
            row=pt.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([p.get("name",""),p.get("type",""),p.get("rationale",""),p.get("route","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)

    # ── Integrated Risk Register ──────────────────────────────────
    h1(doc,"Integrated Risk Register")
    if enr.get("risk_synthesis"):
        rt2=doc.add_table(rows=1,cols=2); rt2.style="Table Grid"
        for i,h in enumerate(["Risk","Description & Mitigation"]):
            c=rt2.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,risk in enumerate(enr["risk_synthesis"]):
            row=rt2.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            parts=str(risk).split(" — ",1) if " — " in str(risk) else [str(risk),"See full analysis"]
            r0=row.cells[0].paragraphs[0].add_run(_xs(parts[0])); r0.font.size=Pt(9.5); r0.bold=True
            r1=row.cells[1].paragraphs[0].add_run(_xs(parts[1] if len(parts)>1 else "")); r1.font.size=Pt(9.5)

    # ── Action Plan ───────────────────────────────────────────────
    h1(doc,"Recommended Action Plan")
    if enr.get("action_plan"):
        for i, step in enumerate(enr["action_plan"], 1):
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
            r=p.add_run(_xs(f"{i}.  {step}")); r.font.size=Pt(10.5)
    if synthesis.get("next_steps"):
        h2(doc,"Additional Next Steps")
        for step in synthesis["next_steps"]: bul(doc, step)

    # ── Quadrant classification summary ──────────────────────────
    h1(doc,"Stage 01 · Innovation Classification")
    kv(doc,"Quadrant", quadrant)
    kv(doc,"Technology level", s1c.get("technology_level",""))
    kv(doc,"Market level", s1c.get("market_level",""))
    kv(doc,"Technology novelty", s1c.get("technology_novelty",""))
    kv(doc,"Market position", s1c.get("market_position",""))
    kv(doc,"Innovation cluster", s1c.get("innovation_cluster",""))
    kv(doc,"Product family", s1c.get("product_family",""))
    if s1c.get("trend_alignment"):
        kv(doc,"Trend alignment", ", ".join(s1c.get("trend_alignment",[])))
    kv(doc,"Project type", s1c.get("project_type",""))
    kv(doc,"Innovation model", s1c.get("innovation_model",""))
    kv(doc,"Reasoning", s1c.get("reasoning",""))

    # ── Footer ────────────────────────────────────────────────────
    doc.add_paragraph()
    ft=doc.add_table(rows=1,cols=1); ft.style="Table Grid"; fc=ft.cell(0,0); set_bg(fc,"1F3864")
    fp=fc.paragraphs[0]; fp.paragraph_format.space_before=Pt(6); fp.paragraph_format.space_after=Pt(6)
    fr=fp.add_run(_xs(f"Schaeffler AI Innovation Research Assistant  ·  Full Innovation Assessment  ·  {datetime.now().strftime('%d %B %Y')}  ·  Capstone Project — Arpan Chowdhury, EBS Universität"))
    fr.font.size=Pt(8); fr.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    _sanitize_doc(doc)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf
def generate_feasibility_report(idea, quadrant, s1c, existence, trl, scores):
    """Generate a Technical Feasibility Word report."""
    feas_ctx = (
        f"Idea: {idea}\nQuadrant: {quadrant}\n"
        f"TRL level: {trl.get('trl_level',3)} — {trl.get('trl_label','')}\n"
        f"TRL rationale: {trl.get('trl_rationale','')}\n"
        f"Existence verdict: {existence.get('existence_verdict','')}\n"
        f"Existence summary: {existence.get('existence_summary','')}\n"
        f"Schaeffler entry readiness: {trl.get('schaeffler_entry_readiness','')}\n"
        f"Time to production readiness: {existence.get('time_to_readiness','')}\n"
        f"Technology core: {existence.get('technology_core','')}\n"
        f"Technology gaps: {chr(59).join(existence.get('technology_gaps',[])[:4])}\n"
        f"Key evidence: {chr(59).join([e.get('title','') + ' (' + e.get('source','') + ')' for e in existence.get('evidence',[])[:4]])}\n"
        f"Analogous Schaeffler tech: {trl.get('analogous_schaeffler_technologies','') or trl.get('analogous_schaeffler_tech','')}\n"
        f"Key technical risks: {chr(59).join([r.get('risk','') for r in trl.get('key_technical_risks',[])[:3]])}\n"
        f"Score: {scores['final_score']}/10"
    )
    extended = call_claude(
        '''You are a Schaeffler R&D director writing a detailed technical feasibility assessment.
Write specific, evidence-based content. Return ONLY valid JSON, no markdown backticks:
{
  "executive_summary": "3-4 full paragraphs: overall feasibility verdict, TRL assessment rationale, key evidence found, and what Schaeffler needs to do to advance the technology.",
  "technology_analysis": "3-4 full paragraphs on the core technology mechanism, current state of the art globally, what has been demonstrated vs what remains theoretical, and the critical technical gaps to bridge.",
  "schaeffler_readiness": "2-3 full paragraphs on Schaeffler-specific technical readiness — which existing competencies (bearings, mechatronics, Vitesco power electronics, tribology) are applicable, what is genuinely missing, and how analogous it is to existing Schaeffler product lines.",
  "development_pathway": "2-3 full paragraphs: specific recommended development pathway from current TRL to TRL 6 and beyond — what milestones, what resources, what partnerships, and realistic timelines.",
  "risks": ["Technical risk 1 with specific mitigation approach", "Technical risk 2 with mitigation", "Technical risk 3", "Technical risk 4"],
  "recommendations": ["Concrete R&D action 1 with timeline", "Concrete action 2", "Concrete action 3", "Concrete action 4"]
}''',
        feas_ctx,
        max_tokens=3500
    )
    raw_e = extended.strip().replace("```json","").replace("```","").strip()
    fb = raw_e.find("{"); lb = raw_e.rfind("}") + 1
    if fb >= 0: raw_e = raw_e[fb:lb]
    try:
        ext = json.loads(raw_e)
    except:
        ext = {
            "executive_summary": f"This innovation idea is assessed at TRL {trl.get('trl_level',3)} — {trl.get('trl_label','')}. Existence verdict: {existence.get('existence_verdict','Research Stage')}. {existence.get('existence_summary','')} Schaeffler entry readiness: {trl.get('schaeffler_entry_readiness','Ready for Innovation')}. Estimated time to production readiness: {existence.get('time_to_readiness','3-5 years')}.",
            "technology_analysis": f"Core technology: {existence.get('technology_core','')}. TRL rationale: {trl.get('trl_rationale','')} Technology gaps to bridge: {chr(44).join(existence.get('technology_gaps',[])[:3])}.",
            "schaeffler_readiness": f"Analogous Schaeffler technology: {trl.get('analogous_schaeffler_technologies','') or trl.get('analogous_schaeffler_tech','')}. {trl.get('entry_rationale','')}",
            "development_pathway": f"From current TRL {trl.get('trl_level',3)}, the recommended development pathway progresses through lab validation, prototype testing in relevant environment, and pilot deployment. Estimated timeline: {existence.get('time_to_readiness','3-5 years')}.",
            "risks": [r.get('risk','') + ' — ' + r.get('mitigation','') for r in trl.get('key_technical_risks',[])[:4]] or ["Technology gaps require systematic R&D investment", "Manufacturing scalability not yet demonstrated", "Integration with existing Schaeffler systems needs validation"],
            "recommendations": ["Commission proof-of-concept study with Schaeffler R&D within 60 days", f"Target TRL {min(trl.get('trl_level',3)+2, 6)} within 18 months via innovation project", "Identify university or research institute partner for early-stage development", "Conduct internal engineering workshop to assess competency gaps"]
        }

    NAVY=RGBColor(0x1F,0x38,0x64); WHITE=RGBColor(0xFF,0xFF,0xFF)
    GREY=RGBColor(0x55,0x55,0x55); BLACK=RGBColor(0x00,0x00,0x00)
    LBLUE=RGBColor(0x60,0xA5,0xFA)

    def set_bg(cell, hx):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); tcPr.append(shd)

    def h1(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8"); bot.set(qn("w:space"),"3"); bot.set(qn("w:color"),"2E75B6")
        pBdr.append(bot); pPr.append(pBdr); r=p.add_run(_xs(text))
        r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY

    def body(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; r=p.add_run(_xs(text)); r.font.size=Pt(10.5)

    def kv(doc, label, value):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r1=p.add_run(_xs(f"{label}: ")); r1.bold=True; r1.font.size=Pt(10.5); r1.font.color.rgb=NAVY
        r2=p.add_run(_xs(value)); r2.font.size=Pt(10.5)

    def bul(doc, text):
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(_xs(text)); r.font.size=Pt(10.5)

    doc=DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # Header
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; c=t.cell(0,0); set_bg(c,"1F3864")
    p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("TECHNICAL FEASIBILITY REPORT"); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=WHITE
    p2=c.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(10)
    r2=p2.add_run("Schaeffler AI Innovation Research Assistant  ·  Stage 04"); r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    doc.add_paragraph()

    # Title
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("Technical Feasibility Assessment"); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
    p2=doc.add_paragraph()
    r2=p2.add_run(_xs(f"Score: {scores['final_score']}/10  ·  TRL {trl.get('trl_level','')}  ·  {datetime.now().strftime('%d %B %Y')}"))
    r2.font.size=Pt(10); r2.italic=True; r2.font.color.rgb=GREY

    # Idea box
    doc.add_paragraph()
    tb=doc.add_table(rows=1,cols=2); tb.style="Table Grid"
    c1=tb.cell(0,0); c2=tb.cell(0,1); set_bg(c1,"1F3864"); set_bg(c2,"EAF1FB"); c1.width=Inches(0.12)
    c1.paragraphs[0].add_run("")
    rp=c2.paragraphs[0]; rp.paragraph_format.space_before=Pt(8); rp.paragraph_format.space_after=Pt(2)
    rb=rp.add_run("Innovation Idea"); rb.bold=True; rb.font.size=Pt(9); rb.font.color.rgb=NAVY
    rp2=c2.add_paragraph(); rp2.paragraph_format.space_before=Pt(0); rp2.paragraph_format.space_after=Pt(8)
    ri=rp2.add_run(_xs(idea)); ri.font.size=Pt(10); ri.italic=True
    doc.add_paragraph()

    # Score summary
    h1(doc,"Score Summary")
    st2=doc.add_table(rows=1,cols=3); st2.style="Table Grid"
    for i,h in enumerate(["Dimension","Score","Weight"]):
        c=st2.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    for i,(dim,score,wt) in enumerate([
        ("TRL Score",f"{scores['trl_score']:.1f}/10","33%"),
        ("Existence Quality",f"{scores['existence_score']:.1f}/10","33%"),
        ("Risk Profile",f"{scores['risk_score']:.1f}/10","33%"),
    ]):
        row=st2.add_row(); fill="EAF1FB" if i%2==0 else "FFFFFF"
        for c in row.cells: set_bg(c,fill)
        for j,val in enumerate([dim,score,wt]):
            r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(10); r.bold=(j==0)
    fr=st2.add_row()
    for c in fr.cells: set_bg(c,"1F3864")
    r=fr.cells[0].paragraphs[0].add_run("FINAL SCORE"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    r2=fr.cells[2].paragraphs[0].add_run(_xs(f"{scores['final_score']}/10")); r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=LBLUE
    doc.add_paragraph()

    # Executive summary
    h1(doc,"Executive Summary"); body(doc,ext.get("executive_summary",""))

    # TRL assessment
    h1(doc,"Technology Readiness Level Assessment")
    kv(doc,"TRL Level",trl.get("trl_label",""))
    kv(doc,"Existence verdict",existence.get("existence_verdict",""))
    kv(doc,"Schaeffler entry readiness",trl.get("schaeffler_entry_readiness",""))
    kv(doc,"Time to production readiness",existence.get("time_to_readiness",""))
    doc.add_paragraph(); body(doc,trl.get("trl_rationale",""))

    # TRL scale table
    doc.add_paragraph()
    trl_tbl=doc.add_table(rows=1,cols=3); trl_tbl.style="Table Grid"
    for i,h in enumerate(["TRL","Level","Description"]):
        c=trl_tbl.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    trl_rows=[
        (1,"Basic principles observed","Theoretical concept only"),
        (2,"Technology concept formulated","Application identified, no testing"),
        (3,"Experimental proof of concept","Lab demonstration, key functions validated"),
        (4,"Technology validated in lab","Component tested, controlled environment"),
        (5,"Validated in relevant environment","Prototype tested, industrial-like conditions"),
        (6,"Demonstrated in relevant environment","System prototype demonstrated"),
        (7,"Prototype in operational environment","Field trial or industrial pilot"),
        (8,"System complete and qualified","Limited production run"),
        (9,"Proven in operational environment","Commercial deployment at scale"),
    ]
    current_trl = trl.get("trl_level",3)
    for lvl,lbl,desc in trl_rows:
        row=trl_tbl.add_row()
        fill="EAF5EA" if lvl==current_trl else ("EAF1FB" if lvl%2==0 else "FFFFFF")
        for c in row.cells: set_bg(c,fill)
        r0=row.cells[0].paragraphs[0].add_run(f"TRL {lvl}"); r0.font.size=Pt(10); r0.bold=(lvl==current_trl)
        r1=row.cells[1].paragraphs[0].add_run(lbl); r1.font.size=Pt(10); r1.bold=(lvl==current_trl)
        r2=row.cells[2].paragraphs[0].add_run(desc); r2.font.size=Pt(9.5)

    # Technology analysis
    h1(doc,"Technology Analysis"); body(doc,ext.get("technology_analysis",""))

    # Evidence table
    h1(doc,"Existing Evidence")
    body(doc,existence.get("existence_summary",""))
    evidence=existence.get("evidence",[])
    if evidence:
        doc.add_paragraph()
        ev_tbl=doc.add_table(rows=1,cols=4); ev_tbl.style="Table Grid"
        for i,h in enumerate(["Type","Title","Relevance","Source"]):
            c=ev_tbl.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,ev in enumerate(evidence):
            row=ev_tbl.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([ev.get("type",""),ev.get("title",""),ev.get("relevance",""),ev.get("source","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==1)
        # Description column
        doc.add_paragraph()
        for ev in evidence:
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            r1=p.add_run(_xs(f"{ev.get('title','')}: ")); r1.bold=True; r1.font.size=Pt(10)
            r2=p.add_run(_xs(ev.get("description",""))); r2.font.size=Pt(10)

    # Technology gaps
    if existence.get("technology_gaps"):
        h1(doc,"Technology Gaps to Bridge")
        for gap in existence["technology_gaps"]: bul(doc,gap)

    # Technical risks
    h1(doc,"Key Technical Risks")
    risks=trl.get("key_technical_risks",[])
    if risks:
        rt=doc.add_table(rows=1,cols=3); rt.style="Table Grid"
        for i,h in enumerate(["Risk","Severity","Mitigation"]):
            c=rt.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,risk in enumerate(risks):
            row=rt.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([risk.get("risk",""),risk.get("severity",""),risk.get("mitigation","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)

    # Schaeffler readiness & development pathway
    h1(doc,"Schaeffler Readiness"); body(doc,ext.get("schaeffler_readiness",""))
    h1(doc,"Development Pathway"); body(doc,ext.get("development_pathway",""))
    h1(doc,"Recommendations")
    for rec in ext.get("recommendations",[]): bul(doc,rec)

    # Footer
    doc.add_paragraph()
    ft=doc.add_table(rows=1,cols=1); ft.style="Table Grid"; fc=ft.cell(0,0); set_bg(fc,"1F3864")
    fp=fc.paragraphs[0]; fp.paragraph_format.space_before=Pt(6); fp.paragraph_format.space_after=Pt(6)
    fr=fp.add_run(_xs(f"Schaeffler AI Innovation Research Assistant  ·  Stage 04: Technical Feasibility  ·  {datetime.now().strftime('%d %B %Y')}  ·  Capstone Project — Arpan Chowdhury, EBS Universität"))
    fr.font.size=Pt(8); fr.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    _sanitize_doc(doc)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


def generate_patent_report(idea, quadrant, s1c, landscape, ansoff_data, scores):
    """Generate a Patent Intelligence Word report."""
    pat_ctx = (
        f"Idea: {idea}\nQuadrant: {quadrant}\n"
        f"Activity level: {landscape.get('activity_level','')}  Trend: {landscape.get('filing_trend','')}\n"
        f"Filing trend rationale: {landscape.get('filing_trend_rationale','')}\n"
        f"Novelty signal: {ansoff_data.get('novelty_signal','')}  IP risk: {ansoff_data.get('ip_risk','')}\n"
        f"Novelty rationale: {ansoff_data.get('novelty_rationale','')}\n"
        f"IP risk rationale: {ansoff_data.get('ip_risk_rationale','')}\n"
        f"Key filers: {chr(44).join([f.get('company','') for f in landscape.get('key_filers',[])[:6]])}\n"
        f"White spaces: {chr(59).join(landscape.get('white_spaces',[]))}\n"
        f"Schaeffler existing IP: {ansoff_data.get('schaeffler_position',{}).get('existing_ip','')}\n"
        f"IP gap addressed: {ansoff_data.get('schaeffler_position',{}).get('gap','')}\n"
        f"Landscape summary: {landscape.get('landscape_summary','')}\nScore: {scores['final_score']}/10"
    )
    extended = call_claude(
        '''You are a senior patent intelligence analyst writing a detailed IP report for Schaeffler Group.
Write specific, actionable content based on the data provided. Return ONLY valid JSON, no markdown backticks:
{
  "executive_summary": "3-4 full paragraphs: overall IP landscape assessment, novelty of the idea, key risks and opportunities, and headline recommendation for Schaeffler IP strategy.",
  "landscape_analysis": "3-4 full paragraphs covering filing activity levels, who is filing and why, filing trends and what they indicate, which quadrant of IP activity is most active, and what it means for new entrants.",
  "ip_strategy": "2-3 full paragraphs: specific recommended IP strategy for Schaeffler — whether to file broadly or narrowly, which white spaces to target, which claims to prioritise, and how to defend against existing filers.",
  "risks": ["Specific IP risk 1 with mitigation strategy", "Specific risk 2 with action", "Specific risk 3 with action", "Specific risk 4"],
  "recommendations": ["Concrete IP action 1 with timeline", "Concrete action 2", "Concrete action 3", "Concrete action 4"]
}''',
        pat_ctx,
        max_tokens=3500
    )
    raw_e = extended.strip().replace("```json","").replace("```","").strip()
    fb = raw_e.find("{"); lb = raw_e.rfind("}") + 1
    if fb >= 0: raw_e = raw_e[fb:lb]
    try:
        ext = json.loads(raw_e)
    except:
        sp = ansoff_data.get('schaeffler_position',{})
        ext = {
            "executive_summary": f"The patent landscape for this innovation shows {landscape.get('activity_level','moderate')} activity with a {landscape.get('filing_trend','stable')} trend. Novelty signal is {ansoff_data.get('novelty_signal','Moderate')} and IP risk is assessed as {ansoff_data.get('ip_risk','Medium')}. {landscape.get('landscape_summary','')} Schaeffler's existing IP: {sp.get('existing_ip','see analysis')}.",
            "landscape_analysis": f"Filing trend is {landscape.get('filing_trend','stable')}. {landscape.get('filing_trend_rationale','')} Key filers include {chr(44).join([f.get('company','') for f in landscape.get('key_filers',[])[:4]])}. White spaces identified: {chr(59).join(landscape.get('white_spaces',[])[:3])}.",
            "ip_strategy": f"Schaeffler should target the identified IP white spaces to establish a defensible position. The gap this idea addresses: {sp.get('gap','see analysis')}. Novelty rationale: {ansoff_data.get('novelty_rationale','')}",
            "risks": [f"IP risk level {ansoff_data.get('ip_risk','Medium')}: {ansoff_data.get('ip_risk_rationale','conduct FTO analysis before R&D commitment')}", "Freedom-to-operate analysis required before external disclosure", "Monitor filing activity of identified key players for blocking patents", "Defensive filing strategy needed to protect white space positions"],
            "recommendations": ["Commission formal FTO (freedom-to-operate) analysis within 60 days", "File provisional patent applications in identified white spaces", "Monitor competitor filing activity via patent watch service", "Engage Schaeffler patent counsel to assess portfolio gap"]
        }

    NAVY  = RGBColor(0x1F,0x38,0x64); BLUE = RGBColor(0x2E,0x75,0xB6)
    WHITE = RGBColor(0xFF,0xFF,0xFF); GREY = RGBColor(0x55,0x55,0x55)
    BLACK = RGBColor(0x00,0x00,0x00); LBLUE= RGBColor(0x60,0xA5,0xFA)

    def set_bg(cell, hex_col):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_col); tcPr.append(shd)

    def h1(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr"); bot=OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8"); bot.set(qn("w:space"),"3"); bot.set(qn("w:color"),"2E75B6")
        pBdr.append(bot); pPr.append(pBdr)
        r=p.add_run(_xs(text)); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY

    def body(doc, text):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        r=p.add_run(text); r.font.size=Pt(10.5)

    def kv(doc, label, value):
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r1=p.add_run(_xs(f"{label}: ")); r1.bold=True; r1.font.size=Pt(10.5); r1.font.color.rgb=NAVY
        r2=p.add_run(_xs(value)); r2.font.size=Pt(10.5)

    def bul(doc, text):
        p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(_xs(text)); r.font.size=Pt(10.5)

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # Header
    t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; c=t.cell(0,0); set_bg(c,"1F3864")
    p=c.paragraphs[0]; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("PATENT INTELLIGENCE REPORT"); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=WHITE
    p2=c.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(10)
    r2=p2.add_run("Schaeffler AI Innovation Research Assistant  ·  Stage 03"); r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    doc.add_paragraph()

    # Title
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run("Patent Intelligence Analysis"); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
    p2=doc.add_paragraph()
    r2=p2.add_run(f"Score: {scores['final_score']}/10  ·  Quadrant: {quadrant}  ·  {datetime.now().strftime('%d %B %Y')}")
    r2.font.size=Pt(10); r2.italic=True; r2.font.color.rgb=GREY

    # Idea box
    doc.add_paragraph()
    tb=doc.add_table(rows=1,cols=2); tb.style="Table Grid"
    c1=tb.cell(0,0); c2=tb.cell(0,1); set_bg(c1,"1F3864"); set_bg(c2,"EAF1FB"); c1.width=Inches(0.12)
    c1.paragraphs[0].add_run("")
    rp=c2.paragraphs[0]; rp.paragraph_format.space_before=Pt(8); rp.paragraph_format.space_after=Pt(2)
    rb=rp.add_run("Innovation Idea"); rb.bold=True; rb.font.size=Pt(9); rb.font.color.rgb=NAVY
    rp2=c2.add_paragraph(); rp2.paragraph_format.space_before=Pt(0); rp2.paragraph_format.space_after=Pt(8)
    ri=rp2.add_run(_xs(idea)); ri.font.size=Pt(10); ri.italic=True
    doc.add_paragraph()

    # Score summary
    h1(doc,"Score Summary")
    st2=doc.add_table(rows=1,cols=3); st2.style="Table Grid"
    for i,h in enumerate(["Dimension","Score","Weight"]):
        c=st2.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    for i,(dim,score,wt) in enumerate([
        ("Landscape Openness",      f"{scores['landscape_score']:.1f}/10",           "25%"),
        ("Novelty Signal",          f"{scores['novelty_score']:.1f}/10",             "25%"),
        ("IP Risk — Idea",          f"{scores.get('ip_idea_score', scores.get('ip_score', 5)):.1f}/10",       "25%"),
        ("IP Risk — Schaeffler",    f"{scores.get('ip_schaeffler_score', scores.get('ip_score', 5)):.1f}/10", "25%"),
    ]):
        row=st2.add_row(); fill="EAF1FB" if i%2==0 else "FFFFFF"
        for j,c in enumerate(row.cells): set_bg(c,fill)
        for j,val in enumerate([dim,score,wt]):
            r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(10); r.bold=(j==0)
    fr=st2.add_row()
    for c in fr.cells: set_bg(c,"1F3864")
    r=fr.cells[0].paragraphs[0].add_run("FINAL PATENT INTELLIGENCE SCORE"); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    r2=fr.cells[2].paragraphs[0].add_run(_xs(f"{scores['final_score']}/10")); r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=LBLUE
    doc.add_paragraph()

    # Executive summary
    h1(doc,"Executive Summary"); body(doc,ext.get("executive_summary",""))

    # Landscape overview
    h1(doc,"Patent Landscape Overview")
    kv(doc,"Activity level",landscape.get("activity_level",""))
    kv(doc,"Filing trend",landscape.get("filing_trend","") + " — " + landscape.get("filing_trend_rationale",""))
    kv(doc,"Novelty signal",ansoff_data.get("novelty_signal","") + " — " + ansoff_data.get("novelty_rationale",""))
    kv(doc,"IP risk",ansoff_data.get("ip_risk","") + " — " + ansoff_data.get("ip_risk_rationale",""))
    if landscape.get("technology_keywords"):
        kv(doc,"Key technology terms",", ".join(landscape["technology_keywords"]))
    doc.add_paragraph(); body(doc,ext.get("landscape_analysis",""))

    # Key filers table
    h1(doc,"Key Patent Filers")
    filers=landscape.get("key_filers",[])
    if filers:
        ft=doc.add_table(rows=1,cols=4); ft.style="Table Grid"
        for i,h in enumerate(["Company","Type","Threat","Focus"]):
            c=ft.cell(0,i); set_bg(c,"1F3864"); r=c.paragraphs[0].add_run(h)
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
        for idx,fi in enumerate(filers):
            row=ft.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([fi.get("company",""),fi.get("type",""),fi.get("threat_level",""),fi.get("focus","")+" "+fi.get("source","")]):
                r=row.cells[j].paragraphs[0].add_run(_xs(val)); r.font.size=Pt(9.5); r.bold=(j==0)

    # Schaeffler IP position
    h1(doc,"Schaeffler IP Position")
    sp=ansoff_data.get("schaeffler_position",{})
    kv(doc,"Existing IP",sp.get("existing_ip",""))
    kv(doc,"IP gap addressed",sp.get("gap",""))
    if landscape.get("white_spaces"):
        doc.add_paragraph()
        p=doc.add_paragraph(); r=p.add_run("IP White Spaces"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=NAVY
        for ws in landscape["white_spaces"]: bul(doc,ws)
    doc.add_paragraph(); body(doc,ext.get("ip_strategy",""))

    # Risks & recommendations
    h1(doc,"IP Risks & Mitigations")
    for risk in ext.get("risks",[]): bul(doc,risk)
    h1(doc,"Recommendations")
    for rec in ext.get("recommendations",[]): bul(doc,rec)

    # Footer
    doc.add_paragraph()
    ft2=doc.add_table(rows=1,cols=1); ft2.style="Table Grid"; fc=ft2.cell(0,0); set_bg(fc,"1F3864")
    fp=fc.paragraphs[0]; fp.paragraph_format.space_before=Pt(6); fp.paragraph_format.space_after=Pt(6)
    fr=fp.add_run(_xs(f"Schaeffler AI Innovation Research Assistant  ·  Stage 03: Patent Intelligence  ·  {datetime.now().strftime('%d %B %Y')}  ·  Capstone Project — Arpan Chowdhury, EBS Universität"))
    fr.font.size=Pt(8); fr.font.color.rgb=RGBColor(0x93,0xC5,0xFD)

    _sanitize_doc(doc)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


def _sanitize_doc(doc):
    """Walk every XML text node and strip XML-illegal control characters."""
    for el in doc.element.iter():
        try:
            if el.text:
                el.text = _XML_CTRL_RE.sub('', el.text)
        except (AttributeError, TypeError):
            pass
        try:
            if el.tail:
                el.tail = _XML_CTRL_RE.sub('', el.tail)
        except (AttributeError, TypeError):
            pass

def generate_market_report(idea, quadrant, s1c, market, comp, sectors, weights, final_score):
    """Generate a formatted Word document market intelligence report."""
    mkt_ctx = (
        f"Idea: {idea}\nQuadrant: {quadrant}\nMarket: {market.get('market_name','')}\n"
        f"Size 2024: {_mval(market.get('market_size_current') or market.get('market_size_2024',''))}  Size 2030: {_mval(market.get('market_size_forecast') or market.get('market_size_2030',''))}\n"
        f"CAGR: {_mval(market.get('cagr',''))}  Maturity: {market.get('market_maturity','')}\n"
        f"Geography: {market.get('geographic_focus','')}\n"
        f"Growth drivers: {chr(59).join(market.get('growth_drivers',[])[:4])}\n"
        f"Competitive intensity: {comp.get('competitive_intensity','')}  White space: {comp.get('white_space','')}\n"
        f"Schaeffler advantage: {comp.get('schaeffler_advantage','')}\n"
        f"Key competitors: {', '.join([c.get('name','') for c in comp.get('competitors',[])[:6]])}\n"
        f"Primary sectors: {', '.join(sectors.get('primary_sectors',[]))}\n"
        f"Sector fit rationale: {sectors.get('sector_fit_rationale','')}\nFinal score: {final_score}/10"
    )
    extended = call_claude(
        '''You are a senior market analyst writing a comprehensive market intelligence report for Schaeffler Group.
Write substantive, specific content using the data provided. Return ONLY valid JSON, no markdown backticks:
{
  "executive_summary": "3-4 full paragraphs summarising market opportunity, competitive position, and strategic fit. Reference exact figures provided.",
  "market_deep_dive": "3-4 full paragraphs covering market dynamics, demand signals, geographic breakdown, growth trajectory and 5-year outlook.",
  "competitive_analysis": "3-4 full paragraphs on key players, their strategies, where white space exists, and how Schaeffler is positioned.",
  "schaeffler_fit": "2-3 full paragraphs on strategic fit — reference Schaeffler electrification agenda, Vitesco merger, E-Mobility growth, and P3 portfolio formula.",
  "risks": ["Specific risk 1 with concrete mitigation", "Specific risk 2 with mitigation", "Specific risk 3 with mitigation", "Specific risk 4"],
  "recommendations": ["Concrete action 1 with timeline", "Concrete action 2", "Concrete action 3", "Concrete action 4"]
}''',
        mkt_ctx,
        max_tokens=3500
    )
    raw_e = extended.strip().replace("```json","").replace("```","").strip()
    fb = raw_e.find("{"); lb = raw_e.rfind("}") + 1
    if fb >= 0: raw_e = raw_e[fb:lb]
    try:
        ext = json.loads(raw_e)
    except:
        ext = {
            "executive_summary": f"The {market.get('market_name','target market')} represents a {market.get('market_maturity','growing')} opportunity. Market size is estimated at {market.get('market_size_2024','significant')} in 2024, projected to reach {market.get('market_size_2030','substantial')} by 2030 at a CAGR of {market.get('cagr','strong growth')}. Competitive intensity is {comp.get('competitive_intensity','moderate')}, with white space identified: {comp.get('white_space','see analysis')}.",
            "market_deep_dive": f"The market is characterised by {market.get('market_maturity','growth')} dynamics. Key growth drivers include: {chr(10).join(market.get('growth_drivers',['strong demand','digital transformation','electrification'])[:3])}. Geographic focus: {market.get('geographic_focus','global')}.",
            "competitive_analysis": f"The competitive landscape shows {comp.get('competitive_intensity','moderate')} intensity. Schaeffler advantage: {comp.get('schaeffler_advantage','precision engineering and OEM relationships')}. White space: {comp.get('white_space','niche segments underserved')}.",
            "schaeffler_fit": f"Primary sector fit is strongest in {chr(44).join(sectors.get('primary_sectors',[]))}, which aligns with Schaeffler's post-Vitesco portfolio strategy. {sectors.get('sector_fit_rationale','See sector analysis for detail.')}",
            "risks": ["Competitive intensity requires freedom-to-operate analysis before R&D commitment — conduct FTO within 60 days", "Market timing risk: validate demand with target OEMs before scaling investment", "Sector fit assumptions require validation with Schaeffler divisional teams", "IP landscape may restrict entry — patent clearance essential"],
            "recommendations": ["Commission internal engineering feasibility review within 30 days", "Conduct FTO IP analysis with Schaeffler patent counsel", f"Identify pilot customer in {chr(44).join(sectors.get('primary_sectors',['target sector'])[:1])} for co-development conversation", "Present to Innovation steering committee with this report as supporting material"]
        }

    NAVY  = RGBColor(0x1F,0x38,0x64)
    BLUE  = RGBColor(0x2E,0x75,0xB6)
    WHITE = RGBColor(0xFF,0xFF,0xFF)
    GREY  = RGBColor(0x55,0x55,0x55)
    BLACK = RGBColor(0x00,0x00,0x00)
    LBLUE = RGBColor(0x60,0xA5,0xFA)

    # Sanitize every string before it touches the XML tree
    def xs(s):
        if s is None: return ""
        return _XML_CTRL_RE.sub('', str(s))

    def set_bg(cell, hex_col):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear")
        shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"), hex_col)
        tcPr.append(shd)

    def h1(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8")
        bot.set(qn("w:space"),"3"); bot.set(qn("w:color"),"2E75B6")
        pBdr.append(bot); pPr.append(pBdr)
        r = p.add_run(xs(text))
        r.bold=True; r.font.size=Pt(14); r.font.color.rgb=NAVY

    def h2(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(xs(text))
        r.bold=True; r.font.size=Pt(11); r.font.color.rgb=NAVY

    def body(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(xs(text))
        r.font.size=Pt(10.5)

    def kv(doc, label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(xs(f"{label}: "))
        r1.bold=True; r1.font.size=Pt(10.5); r1.font.color.rgb=NAVY
        r2 = p.add_run(xs(value))
        r2.font.size=Pt(10.5)

    def bul(doc, text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
        r = p.add_run(xs(text)); r.font.size=Pt(10.5)

    def hdr_row(tbl, headers):
        row = tbl.rows[0]
        for i, h in enumerate(headers):
            c = row.cells[i]; set_bg(c,"1F3864")
            r = c.paragraphs[0].add_run(xs(h))
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)

    # ── Header banner ─────────────────────────────────────────
    t = doc.add_table(rows=1,cols=1); t.style="Table Grid"
    c = t.cell(0,0); set_bg(c,"1F3864")
    p = c.paragraphs[0]
    p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    r = p.add_run("MARKET INTELLIGENCE REPORT")
    r.bold=True; r.font.size=Pt(9); r.font.color.rgb=WHITE
    p2 = c.add_paragraph()
    p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(10)
    r2 = p2.add_run("Schaeffler AI Innovation Research Assistant  ·  Stage 02")
    r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(0x93,0xC5,0xFD)
    doc.add_paragraph()

    # ── Title ─────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after=Pt(2)
    r = p.add_run(xs(market.get("market_name", idea[:80])))
    r.bold=True; r.font.size=Pt(18); r.font.color.rgb=NAVY
    p2 = doc.add_paragraph()
    r2 = p2.add_run(xs(f"Score: {final_score}/10  ·  Quadrant: {quadrant}  ·  {datetime.now().strftime('%d %B %Y')}"))
    r2.font.size=Pt(10); r2.italic=True; r2.font.color.rgb=GREY

    # ── Idea box ──────────────────────────────────────────────
    doc.add_paragraph()
    tb = doc.add_table(rows=1,cols=2); tb.style="Table Grid"
    c1=tb.cell(0,0); c2=tb.cell(0,1)
    set_bg(c1,"1F3864"); set_bg(c2,"EAF1FB")
    c1.width=Inches(0.12)
    c1.paragraphs[0].add_run("")
    rp=c2.paragraphs[0]
    rp.paragraph_format.space_before=Pt(8); rp.paragraph_format.space_after=Pt(2)
    rp.add_run("Innovation Idea").bold=True
    rb=rp.runs[0]; rb.font.size=Pt(9); rb.font.color.rgb=NAVY
    rp2=c2.add_paragraph()
    rp2.paragraph_format.space_before=Pt(0); rp2.paragraph_format.space_after=Pt(8)
    ri=rp2.add_run(xs(idea)); ri.font.size=Pt(10); ri.italic=True
    doc.add_paragraph()

    # ── Score summary ─────────────────────────────────────────
    h1(doc,"Score Summary")
    st2=doc.add_table(rows=1,cols=4); st2.style="Table Grid"
    hdr_row(st2,["Dimension","Score","Weight","Weighted"])
    rows_data=[
        ("Market Attractiveness",f"{weights['Market Attractiveness'][0]:.1f}/10","40%",f"{weights['Market Attractiveness'][0]*0.4:.1f}"),
        ("Sector Fit",f"{weights['Sector Fit'][0]:.1f}/10","35%",f"{weights['Sector Fit'][0]*0.35:.1f}"),
        ("Competition Opportunity",f"{weights['Competition Opportunity'][0]:.1f}/10","25%",f"{weights['Competition Opportunity'][0]*0.25:.1f}"),
    ]
    for i,(dim,score,weight,weighted) in enumerate(rows_data):
        row=st2.add_row(); fill="EAF1FB" if i%2==0 else "FFFFFF"
        for j,val in enumerate([dim,score,weight,weighted]):
            c=row.cells[j]; set_bg(c,fill)
            r=c.paragraphs[0].add_run(xs(val)); r.font.size=Pt(10); r.bold=(j==0)
    fr=st2.add_row()
    for c in fr.cells: set_bg(c,"1F3864")
    r=fr.cells[0].paragraphs[0].add_run("FINAL SCORE")
    r.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    r2=fr.cells[3].paragraphs[0].add_run(xs(f"{final_score}/10"))
    r2.bold=True; r2.font.size=Pt(11); r2.font.color.rgb=LBLUE
    doc.add_paragraph()

    # ── Executive summary ─────────────────────────────────────
    h1(doc,"Executive Summary")
    body(doc, ext.get("executive_summary",""))

    # ── Market size ───────────────────────────────────────────
    h1(doc,"Market Size & Growth")
    kv(doc,"Market size (2024)",_mval(market.get("market_size_current") or market.get("market_size_2024","N/A")))
    kv(doc,"Projected (2030)",_mval(market.get("market_size_forecast") or market.get("market_size_2030","N/A")))
    kv(doc,"CAGR",_mval(market.get("cagr","N/A")))
    kv(doc,"Maturity",market.get("market_maturity","N/A"))
    kv(doc,"Geography",market.get("geographic_focus","N/A"))
    doc.add_paragraph()
    body(doc, ext.get("market_deep_dive",""))
    if market.get("growth_drivers"):
        h2(doc,"Growth Drivers")
        for d in market["growth_drivers"]: bul(doc,d)

    # ── Sector clusters ───────────────────────────────────────
    h1(doc,"Schaeffler Sector Cluster Fit")
    primary=sectors.get("primary_sectors",[])
    body(doc,f"Primary sectors: {', '.join(primary)}")
    body(doc,sectors.get("sector_fit_rationale",""))
    doc.add_paragraph()
    st3=doc.add_table(rows=1,cols=3); st3.style="Table Grid"
    hdr_row(st3,["Sector","Score","Rationale"])
    for idx,(sector,data) in enumerate(sectors.get("sector_scores",{}).items()):
        row=st3.add_row()
        fill="EAF5EA" if sector in primary else ("EAF1FB" if idx%2==0 else "FFFFFF")
        for c in row.cells: set_bg(c,fill)
        r0=row.cells[0].paragraphs[0].add_run(xs(sector))
        r0.font.size=Pt(10); r0.bold=(sector in primary)
        r1=row.cells[1].paragraphs[0].add_run(xs(f"{data.get('score',0)}/10"))
        r1.font.size=Pt(10); r1.bold=True
        r1.font.color.rgb=BLUE if sector in primary else BLACK
        r2=row.cells[2].paragraphs[0].add_run(xs(data.get("rationale","")))
        r2.font.size=Pt(9.5)

    # ── Competitive landscape ─────────────────────────────────
    h1(doc,"Competitive Landscape")
    kv(doc,"Competitive intensity",comp.get("competitive_intensity",""))
    kv(doc,"White space",comp.get("white_space",""))
    kv(doc,"Schaeffler advantage",comp.get("schaeffler_advantage",""))
    doc.add_paragraph()
    body(doc, ext.get("competitive_analysis",""))
    if comp.get("competitors"):
        h2(doc,"Key Players")
        ct=doc.add_table(rows=1,cols=3); ct.style="Table Grid"
        hdr_row(ct,["Company","Type","Relevance"])
        for idx,ci in enumerate(comp["competitors"]):
            row=ct.add_row(); fill="EAF1FB" if idx%2==0 else "FFFFFF"
            for c in row.cells: set_bg(c,fill)
            for j,val in enumerate([ci.get("name",""),ci.get("type",""),ci.get("relevance","")+" "+ci.get("source","")]):
                r=row.cells[j].paragraphs[0].add_run(xs(val))
                r.font.size=Pt(9.5); r.bold=(j==0)

    # ── Strategic fit ─────────────────────────────────────────
    h1(doc,"Schaeffler Strategic Fit")
    body(doc, ext.get("schaeffler_fit",""))

    # ── Risks & recommendations ───────────────────────────────
    h1(doc,"Risks & Mitigations")
    risks_list = [
        f"Competitive intensity is {comp.get('competitive_intensity','unknown')} — monitor key players and conduct freedom-to-operate analysis before committing R&D budget",
        "Market timing risk — validate demand with target customers before scaling investment",
        "Sector fit assumptions should be validated with Schaeffler divisional teams before resource allocation"
    ]
    for risk in risks_list: bul(doc, risk)
    h1(doc,"Recommendations")
    recs_list = [
        "Commission internal engineering feasibility review within 30 days",
        "Conduct freedom-to-operate IP analysis with Schaeffler patent team",
        f"Identify pilot customer in {', '.join(sectors.get('primary_sectors', ['target sector'])[:1])} for co-development",
        "Present to Innovation steering committee with this report as supporting material"
    ]
    for rec in recs_list: bul(doc, rec)

    # ── Footer ────────────────────────────────────────────────
    doc.add_paragraph()
    ft=doc.add_table(rows=1,cols=1); ft.style="Table Grid"
    fc=ft.cell(0,0); set_bg(fc,"1F3864")
    fp=fc.paragraphs[0]
    fp.paragraph_format.space_before=Pt(6); fp.paragraph_format.space_after=Pt(6)
    fr=fp.add_run(xs(f"Schaeffler AI Innovation Research Assistant  ·  Stage 02: Market Intelligence  ·  {datetime.now().strftime('%d %B %Y')}  ·  Capstone Project — Arpan Chowdhury, EBS Universität"))
    fr.font.size=Pt(8); fr.font.color.rgb=RGBColor(0x93,0xC5,0xFD)

    _sanitize_doc(doc)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf



# ════════════════════════════════════════════════════════════
# RUN FULL ANALYSIS — sequential pipeline helper
# ════════════════════════════════════════════════════════════
def run_stage2(idea, quadrant, s1c):
    """Run Stage 02 Market Intelligence and store results in session state."""
    web_ctx = ""
    system_market = """You are a senior market analyst. Analyse the market for this innovation idea.
RULES:
- Use only the most credible sources: McKinsey Global Institute, Gartner, Frost & Sullivan, BloombergNEF, IEA, Roland Berger, Statista, MarketsandMarkets, Grand View Research, Allied Market Research, Mordor Intelligence, Fortune Business Insights, IDC, Wood Mackenzie, S&P Global.
- For each market figure, provide structured source objects. Each source MUST have its own entry in the sources array.
- For URLs: provide the EXACT deep-link URL. If you do not know it, omit the url field entirely.
- market_score: integer 1-10. Guide: 9-10=large fast-growing; 7-8=strong; 5-6=moderate; 3-4=niche; 1-2=tiny/declining.
Return ONLY valid JSON with NO extra text, NO inline comments:
{"market_name":"string",
"market_size_current":{"value":"$X.XB","year":"2024","sources":[{"org":"OrgName","title":"Report Title Year","year":"2024","url":"https://url-or-omit-field"}]},
"market_size_forecast":{"value":"$X.XB","year":"2030","sources":[{"org":"OrgName","title":"Report Title Year","year":"2024","url":"https://url-or-omit-field"}]},
"cagr":{"value":"X%","period":"2024-2030","sources":[{"org":"OrgName","title":"Report Title Year","year":"2024","url":"https://url-or-omit-field"}]},
"growth_drivers":["driver 1","driver 2","driver 3"],"market_maturity":"Emerging/Growing/Mature/Declining",
"geographic_focus":"string","market_score":7,"market_score_rationale":"2 sentences"}"""
    raw = call_claude(system_market, f"Idea: {idea}\nQuadrant: {quadrant}\nTech: {s1c.get('technology_novelty','')}", max_tokens=1400)
    try:
        market = _parse_json(raw)
    except Exception:
        market = {"market_name":"N/A","market_size_current":{"value":"N/A","year":"2024","sources":[]},"market_size_forecast":{"value":"N/A","year":"2030","sources":[]},"cagr":{"value":"N/A","period":"","sources":[]},"growth_drivers":[],"market_maturity":"N/A","geographic_focus":"N/A","market_score":5,"market_score_rationale":""}

    system_comp = """You are a competitive intelligence analyst. Identify key competitors for this idea.
Every company must have a source. Return ONLY valid JSON with NO inline comments or annotations on numeric fields:
{"competitors":[{"name":"string","type":"Incumbent/Startup/Research","relevance":"one sentence","source":"Source: X, Y"}],
"competitive_intensity":"Low/Medium/High/Very High","white_space":"one sentence","schaeffler_advantage":"one sentence",
"competition_score":7,"competition_score_rationale":"2 sentences"}
Scoring guide (do NOT include this line or any annotations in the JSON): 9-10=very open/few players; 7-8=some room; 5-6=moderate; 3-4=crowded; 1-2=saturated."""
    raw = call_claude(system_comp, f"Idea: {idea}\nMarket: {market.get('market_name','')}\nQuadrant: {quadrant}", max_tokens=2000)
    try:
        comp = _parse_json(raw)
    except Exception:
        comp = {"competitors":[],"competitive_intensity":"N/A","white_space":"N/A","schaeffler_advantage":"N/A","competition_score":5,"competition_score_rationale":""}

    system_sectors = """You are a Schaeffler strategist. Score fit against Schaeffler's 10 sector clusters (0-10 each).
Clusters: Passenger Cars, Commercial Vehicles, Industrial Machinery, Rail, Aerospace, Two-Wheelers, Construction & Agriculture, Medical Equipment, Conventional Energy, Renewable Energy.
Return ONLY valid JSON:
{"sector_scores":{"Passenger Cars":{"score":0-10,"rationale":"one sentence"},"Commercial Vehicles":{"score":0-10,"rationale":"one sentence"},"Industrial Machinery":{"score":0-10,"rationale":"one sentence"},"Rail":{"score":0-10,"rationale":"one sentence"},"Aerospace":{"score":0-10,"rationale":"one sentence"},"Two-Wheelers":{"score":0-10,"rationale":"one sentence"},"Construction & Agriculture":{"score":0-10,"rationale":"one sentence"},"Medical Equipment":{"score":0-10,"rationale":"one sentence"},"Conventional Energy":{"score":0-10,"rationale":"one sentence"},"Renewable Energy":{"score":0-10,"rationale":"one sentence"}},
"primary_sectors":["top 2-3 sector names"],"sector_fit_score":0-10,"sector_fit_rationale":"2 sentences"}"""
    raw = call_claude(system_sectors, f"Idea: {idea}\nQuadrant: {quadrant}", max_tokens=1000)
    try:
        sectors = _parse_json(raw)
    except:
        sectors = {"sector_scores":{},"primary_sectors":[],"sector_fit_score":5,"sector_fit_rationale":""}

    ms = float(market.get("market_score",5))
    ss = float(sectors.get("sector_fit_score",5))
    cs = float(comp.get("competition_score",5))
    final = round(ms*0.40 + ss*0.35 + cs*0.25, 1)

    st.session_state.s2_data = {
        "market": market, "comp": comp, "sectors": sectors,
        "weights": {"Market Attractiveness":(ms,0.40),"Sector Fit":(ss,0.35),"Competition Opportunity":(cs,0.25)},
        "final_score": final, "web_results": []
    }
    st.session_state.s2_step = "done"


def _parse_json_robust(raw):
    """Robustly parse JSON from Claude output. Returns None on failure."""
    if not raw:
        return None
    try:
        return _parse_json(raw)
    except Exception:
        return None


def _compute_ip_proximity_risk(filer_positions, ref_x, ref_y, threshold=2.0):
    """
    Deterministically compute IP risk from Ansoff matrix filer proximity.
    ≥5 nearby → High risk → score 2
    2–4 nearby → Medium risk → score 5
    <2 nearby  → Low risk  → score 8
    """
    nearby = sum(
        1 for f in filer_positions
        if abs(float(f.get("x_score", 5)) - ref_x) <= threshold
        and abs(float(f.get("y_score", 5)) - ref_y) <= threshold
    )
    if nearby >= 5:
        return "High", 2, nearby
    elif nearby >= 2:
        return "Medium", 5, nearby
    else:
        return "Low", 8, nearby


def run_stage3(idea, quadrant, s1c):
    """Run Stage 03 Patent Intelligence and store results in session state."""
    system_landscape = """You are a patent intelligence analyst specialising in industrial technology.
Analyse the external patent landscape for this innovation idea.

RULES:
- Focus on COMPANIES filing patents, not individual patents
- Classify each company as: Competitor / Customer / Research Institution / Patent Troll / Adjacent Player
- Every claim must have [Source: org, year] where possible
- Be specific about technology sub-areas

Return ONLY valid JSON:
{
  "technology_keywords": ["3-5 key patent search terms for this idea"],
  "landscape_summary": "2-3 sentences on overall patent activity in this space",
  "activity_level": "Low / Moderate / High / Very High",
  "filing_trend": "Increasing / Stable / Decreasing",
  "filing_trend_rationale": "one sentence",
  "patent_landscape_score": integer 1-10 for landscape openness. Use: 9-10=very few filings, open territory; 7-8=some activity but clear gaps; 5-6=moderate filing density; 3-4=dense filing landscape; 1-2=saturated with patents by large incumbents,
  "key_filers": [
    {
      "company": "company name",
      "type": "Competitor / Customer / Research Institution / Adjacent Player",
      "focus": "one sentence on what they are patenting",
      "threat_level": "Low / Medium / High",
      "schaeffler_relationship": "Direct competitor / Potential customer / Partner / Unknown",
      "source": "Source: X, Year"
    }
  ],
  "white_spaces": ["white space opportunity 1", "white space opportunity 2", "white space opportunity 3"]
}"""
    raw = call_claude(system_landscape, f"Idea: {idea}\nQuadrant: {quadrant}\nTech novelty: {s1c.get('technology_novelty','')}", max_tokens=2000)
    try:
        landscape = _parse_json(raw)
    except:
        landscape = {"technology_keywords":[],"landscape_summary":"N/A","activity_level":"N/A","filing_trend":"N/A","filing_trend_rationale":"","patent_landscape_score":5,"key_filers":[],"white_spaces":[]}

    key_filers_run3 = landscape.get("key_filers", [])
    filers_full_run3 = json.dumps([
        {"company": f.get("company",""), "type": f.get("type",""), "focus": f.get("focus",""), "threat_level": f.get("threat_level","")}
        for f in key_filers_run3
    ])
    system_ansoff = """You are a Schaeffler Group patent strategist.
Map patent filing companies onto Schaeffler's modified Ansoff matrix based on where their patents sit.

The matrix axes (MUST match Schaeffler's Stage 1 Innovation Framework):
- X axis: Technology Dimension (0=Existing/Established Technology → 10=New to the World)
- Y axis: Market Dimension (0=Existing/Established Market → 10=New to the World)

Quadrant positions (identical to Schaeffler's quadrant classifier):
- EXPLOIT  (bottom-left):  existing tech  + existing market  — x_score 0–5,  y_score 0–5
- EXTEND   (top-left):     existing tech  + new market       — x_score 0–5,  y_score 5–10
- RADICAL  (top-right):    new tech       + new market       — x_score 5–10, y_score 5–10
- DISRUPT  (bottom-right): new tech       + existing market  — x_score 5–10, y_score 0–5

For each company, assign:
- matrix_position: which quadrant their patent activity sits in
- x_score: 0-10 (0=existing technology, 10=new to the world technology)
- y_score: 0-10 (0=existing market, 10=new to the world market)

Also map where SCHAEFFLER's own known IP sits relative to this idea.

Return ONLY valid JSON:
{
  "filer_positions": [
    {
      "company": "company name",
      "matrix_position": "EXPLOIT/EXTEND/RADICAL/DISRUPT",
      "x_score": 0-10,
      "y_score": 0-10,
      "rationale": "one sentence"
    }
  ],
  "schaeffler_position": {
    "matrix_position": "EXPLOIT/EXTEND/RADICAL/DISRUPT",
    "x_score": 0-10,
    "y_score": 0-10,
    "existing_ip": "one sentence on what Schaeffler already has in this space",
    "gap": "one sentence on the IP gap this idea addresses"
  },
  "idea_position": {
    "x_score": 0-10,
    "y_score": 0-10
  },
  "novelty_signal": "Strong / Moderate / Weak",
  "novelty_rationale": "one sentence",
  "ip_risk": "Low / Medium / High",
  "ip_risk_rationale": "one sentence"
}"""
    raw2 = call_claude(system_ansoff,
        f"Idea: {idea}\nQuadrant: {quadrant}\nTech keywords: {landscape.get('technology_keywords','')}\n"
        f"IMPORTANT: You MUST map ALL {len(key_filers_run3)} filers listed below. Do not skip any.\n"
        f"Key filers (map every single one): {filers_full_run3}",
        max_tokens=max(2500, len(key_filers_run3) * 250 + 1200))
    try:
        ansoff_data = _parse_json(raw2)
    except:
        ansoff_data = {"filer_positions":[],"schaeffler_position":{"matrix_position":"EXPLOIT","x_score":2,"y_score":2,"existing_ip":"N/A","gap":"N/A"},"idea_position":{"x_score":7,"y_score":7},"novelty_signal":"Moderate","novelty_rationale":"","ip_risk":"Medium","ip_risk_rationale":""}

    # Guarantee every key_filer has a position
    positioned_run3 = {fp.get("company","").lower() for fp in ansoff_data.get("filer_positions", [])}
    type_defaults_run3 = {
        "Competitor":          ("EXPLOIT", 3.0, 3.0),
        "Customer":            ("EXTEND",  2.5, 6.5),
        "Research Institution":("RADICAL", 7.0, 7.5),
        "Adjacent Player":     ("DISRUPT", 6.5, 3.5),
        "Patent Troll":        ("EXPLOIT", 2.0, 2.0),
    }
    for i, f in enumerate(key_filers_run3):
        name = f.get("company","")
        if name.lower() not in positioned_run3 and name:
            quad, bx, by = type_defaults_run3.get(f.get("type","Adjacent Player"), ("EXPLOIT", 4.0, 4.0))
            nudge_x = ((i * 0.7) % 2.0) - 1.0
            nudge_y = ((i * 1.1) % 2.0) - 1.0
            ansoff_data.setdefault("filer_positions", []).append({
                "company": name, "type": f.get("type","Adjacent Player"),
                "matrix_position": quad,
                "x_score": round(min(9.5, max(0.5, bx + nudge_x)), 1),
                "y_score": round(min(9.5, max(0.5, by + nudge_y)), 1),
                "rationale": f.get("focus","Auto-placed based on filer type")
            })

    # ── Landscape score: from LLM patent_landscape_score ─────────────────────
    landscape_score = float(landscape.get("patent_landscape_score", 5))

    # ── Novelty score: separate domain-expert Claude call (more reliable) ────
    filer_summary = ", ".join(
        f"{f.get('company','')} ({f.get('type','')})"
        for f in key_filers_run3[:8]
    )
    system_novelty = """You are a domain expert academic assessing patent novelty.
Evaluate the novelty signal based on the patent landscape and filer data provided.

Novelty Signal definitions (apply strictly):
- High:   Blue ocean. Very few patent filers in this specific technology/market combination.
- Medium: Oligopoly. A handful of established players dominate but meaningful gaps exist.
- Low:    Red ocean. Many patent filers, technology is well-covered.

Return ONLY valid JSON:
{
  "novelty_signal": "High / Medium / Low",
  "novelty_rationale": "2-3 sentences from a domain expert perspective",
  "novelty_score": <integer 1-10>
}
Scoring: 9-10=High (uncontested); 7-8=High-leaning; 5-6=Medium; 3-4=Medium-low; 1-2=Low (saturated)"""

    novelty_ctx = (
        f"Idea: {idea}\nQuadrant: {quadrant}\n"
        f"Landscape summary: {landscape.get('landscape_summary','')}\n"
        f"Activity level: {landscape.get('activity_level','')}\n"
        f"Filing trend: {landscape.get('filing_trend','')}\n"
        f"Key filers: {filer_summary}"
    )
    try:
        raw_nov = call_claude(system_novelty, novelty_ctx, max_tokens=500)
        novelty_data = _parse_json_robust(raw_nov)
        if not novelty_data or not isinstance(novelty_data, dict):
            novelty_data = {"novelty_signal": "Medium", "novelty_rationale": "Assessment unavailable.", "novelty_score": 6}
    except Exception:
        novelty_data = {"novelty_signal": "Medium", "novelty_rationale": "Assessment unavailable.", "novelty_score": 6}
    novelty_score = float(novelty_data.get("novelty_score", 6))

    # ── IP Risk scores: deterministic from Ansoff matrix proximity ────────────
    filer_positions_r3 = ansoff_data.get("filer_positions", [])
    idea_x  = float(ansoff_data.get("idea_position", {}).get("x_score", 7.0))
    idea_y  = float(ansoff_data.get("idea_position", {}).get("y_score", 7.0))
    sch_x   = float(ansoff_data.get("schaeffler_position", {}).get("x_score", 3.0))
    sch_y   = float(ansoff_data.get("schaeffler_position", {}).get("y_score", 3.0))

    ip_idea_label, ip_idea_score, ip_idea_nearby           = _compute_ip_proximity_risk(filer_positions_r3, idea_x, idea_y)
    ip_sch_label,  ip_sch_score,  ip_sch_nearby            = _compute_ip_proximity_risk(filer_positions_r3, sch_x, sch_y)

    # ── Final score: 4 components, 25% each ──────────────────────────────────
    final_patent = round((landscape_score + novelty_score + ip_idea_score + ip_sch_score) / 4, 1)

    st.session_state.s3_data = {
        "landscape":           landscape,
        "ansoff_data":         ansoff_data,
        "novelty_data":        novelty_data,
        "landscape_score":     landscape_score,
        "novelty_score":       novelty_score,
        "ip_idea_label":       ip_idea_label,
        "ip_idea_score":       ip_idea_score,
        "ip_idea_nearby":      ip_idea_nearby,
        "ip_schaeffler_label": ip_sch_label,
        "ip_schaeffler_score": ip_sch_score,
        "ip_sch_nearby":       ip_sch_nearby,
        "final_score":         final_patent,
    }
    st.session_state.s3_step = "done"


def run_stage4(idea, quadrant, s1c):
    """Run Stage 04 Technical Feasibility and store results in session state."""
    system_existence = """You are a technology intelligence analyst specialising in industrial and automotive R&D.
Check whether the core technology behind this innovation idea has been demonstrated anywhere.
Look for evidence in: academic research, university labs, startup products, government programmes,
industry pilots, defence/aerospace, and adjacent industries.

Return ONLY valid JSON:
{
  "technology_core": "one sentence describing the core technology mechanism",
  "existence_verdict": "Demonstrated / Partially Demonstrated / Research Stage / Theoretical",
  "existence_summary": "2-3 sentences on the state of the art",
  "evidence": [
    {
      "type": "Academic Paper / Startup / Pilot / Government Programme / Industry Deployment",
      "title": "name of the paper, company, or programme",
      "description": "one sentence on what was demonstrated",
      "source": "Source: org/journal, year",
      "confidence": "High / Medium / Low",
      "relevance": "Direct / Adjacent / Analogous"
    }
  ],
  "technology_gaps": ["gap 1 between current state and full deployment", "gap 2", "gap 3"],
  "time_to_readiness": "estimated years to production readiness",
  "keywords": ["6-10 key technical terms from this domain for keyword map"]
}"""
    try:
        raw = call_claude(system_existence,
            f"Idea: {idea}\nQuadrant: {quadrant}\nTech: {s1c.get('technology_novelty','')}", max_tokens=3000)
        existence = _parse_json_robust(raw)
        if not existence or not isinstance(existence, dict):
            raise ValueError("Parse failed")
    except Exception:
        existence = {"technology_core": "N/A", "existence_verdict": "Research Stage", "existence_summary": "N/A",
                     "evidence": [], "technology_gaps": [], "time_to_readiness": "Not yet estimated", "keywords": []}

    system_trl = """You are a Schaeffler R&D director assessing technology maturity.
Use the Schaeffler-adapted TRL framework (modified from NASA TRL for industrial/automotive context):

TRL 1 — Basic principles observed (theoretical concept only)
TRL 2 — Technology concept formulated (application identified, no testing)
TRL 3 — Experimental proof of concept (lab demonstration, key functions validated)
TRL 4 — Technology validated in lab (component tested in controlled environment)
TRL 5 — Validated in relevant environment (prototype tested in industrial-like conditions)
TRL 6 — Demonstrated in relevant environment (system prototype demonstrated)
TRL 7 — System prototype in operational environment (field trial or industrial pilot)
TRL 8 — System complete and qualified (full production design, limited production run)
TRL 9 — Proven in operational environment (commercial deployment at scale)

Return ONLY valid JSON:
{
  "trl_level": 1-9,
  "trl_label": "TRL X — label from framework above",
  "trl_rationale": "2-3 sentences justifying this TRL rating based on evidence",
  "schaeffler_entry_readiness": "Too Early / Ready for Innovation / Ready for Product Development",
  "entry_rationale": "one sentence",
  "key_technical_risks": [
    {"risk": "technical risk description", "severity": "High/Medium/Low", "mitigation": "one sentence"}
  ],
  "analogous_schaeffler_technologies": "one sentence on which of Schaeffler's 8 Motion Product Families (Guide Motion/Transmit Motion/Control Motion/Generate Motion/Power Motion/Drive Motion/Energize Motion/Sustain Motion) this is closest to"
}"""
    try:
        raw2 = call_claude(system_trl,
            f"Idea: {idea}\nExistence verdict: {existence.get('existence_verdict','')}\nEvidence count: {len(existence.get('evidence',[]))}\nGaps: {existence.get('technology_gaps',[])}",
            max_tokens=1500)
        trl = _parse_json_robust(raw2)
        if not trl or not isinstance(trl, dict):
            raise ValueError("Parse failed")
    except Exception:
        trl = {"trl_level": 0, "trl_label": "TRL — parse failed, re-run",
               "trl_rationale": "", "schaeffler_entry_readiness": "Too Early",
               "entry_rationale": "", "key_technical_risks": [], "analogous_schaeffler_technologies": ""}

    # ── TRL score: fully deterministic lookup — no LLM variance ──────────────
    _TRL_SCORE_MAP = {1: 1.0, 2: 2.0, 3: 3.5, 4: 5.0, 5: 6.0, 6: 7.0, 7: 8.0, 8: 9.0, 9: 10.0}
    trl_level_val = int(trl.get("trl_level", 0))
    trl_score = _TRL_SCORE_MAP.get(trl_level_val, 0.0)

    # ── Existence score: deterministic map ────────────────────────────────────
    ev_map = {"Demonstrated": 9, "Partially Demonstrated": 6, "Research Stage": 3, "Theoretical": 1}
    existence_score = float(ev_map.get(existence.get("existence_verdict", "Research Stage"), 0.0))

    # ── Risk score: High=2, Medium=5, Low=8 (inverted — lower risk = higher score) ──
    _sev_safety = {"High": 2, "Medium": 5, "Low": 8}
    all_risks = trl.get("key_technical_risks", [])
    risk_score = round(
        sum(_sev_safety.get(r.get("severity", "Medium"), 5) for r in all_risks) / max(len(all_risks), 1), 1
    ) if all_risks else 0.0

    # ── Final: equal 33.3% weights ───────────────────────────────────────────
    final_feasibility = round((trl_score + existence_score + risk_score) / 3, 1)

    st.session_state.s4_data = {
        "existence":       existence,
        "trl":             trl,
        "trl_score":       trl_score,
        "existence_score": existence_score,
        "risk_score":      risk_score,
        "final_score":     final_feasibility
    }
    st.session_state.s4_step = "done"


def run_stage5(idea, quadrant, s1c):
    """Run Stage 05 P³ Perspective and store results in session state."""
    # ── Safe variable preparation — guard against None from Claude-stored nulls ──
    s3_landscape = st.session_state.get("s3_data") or {}
    s3_landscape = s3_landscape.get("landscape") or {}
    s4_data      = st.session_state.get("s4_data") or {}
    s4_existence = s4_data.get("existence") or {}
    s4_trl       = s4_data.get("trl") or {}

    raw_filers  = s3_landscape.get("key_filers") or []
    raw_sources = s4_existence.get("evidence") or []
    prior_filers          = [str(f.get("company","")) for f in raw_filers if isinstance(f, dict)]
    prior_evidence_sources= [str(e.get("source","")) for e in raw_sources if isinstance(e, dict)]

    trl_level          = int(s4_trl.get("trl_level") or 3)
    innovation_cluster = str(s1c.get("innovation_cluster") or "")
    product_family     = str(s1c.get("product_family") or "")
    raw_trends         = s1c.get("trend_alignment")
    trend_alignment    = list(raw_trends) if isinstance(raw_trends, (list, tuple)) else []
    innovation_model   = str(s1c.get("innovation_model") or "Integrated")

    filers_str  = ", ".join(prior_filers[:6])   or "none identified"
    sources_str = ", ".join(prior_evidence_sources[:5]) or "none identified"
    trends_str  = ", ".join(str(t) for t in trend_alignment) or "none"

    system_readiness = (
        "You are a senior Schaeffler innovation strategist assessing internal P³ Perspective.\n"
        "Schaeffler P³ formula: Performance = Portfolio x People x Process.\n"
        f"Innovation cluster: {innovation_cluster} | Product family: {product_family} | "
        f"Strategic trends: {trends_str} | Current TRL: {trl_level}\n"
        f"Patent filers from Stage 03: {filers_str}\n"
        f"Evidence sources from Stage 04: {sources_str}\n"
        "Schaeffler competencies: precision bearings, mechatronics, power electronics (Vitesco merger), "
        "tribology, EV drivetrains, embedded sensors, ASPICE/ISO 26262, OEM Tier 1 supply chain.\n\n"
        f"Innovation model: {innovation_model}\n"
        "- RADICAL (Integrated model): assess P3 readiness for Schaeffler FIP-VEP-PEP process, "
        "leveraging OEM relationships, manufacturing scale, internal R&D.\n"
        "- DISRUPTIVE (Accelerator model): assess P3 readiness for the Accelerator/VC track — "
        "external co-development, startup partnerships, VC co-investment. Internal P3 gaps are expected.\n\n"
        "SCORING RUBRICS (apply strictly):\n"
        "Portfolio score (strategic fit):\n"
        "  9-10 = Direct alignment with a Schaeffler innovation cluster, addresses a defined strategic trend, clear product family fit\n"
        "  7-8  = Good alignment with one cluster, moderate trend relevance, identifiable product family\n"
        "  5-6  = Partial fit, indirect relevance to cluster or trend\n"
        "  3-4  = Weak strategic fit, marginal cluster relevance\n"
        "  1-2  = No clear fit with any innovation cluster or strategic direction\n"
        "People score (competency readiness):\n"
        "  9-10 = Core competencies fully matched within Schaeffler, no critical gaps, can execute immediately\n"
        "  7-8  = Most competencies matched, one manageable gap with a clear closure route (hire/upskill)\n"
        "  5-6  = Partial match, 1-2 significant gaps requiring external sourcing or partnership\n"
        "  3-4  = Major competency gaps, requires significant hiring, acquisition, or JDA\n"
        "  1-2  = Fundamental capability mismatch, no relevant expertise at Schaeffler\n"
        "Process score (infrastructure & asset readiness):\n"
        "  9-10 = Existing Schaeffler processes and assets directly applicable, minimal new investment needed\n"
        "  7-8  = Most processes applicable, some adaptation or moderate investment required\n"
        "  5-6  = Moderate process fit, meaningful investment and new tooling required\n"
        "  3-4  = Few applicable processes, significant new infrastructure needed\n"
        "  1-2  = No applicable processes, requires building capability from scratch\n\n"
        "Return ONLY valid JSON with exactly these keys. Use real integer scores 1-10, real strings, real arrays:\n"
        '{"p3_portfolio":{"score":7,"rationale":"two sentences","cluster_fit":"one sentence",'
        '"strengths":["strength 1","strength 2"],"gaps":["gap 1"]},'
        '"p3_people":{"score":6,"rationale":"two sentences","matched_competencies":["comp 1","comp 2","comp 3"],'
        '"competency_gap":"the single critical missing competency","sourcing_route":"how to close the gap"},'
        '"p3_process":{"score":6,"rationale":"two sentences","applicable_assets":["asset 1","asset 2"],'
        '"investment_required":"what needs to be built or acquired","time_to_close":"estimated months"},'
        '"partnership_candidates":[{"name":"org name","type":"Startup","rationale":"why them","route":"Co-develop"}],'
        '"org_gaps":[{"gap":"gap name","severity":"High","closure_route":"how to close","timeline":"6 months"}],'
        '"build_or_partner":{"recommendation":"Co-develop","rationale":"two to three sentences",'
        '"time_to_trl6_internal":"30 months","time_to_trl6_partner":"18 months"},'
        '"p3_perspective_score":6}'
    )

    _p3_fallback = {
        "p3_portfolio":{"score":5,"rationale":"Analysis unavailable — re-run Stage 05.","cluster_fit":"N/A","strengths":[],"gaps":[]},
        "p3_people":{"score":5,"rationale":"Analysis unavailable — re-run Stage 05.","matched_competencies":[],"competency_gap":"N/A","sourcing_route":"N/A"},
        "p3_process":{"score":5,"rationale":"Analysis unavailable — re-run Stage 05.","applicable_assets":[],"investment_required":"N/A","time_to_close":"N/A"},
        "partnership_candidates":[],"org_gaps":[],
        "build_or_partner":{"recommendation":"Co-develop","rationale":"Analysis unavailable.","time_to_trl6_internal":"N/A","time_to_trl6_partner":"N/A"},
        "p3_perspective_score":5
    }

    try:
        raw = call_claude(system_readiness,
                          f"Innovation idea: {idea}\nQuadrant: {quadrant}\nTRL level: {trl_level}",
                          max_tokens=2500)
        org_data = _parse_json(raw)
        # Light sanity check — if top-level keys missing, retry once
        if "p3_portfolio" not in org_data or "p3_people" not in org_data:
            raise ValueError("Missing required P3 keys")
    except Exception as _e5:
        try:
            raw2 = call_claude(
                "Return ONLY valid JSON. No markdown, no commentary. "
                "Required top-level keys: p3_portfolio, p3_people, p3_process, "
                "partnership_candidates, org_gaps, build_or_partner, p3_perspective_score. "
                "Each score field must be an integer 1-10. Each rationale must be a real string.",
                f"Idea: {idea}\nQuadrant: {quadrant}\nTRL: {trl_level}\n"
                f"Cluster: {innovation_cluster}\nModel: {innovation_model}\n"
                "Assess Schaeffler's internal P3 readiness (Portfolio, People, Process) for this idea.",
                max_tokens=2500
            )
            org_data = _parse_json(raw2)
        except Exception as _e5_retry:
            import traceback
            st.error(f"⚠️ Stage 5 API Error (retry failed): {str(_e5_retry)}")
            st.info("Falling back to default assessment. Check API key and Claude connectivity.")
            st.caption(f"Debug: {traceback.format_exc()[:200]}")
            org_data = _p3_fallback

    p_portfolio = float(org_data.get("p3_portfolio",{}).get("score",5))
    p_people    = float(org_data.get("p3_people",{}).get("score",5))
    p_process   = float(org_data.get("p3_process",{}).get("score",5))
    final_org   = round((p_portfolio + p_people + p_process) / 3, 1)

    st.session_state.s5_data = {
        "org_data": org_data,
        "p_portfolio": p_portfolio,
        "p_people": p_people,
        "p_process": p_process,
        "final_score": final_org
    }
    st.session_state.s5_step = "done"


def run_stage6_synthesis(idea, quadrant, s1c):
    """Run Stage 06 synthesis with default weights and store results."""
    market_score      = st.session_state.s2_data.get("final_score", 5.0)
    patent_score      = st.session_state.s3_data.get("final_score", 5.0)
    feasibility_score = st.session_state.s4_data.get("final_score", 5.0)
    org_score         = st.session_state.s5_data.get("final_score", 5.0)

    weights = {"market":25,"patent":25,"feasibility":25,"org":25}
    ipi = round((market_score + patent_score + feasibility_score + org_score) / 4, 1)

    s2d = st.session_state.s2_data
    s3d = st.session_state.s3_data
    s4d = st.session_state.s4_data
    s5d = st.session_state.s5_data
    org_d = s5d.get("org_data",{})

    synthesis_context = f"""Idea: {idea}\nQuadrant: {quadrant}\nIPI: {ipi}/10
Market ({weights['market']}%): {market_score}/10 — {s2d.get('market',{}).get('market_name','')}
Patent ({weights['patent']}%): {patent_score}/10 — Novelty: {s3d.get('ansoff_data',{}).get('novelty_signal','')} IP risk: {s3d.get('ansoff_data',{}).get('ip_risk','')}
Feasibility ({weights['feasibility']}%): {feasibility_score}/10 — TRL {s4d.get('trl',{}).get('trl_level','')} {s4d.get('trl',{}).get('schaeffler_entry_readiness','')}
{('⚠️ SENSING PHASE (FUTURE OPTIONS TRACK): This idea was classified as a Future Option at Stage 1 — technology novelty in the borderline 5.0–6.5 band or no innovation cluster assigned yet. Frame the recommendation and next steps around what evidence is needed to graduate this idea to the full pipeline. Use "SENSING PHASE — CONTINUE MONITORING" as the recommendation where the data supports further maturation rather than immediate investment or rejection.' if s1c.get('route') == 'FUTURE_OPTIONS' else '')}
P³ Score ({weights['org']}%): {org_score}/10 — {org_d.get('build_or_partner',{}).get('recommendation','')}
Innovation Model: {s1c.get('innovation_model','Integrated') or 'Integrated'} · Pipeline Route: {s1c.get('pipeline_route','')}"""

    system_structured = """You are a senior Schaeffler innovation strategist. Return ONLY valid JSON:
{"headline":"one direct sentence","recommendation":"PROCEED or PROCEED WITH CONDITIONS or SENSING PHASE — CONTINUE MONITORING or DEFER or REJECT",
"recommendation_rationale":"2-3 sentences","strongest_signals":["signal 1","signal 2","signal 3"],
"key_concerns":["concern 1","concern 2","concern 3"],"conditions":["condition 1","condition 2"],
"strategic_fit":"2-3 sentences referencing Schaeffler P³, electrification, Vitesco merger",
"risks":["risk 1 with mitigation","risk 2","risk 3"],
"next_steps":["action 1","action 2","action 3","action 4"]}"""

    raw1 = call_claude(system_structured, synthesis_context, max_tokens=2000)
    raw1_clean = raw1.strip().replace("```json","").replace("```","").strip()
    fb = raw1_clean.find("{"); lb = raw1_clean.rfind("}") + 1
    if fb >= 0: raw1_clean = raw1_clean[fb:lb]
    try:
        synthesis_structured = json.loads(raw1_clean)
    except:
        synthesis_structured = {"headline":f"IPI {ipi}/10","recommendation":"PROCEED WITH CONDITIONS" if ipi>=5 else "DEFER","recommendation_rationale":"Based on pipeline analysis.","strongest_signals":[],"key_concerns":[],"conditions":[],"strategic_fit":"","risks":[],"next_steps":[]}

    system_narrative = "Write a 4-paragraph narrative synthesis for this Schaeffler innovation assessment. Flowing prose, no bullets. Cover: market opportunity, IP landscape, technical maturity, P³ Perspective, and recommendation. Reference Schaeffler P³ formula and electrification context."
    raw2 = call_claude(system_narrative, synthesis_context + f"\nRecommendation: {synthesis_structured.get('recommendation','')}\nIPI: {ipi}/10", max_tokens=1400)
    narrative_text = raw2.strip().replace("```","").strip()

    synthesis = {**synthesis_structured, "narrative": narrative_text}

    st.session_state.s6_data = {
        "ipi": ipi, "weights": weights, "synthesis": synthesis,
        "scores": {"market":market_score,"patent":patent_score,"feasibility":feasibility_score,"p3":org_score}
    }
    st.session_state.s6_step = "done"

    # Auto-save fires here so it runs on the full-run path (active_stage stays 1,
    # Stage 6 UI never renders, so the save in that block is never reached).
    if not st.session_state.get("_s6_saved_to_sheets"):
        qa_pairs = []
        for q, a in zip(st.session_state.get("s1_questions", []), st.session_state.get("s1_answers", [])):
            qa_pairs.append(f"Q: {q} / A: {a}")
        row = {
            "Date":                  datetime.now().strftime("%Y-%m-%d %H:%M"),
            "Submitter Name":        st.session_state.get("user_name", ""),
            "Position":              st.session_state.get("user_position", ""),
            "Department":            st.session_state.get("user_dept", ""),
            "Full Idea Description": idea,
            "Clarifying Q&A":        " | ".join(qa_pairs),
            "Quadrant":              quadrant,
            "Innovation Cluster":    s1c.get("innovation_cluster", ""),
            "Product Family":        s1c.get("product_family", ""),
            "Market Score":          str(market_score),
            "Patent Score":          str(patent_score),
            "Feasibility Score":     str(feasibility_score),
            "P³ Score":         str(org_score),
            "IPI Score":             str(ipi),
            "Recommendation":        synthesis.get("recommendation", ""),
            "Key Concerns":          " | ".join(synthesis.get("key_concerns", [])[:3]),
            "Next Steps":            " | ".join(synthesis.get("next_steps", [])[:4]),
            "Market Name":           s2d.get("market", {}).get("market_name", ""),
            "Market Size 2024":      _mval(s2d.get("market", {}).get("market_size_current") or s2d.get("market", {}).get("market_size_2024", "")),
            "CAGR":                  _mval(s2d.get("market", {}).get("cagr", "")),
            "TRL Level":             str(s4d.get("trl", {}).get("trl_level", "")),
            "Build Strategy":        s5d.get("org_data", {}).get("build_or_partner", {}).get("recommendation", ""),
        }
        saved, _err = save_idea_to_sheets(row)
        if saved:
            st.session_state["_s6_saved_to_sheets"] = True


if st.session_state.active_stage == 1:
    st.markdown(f"## {T('s1_title')}")
    st.markdown(f"""<div style="background:#1a2d45;border-radius:6px;padding:14px 18px;margin-bottom:16px;border-left:4px solid #2E75B6;">
<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;">{T('s1_what_label')}</div>
<div style="color:#e2e8f0;font-size:13px;margin-top:6px;">{T('s1_what')}</div>
<div style="color:#94a3b8;font-size:12px;margin-top:8px;">{T('s1_you_get')}</div>
</div>""", unsafe_allow_html=True)
    st.markdown("---")

    # Step 1 — Input
    if st.session_state.s1_step == 1:
        st.subheader(T("s1_step1"))

        ucol1, ucol2, ucol3 = st.columns(3)
        with ucol1:
            user_name = st.text_input(T("s1_name"), value=st.session_state.user_name, placeholder=T("s1_name_ph"))
        with ucol2:
            user_position = st.text_input(T("s1_role"), value=st.session_state.user_position, placeholder=T("s1_role_ph"))
        with ucol3:
            user_dept = st.text_input(T("s1_dept"), value=st.session_state.user_dept, placeholder=T("s1_dept_ph"))

        st.markdown("---")
        idea = st.text_area(T("s1_idea_label"), height=150, placeholder=T("s1_idea_ph"))

        if st.button(T("s1_submit"), type="primary"):
            if not user_name.strip() or not user_position.strip() or not user_dept.strip():
                st.warning(T("s1_warn_identity"))
            elif not idea.strip():
                st.warning(T("s1_warn_empty"))
            elif len(idea.split()) < 15:
                st.warning(T("s1_warn_brief"))
            else:
                # Save user identity
                st.session_state.user_name     = user_name.strip()
                st.session_state.user_position = user_position.strip()
                st.session_state.user_dept     = user_dept.strip()
                with st.spinner(T("s1_spinner_check")):
                    check = call_claude(
                        'Check if this innovation idea has enough detail to classify. Reply ONLY with JSON: {"sufficient": true/false, "missing": "one short sentence or empty string"}',
                        idea
                    )
                    try:
                        result = json.loads(check.strip().replace("```json","").replace("```","").strip())
                    except:
                        result = {"sufficient": True, "missing": ""}
                if not result.get("sufficient", True):
                    st.warning(f"A bit more detail needed: {result.get('missing','')}")
                else:
                    st.session_state.s1_idea = idea
                    # Check for similar past ideas
                    with st.spinner(T("s1_spinner_similar")):
                        past = load_past_ideas()
                        similar = check_similar_ideas(idea, past)
                        st.session_state.s1_similar_ideas = similar
                    st.session_state.s1_step = 2
                    st.rerun()

    # Step 2 — Forced-choice enrichment
    if st.session_state.s1_step == 2:
        st.markdown("---")
        st.subheader(T("s1_step2"))
        st.info(f"**Your idea:** {st.session_state.s1_idea}")

        # ── Similar ideas alert ───────────────────────────────
        similar = st.session_state.get("s1_similar_ideas", [])
        if similar:
            st.warning(T("s1_similar_warn").format(n=len(similar)))
            for s in similar:
                sim_col = "#f59e0b" if s.get("similarity") == "High" else "#60a5fa"
                st.markdown(f"""
<div style="background:#1a2d45;border-left:4px solid {sim_col};border-radius:4px;padding:12px 16px;margin:6px 0;">
  <div style="color:{sim_col};font-size:11px;font-weight:600;letter-spacing:1px;">{s.get('similarity','').upper()} SIMILARITY — {s.get('date','')} · {s.get('submitter','')} ({s.get('department','')})</div>
  <div style="color:#e2e8f0;font-size:13px;margin-top:4px;">{s.get('idea_snippet','')}</div>
  <div style="color:#94a3b8;font-size:12px;margin-top:4px;">Quadrant: <b>{s.get('quadrant','')}</b> · IPI: <b>{s.get('ipi','')}</b> · Outcome: <b>{s.get('recommendation','')}</b></div>
  <div style="color:#94a3b8;font-size:12px;margin-top:2px;">Similarity reason: {s.get('reason','')}</div>
</div>
""", unsafe_allow_html=True)
            st.markdown("---")
        st.caption(T("s1_step2_caption"))

        # Q1 — Technology novelty
        st.markdown(T("s1_q1"))
        q1_choice = st.radio("", [T("s1_q1a"), T("s1_q1b")], key="q1_radio", label_visibility="collapsed")
        q1_detail = st.text_input(T("s1_q1_detail"), key="q1_detail", placeholder=T("s1_q1_ph"))

        st.markdown("---")
        # Q2 — Market familiarity
        st.markdown(T("s1_q2"))
        st.caption(T("s1_q2_caption"))
        q2_choice = st.radio("", [T("s1_q2a"), T("s1_q2b")], key="q2_radio", label_visibility="collapsed")
        q2_detail = st.text_input(T("s1_q2_detail"), key="q2_detail", placeholder=T("s1_q2_ph"))

        st.markdown("---")
        # Q3 — Problem clarity
        st.markdown(T("s1_q3"))
        q3_choice = st.radio("", [T("s1_q3a"), T("s1_q3b")], key="q3_radio", label_visibility="collapsed")
        q3_detail = st.text_input(T("s1_q3_detail"), key="q3_detail", placeholder=T("s1_q3_ph"))

        if st.button(T("s1_classify_btn"), type="primary"):
            a1 = q1_choice + (f" — {q1_detail}" if q1_detail.strip() else "")
            a2 = q2_choice + (f" — {q2_detail}" if q2_detail.strip() else "")
            a3 = q3_choice + (f" — {q3_detail}" if q3_detail.strip() else "")
            st.session_state.s1_questions = [
                "Has the core technology been demonstrated anywhere?",
                "Does this idea target markets Schaeffler currently operates in?",
                "Is the problem this idea solves already recognised by industry?"
            ]
            st.session_state.s1_answers = [a1, a2, a3]
            st.session_state.s1_step = 3
            st.rerun()

    # Step 3 — Classify (70% idea · 30% Q&A weighted, fully deterministic)
    if st.session_state.s1_step == 3 and not st.session_state.s1_classification:
        with st.spinner(T("s1_spinner_classify")):
            q = st.session_state.s1_questions
            a = st.session_state.s1_answers

            # ── Pass 1: idea-only axis scoring (Q&A not shown to LLM) ─────────
            system_idea = """You are a senior Schaeffler innovation strategist.
Score this innovation idea on Schaeffler's Modified Innovation Matrix axes (Lau et al. ISPIM 2023).

AXIS DEFINITIONS (0-10 each):
Technology axis:
  0-2.5 = Established (commercially deployed, off-the-shelf components or standard industrial processes)
  2.5-5 = Adjacent (technology is commercially known but novel in this specific application)
  5-7.5 = New to Schaeffler (novel mechanism proven in research/lab but not commercialised at scale)
  7.5-10 = New to the World (theoretical or experimental only; no commercial deployment exists anywhere)

Market axis:
  0-2.5 = Established (core Schaeffler markets: automotive ICE/EV powertrains, industrial bearings, rail, conventional energy)
  2.5-5 = Adjacent (sectors Schaeffler partially serves: aerospace, two-wheelers, construction & agriculture, renewable energy)
  5-7.5 = New to Schaeffler (outside current scope: medical devices, consumer electronics, defence)
  7.5-10 = New to the World (entirely new market category with no established demand or business models)

Assign innovation metadata based on the idea content.

Return ONLY valid JSON:
{
  "tech_score": <float 0-10>,
  "market_score": <float 0-10>,
  "technology_level": "Established or Adjacent or New to Schaeffler or New to the World",
  "market_level": "Established or Adjacent or New to Schaeffler or New to the World",
  "technology_novelty": "one sentence describing what makes the technology novel or established",
  "market_position": "one sentence describing the target market relative to Schaeffler's current scope",
  "idea_reasoning": "2-3 sentences explaining the axis scores from the idea description alone",
  "confidence": "High/Medium/Low",
  "innovation_cluster": "one of: Energy Solutions, Material Solutions, Mobility Solutions, E-Drive Solutions, Robotics Solutions, Digital Solutions, Advanced Manufacturing, New Production Concepts — or empty string",
  "trend_alignment": ["1-2 of: Sustainability & Climate Change, New Mobility & Electrification, Autonomous Production, Data Economy & Digitalization, Demographic Change"],
  "product_family": "one of: Guide Motion, Transmit Motion, Control Motion, Generate Motion, Power Motion, Drive Motion, Energize Motion, Sustain Motion — or empty string",
  "project_type": "FIP or VEP or empty string",
  "innovation_model": "Integrated or Accelerator or empty string"
}"""
            try:
                raw_idea = call_claude(system_idea,
                    f"IDEA DESCRIPTION:\n{st.session_state.s1_idea}",
                    max_tokens=1000)
                raw_idea_c = raw_idea.strip().replace("```json","").replace("```","").strip()
                fbi = raw_idea_c.find("{"); lbi = raw_idea_c.rfind("}") + 1
                if fbi >= 0: raw_idea_c = raw_idea_c[fbi:lbi]
                idea_scores = json.loads(raw_idea_c)
            except Exception as e:
                st.error(f"Classification error: {e}")
                st.stop()

            idea_tech   = float(idea_scores.get("tech_score",   5.0))
            idea_market = float(idea_scores.get("market_score", 5.0))

            # ── Pass 2: Q&A → implied axis scores (deterministic, no LLM) ────
            # Q1 "Has the core technology been demonstrated anywhere?"
            #   Yes (demonstrated) → suggests established tech  → qa_tech = 3.5
            #   No  (novel)        → suggests new-to-world tech → qa_tech = 7.5
            qa_tech = 3.5 if a[0].startswith("Yes") else 7.5

            # Q2 "Does this target markets Schaeffler currently operates in?"
            #   Yes (existing)     → suggests established market → qa_market = 3.0
            #   No  (new)          → suggests new market         → qa_market = 7.5
            qa_market = 3.0 if a[1].startswith("Yes") else 7.5

            # Q3 "Is the problem already recognised by industry?" (minor secondary ±0.4)
            #   Yes (recognised problem) → slightly more established on both axes
            #   No  (novel problem)      → slightly more novel on both axes
            qa_adj = -0.4 if a[2].startswith("Yes") else 0.4

            # ── Weighted combination: 70% idea · 30% Q&A ─────────────────────
            # Q&A can only shift a borderline idea (scoring 4-6) across the threshold.
            # A clearly novel idea (score 8+) or clearly established idea (score 2-)
            # cannot be flipped by Q&A answers at 30% weight.
            final_tech   = round(min(9.5, max(0.5,
                idea_tech   * 0.70 + (qa_tech   + qa_adj) * 0.30)), 1)
            final_market = round(min(9.5, max(0.5,
                idea_market * 0.70 + (qa_market + qa_adj) * 0.30)), 1)

            # ── Axis label from final numeric score ───────────────────────────
            def _score_to_level(s):
                if s < 2.5: return "Established"
                if s < 5.0: return "Adjacent"
                if s < 7.5: return "New to Schaeffler"
                return "New to the World"

            # ── Quadrant mapping (Lau 2023 · OnePager definitions) ────────────
            # RADICAL    = New tech (≥5) + Established/Adjacent market (<5)
            # DISRUPTIVE = New tech (≥5) + New market (≥5)
            # EXTEND     = Established tech (<5) + New market (≥5)
            # EXPLOIT    = Established tech (<5) + Established market (<5)
            tech_new   = final_tech   >= 5.0
            market_new = final_market >= 5.0
            # Lau 2023 / Schaeffler paper definitions:
            # RADICAL    = new tech (≥5) + new market (≥5)       → top-right
            # DISRUPTIVE = new tech (≥5) + established market (<5) → bottom-right
            # EXTEND     = established tech (<5) + new market (≥5) → top-left
            # EXPLOIT    = established tech (<5) + established market (<5) → bottom-left
            if   tech_new and market_new:        q_result = "RADICAL"
            elif tech_new and not market_new:    q_result = "DISRUPTIVE"
            elif not tech_new and market_new:    q_result = "EXTEND"
            else:                                q_result = "EXPLOIT"

            # 3-state routing: PIPELINE / FUTURE_OPTIONS / PRODUCT_DIVISION
            if q_result in ("RADICAL", "DISRUPTIVE"):
                _cluster = idea_scores.get("innovation_cluster", "")
                # Future Options: borderline tech score (5.0–6.5, "New to Schaeffler" but not "New to World")
                # OR no innovation cluster assigned yet (idea needs more sensing before full pipeline)
                _borderline_tech = 5.0 <= final_tech <= 6.5
                if not _cluster or _borderline_tech:
                    route = "FUTURE_OPTIONS"
                else:
                    route = "PIPELINE"
            else:
                route = "PRODUCT_DIVISION"

            proceed = route == "PIPELINE"

            # ── Q&A impact note (shown in reasoning for transparency) ─────────
            tech_shift   = round(final_tech   - idea_tech,   1)
            market_shift = round(final_market - idea_market, 1)
            sign = lambda v: ("+" if v >= 0 else "") + str(v)
            qa_note = (
                f"Idea-only axis scores: tech {idea_tech}/10 · market {idea_market}/10. "
                f"Q&A signals (30% weight) adjusted tech to {final_tech} ({sign(tech_shift)}) "
                f"and market to {final_market} ({sign(market_shift)})."
            )

            # ── Assemble final classification dict ────────────────────────────
            classification = {
                "quadrant":            q_result,
                "confidence":          idea_scores.get("confidence", "Medium"),
                "technology_level":    _score_to_level(final_tech),
                "market_level":        _score_to_level(final_market),
                "tech_score_raw":      idea_tech,
                "market_score_raw":    idea_market,
                "tech_score_final":    final_tech,
                "market_score_final":  final_market,
                "technology_novelty":  idea_scores.get("technology_novelty", ""),
                "market_position":     idea_scores.get("market_position", ""),
                "reasoning":           idea_scores.get("idea_reasoning", "") + " " + qa_note,
                "route":               route,
                "proceed":             proceed,
                "schaeffler_division": "" if proceed else (
                    "E-Mobility or Vehicle Lifetime Solutions" if market_new
                    else "Bearings & Industrial Solutions or Powertrain & Chassis"
                ),
                "redirect_message": "" if proceed else (
                    f"This idea applies {'adjacent' if final_tech < 5 else 'established'} technology "
                    f"to {'new-to-Schaeffler' if market_new else 'existing'} markets — "
                    f"best evaluated by Schaeffler Product Development."
                ),
                "innovation_cluster": idea_scores.get("innovation_cluster", ""),
                "trend_alignment":    idea_scores.get("trend_alignment", []),
                "product_family":     idea_scores.get("product_family", ""),
                "project_type":       idea_scores.get("project_type", ""),
                "innovation_model":   idea_scores.get("innovation_model", ""),
                # pipeline_route populated from innovation_model so Stage 5/6 can branch on it
                "pipeline_route":     idea_scores.get("innovation_model", ""),
            }
            st.session_state.s1_classification = classification
        st.rerun()

    # Step 3 — Show result
    if st.session_state.s1_step == 3 and st.session_state.s1_classification:
        c = st.session_state.s1_classification
        quadrant = c.get("quadrant","")
        proceed  = c.get("proceed", False)
        route    = c.get("route", "PIPELINE" if proceed else "PRODUCT_DIVISION")

        st.markdown("---")
        st.subheader(T("s1_result"))
        st.info(f"**Your idea:** {st.session_state.s1_idea}")

        if route == "PRODUCT_DIVISION":
            division = c.get("schaeffler_division","Product Development")
            st.warning(f"**{quadrant}** — {c.get('redirect_message','')}")
            st.markdown(f"→ Suggested home: **{division}**")
            st.markdown(f"""
<div style="background:#1a2d45;border-radius:8px;padding:14px 18px;margin-top:12px;border-left:3px solid #f59e0b;">
<div style="color:#f59e0b;font-size:11px;font-weight:600;letter-spacing:1px;margin-bottom:6px;">WHY THIS IDEA DOESN'T ENTER THE INNOVATION PIPELINE</div>
<div style="color:#e2e8f0;font-size:13px;">
<b>EXPLOIT</b> and <b>EXTEND</b> ideas use established or adjacent technology — they belong in Schaeffler's Product Development divisions, not the Innovation Pipeline, because the core technology risk has already been resolved.<br><br>
The Innovation Pipeline (Stages 02–06) is reserved for <b>RADICAL</b> (breakthrough tech, new market) and <b>DISRUPTIVE</b> (breakthrough tech, existing market) ideas where the technology itself is genuinely novel and unproven.
</div>
<div style="color:#94a3b8;font-size:12px;margin-top:8px;">Technology level: <b>{c.get('technology_level','')}</b> · Market level: <b>{c.get('market_level','')}</b></div>
</div>
""", unsafe_allow_html=True)

        elif route == "FUTURE_OPTIONS":
            st.markdown(f"""
<div style="background:#0f1e35;border-radius:8px;padding:16px 20px;margin-top:8px;border-left:4px solid #818cf8;border:1px solid #818cf844;">
<div style="color:#818cf8;font-size:11px;font-weight:600;letter-spacing:1px;margin-bottom:8px;">⬡ FUTURE OPTIONS TRACK — SENSING PHASE</div>
<div style="color:#e2e8f0;font-size:14px;font-weight:600;margin-bottom:8px;">{quadrant} · Technology score {c.get('tech_score_final',5):.1f}/10 — {c.get('technology_level','')}</div>
<div style="color:#cbd5e1;font-size:13px;line-height:1.6;">
This idea shows genuine technological novelty but sits in the <b>"New to Schaeffler"</b> technology range (5.0–6.5) rather than "New to the World" (>6.5), or the innovation cluster is not yet clearly defined.
<br><br>
In Schaeffler's process, this means the idea enters the <b>sensing phase</b> — not the full Innovation Pipeline yet. The sensing phase involves: structured customer and technology observation, trend monitoring, and periodic reassessment as the technology matures.
</div>
<div style="color:#94a3b8;font-size:12px;margin-top:10px;">
Recommended next step: Submit to the <b>Innovation Cluster</b> lead for {c.get('innovation_cluster','the relevant cluster') or 'the relevant cluster'} for sensing-phase tracking.
<br>Technology level: <b>{c.get('technology_level','')}</b> · Market level: <b>{c.get('market_level','')}</b> · Quadrant: <b>{quadrant}</b>
</div>
</div>
""", unsafe_allow_html=True)

        else:
            # route == "PIPELINE"
            emoji = "🔬" if quadrant == "DISRUPTIVE" else "🚀"
            st.success(f"{emoji} **{quadrant}** — {c.get('reasoning','')}")
            st.caption(f"Confidence: {c.get('confidence','')}")

        tech_score, market_score = (
            c.get("tech_score_final"),
            c.get("market_score_final")
        )
        if tech_score is None or market_score is None:
            tech_score, market_score = get_dot_position(quadrant, c.get("confidence","Medium"))
        st.plotly_chart(ansoff_chart(quadrant, tech_score, market_score), use_container_width=True)

        # ── 4-level axis labels ───────────────────────────────
        if proceed:
            col_t, col_m = st.columns(2)
            col_t.markdown(f'<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:4px 0;"><div style="color:#94a3b8;font-size:11px;letter-spacing:1px;">TECHNOLOGY LEVEL</div><div style="color:#60a5fa;font-size:14px;font-weight:600;">{c.get("technology_level","")}</div></div>', unsafe_allow_html=True)
            col_m.markdown(f'<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:4px 0;"><div style="color:#94a3b8;font-size:11px;letter-spacing:1px;">MARKET LEVEL</div><div style="color:#60a5fa;font-size:14px;font-weight:600;">{c.get("market_level","")}</div></div>', unsafe_allow_html=True)
            col_a, col_b, col_c = st.columns(3)
            col_a.markdown(f'<div style="background:#1a2d45;border-radius:6px;padding:8px 12px;margin:4px 0;"><div style="color:#94a3b8;font-size:10px;letter-spacing:1px;">INNOVATION CLUSTER</div><div style="color:#e2e8f0;font-size:12px;font-weight:600;margin-top:2px;">{c.get("innovation_cluster","")}</div></div>', unsafe_allow_html=True)
            col_b.markdown(f'<div style="background:#1a2d45;border-radius:6px;padding:8px 12px;margin:4px 0;"><div style="color:#94a3b8;font-size:10px;letter-spacing:1px;">PRODUCT FAMILY</div><div style="color:#e2e8f0;font-size:12px;font-weight:600;margin-top:2px;">{c.get("product_family","")}</div></div>', unsafe_allow_html=True)
            col_c.markdown(f'<div style="background:#1a2d45;border-radius:6px;padding:8px 12px;margin:4px 0;"><div style="color:#94a3b8;font-size:10px;letter-spacing:1px;">TREND ALIGNMENT</div><div style="color:#e2e8f0;font-size:12px;margin-top:2px;">{", ".join(c.get("trend_alignment",[]))}</div></div>', unsafe_allow_html=True)
            pipeline = c.get("pipeline_route","")
            if pipeline:
                route_col = "#22c55e" if "Integrated" in pipeline else "#60a5fa"
                st.markdown(f'<div style="background:#0f1e35;border:1px solid {route_col}44;border-radius:6px;padding:8px 14px;margin:8px 0;"><span style="color:{route_col};font-size:12px;font-weight:600;">🔀 Pipeline route: {pipeline}</span></div>', unsafe_allow_html=True)

        col1, col2, col3 = st.columns(3)
        col1.metric("Technology axis", f"{tech_score} / 10")
        col2.metric("Market axis",     f"{market_score} / 10")
        col3.metric("Confidence",       c.get("confidence",""))

        # ── Scoring transparency: show idea vs Q&A contribution ───
        raw_t  = c.get("tech_score_raw",  tech_score)
        raw_m  = c.get("market_score_raw", market_score)
        if raw_t != tech_score or raw_m != market_score:
            sign = lambda v: ("+" if v >= 0 else "") + f"{v:.1f}"
            st.markdown(f"""
<div style="background:#0f1e35;border:1px solid #2a4a70;border-radius:6px;padding:10px 16px;margin:8px 0;">
  <div style="color:#94a3b8;font-size:10px;letter-spacing:1.2px;font-weight:600;margin-bottom:8px;">CLASSIFICATION WEIGHTING  ·  70% Idea · 30% Q&A</div>
  <div style="display:flex;gap:24px;flex-wrap:wrap;">
    <div>
      <div style="color:#64748b;font-size:10px;margin-bottom:2px;">TECHNOLOGY AXIS</div>
      <div style="color:#e2e8f0;font-size:12px;">Idea: <b style="color:#60a5fa;">{raw_t:.1f}</b> &nbsp;·&nbsp; Q&A adjustment: <b style="color:{'#f59e0b' if (tech_score-raw_t)!=0 else '#64748b'};">{sign(tech_score-raw_t)}</b> &nbsp;·&nbsp; Final: <b style="color:#22c55e;">{tech_score:.1f}</b></div>
    </div>
    <div>
      <div style="color:#64748b;font-size:10px;margin-bottom:2px;">MARKET AXIS</div>
      <div style="color:#e2e8f0;font-size:12px;">Idea: <b style="color:#60a5fa;">{raw_m:.1f}</b> &nbsp;·&nbsp; Q&A adjustment: <b style="color:{'#f59e0b' if (market_score-raw_m)!=0 else '#64748b'};">{sign(market_score-raw_m)}</b> &nbsp;·&nbsp; Final: <b style="color:#22c55e;">{market_score:.1f}</b></div>
    </div>
  </div>
</div>""", unsafe_allow_html=True)

        # Continue button — PIPELINE and FUTURE_OPTIONS both proceed; PRODUCT_DIVISION is the only hard stop
        st.markdown("---")
        if route in ("PIPELINE", "FUTURE_OPTIONS"):
            if route == "FUTURE_OPTIONS":
                st.info("⬡ **Sensing Phase** — This idea enters the research pipeline as a Future Option. Scores at Stage 06 will reflect sensing-phase maturity and frame graduation criteria rather than a standard investment verdict.")
            else:
                st.success(T("s1_qualifies"))
            btn_col1, btn_col2 = st.columns(2)
            with btn_col1:
                if st.button(T("s1_continue"), type="primary", key="s1_continue"):
                    st.session_state.active_stage = 2
                    st.rerun()
            with btn_col2:
                if st.button(T("s1_full_run"), type="secondary", key="s1_full_run"):
                    idea_fa = st.session_state.s1_idea
                    s1c_fa  = st.session_state.s1_classification
                    quad_fa = s1c_fa.get("quadrant","RADICAL")
                    _prog_fr = st.progress(0, text="⚡ Full pipeline starting…")
                    _stat_fr = st.empty()
                    _stat_fr.markdown("🔍 **Stage 02:** Market Intelligence…")
                    _prog_fr.progress(8)
                    run_stage2(idea_fa, quad_fa, s1c_fa)
                    _prog_fr.progress(25)
                    _stat_fr.markdown("🔬 **Stage 03:** Patent Intelligence…")
                    run_stage3(idea_fa, quad_fa, s1c_fa)
                    _prog_fr.progress(45)
                    _stat_fr.markdown("⚙️ **Stage 04:** Technical Feasibility…")
                    run_stage4(idea_fa, quad_fa, s1c_fa)
                    _prog_fr.progress(65)
                    _stat_fr.markdown("🏭 **Stage 05:** P³ Perspective…")
                    run_stage5(idea_fa, quad_fa, s1c_fa)
                    _prog_fr.progress(85)
                    _stat_fr.markdown("📊 **Stage 06:** Scoring & Synthesis…")
                    run_stage6_synthesis(idea_fa, quad_fa, s1c_fa)
                    _prog_fr.progress(100, text="✓ Full pipeline complete!")
                    time.sleep(0.4)
                    _prog_fr.empty()
                    _stat_fr.empty()
                    st.session_state.active_stage = 6
                    st.rerun()

        # Post-result chat
        st.markdown("---")
        st.subheader(T("s1_chat_header"))
        for msg in st.session_state.s1_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        user_q = st.chat_input(T("s1_chat_ph"))
        if user_q:
            st.session_state.s1_chat.append({"role":"user","content":user_q})
            with st.chat_message("user"):
                st.markdown(user_q)
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    context = f"""You are a senior Schaeffler innovation expert discussing a classification result.
Idea: {st.session_state.s1_idea}
Classification: {quadrant} (confidence: {c.get('confidence','')})
Reasoning: {c.get('reasoning','')}
Tech novelty: {c.get('technology_novelty','')}
Market position: {c.get('market_position','')}
{"Division: " + c.get('schaeffler_division','') if not proceed else "Qualifies for Innovation pipeline."}
Be specific and concise — 2-4 sentences. Reference Schaeffler's context (electrification, Vitesco, E-Mobility) where relevant."""
                    history = [{"role":m["role"],"content":m["content"]} for m in st.session_state.s1_chat]
                    reply = call_claude_chat(context, history)
                    st.markdown(reply)
                    st.session_state.s1_chat.append({"role":"assistant","content":reply})

        if st.button(T("s1_startover"), key="s1_startover"):
            # Reset Stage 01 inputs and identity
            for k in ["s1_step","s1_idea","s1_questions","s1_answers","s1_classification","s1_chat","s1_similar_ideas","user_name","user_position","user_dept"]:
                st.session_state[k] = defaults.get(k, "" if k in ("user_name","user_position","user_dept") else defaults.get(k))
            # Clear all downstream stages and save flag so new idea gets its own row
            for k in ["s2_data","s2_step","s2_chat",
                      "s3_data","s3_step","s3_chat",
                      "s4_data","s4_step","s4_chat",
                      "s5_data","s5_step","s5_chat",
                      "s6_data","s6_step","s6_chat","s6_action_brief",
                      "_s6_saved_to_sheets","s6_report_buf"]:
                if k in st.session_state:
                    del st.session_state[k]
            st.session_state.active_stage = 1
            st.rerun()

# ════════════════════════════════════════════════════════════
# STAGE 02 — MARKET INTELLIGENCE
# ════════════════════════════════════════════════════════════
elif st.session_state.active_stage == 2:
    st.markdown(f"## {T('s2_title')}")
    st.markdown(f"""<div style="background:#1a2d45;border-radius:6px;padding:14px 18px;margin-bottom:16px;border-left:4px solid #2E75B6;">
<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;">{T('stage_what_label')}</div>
<div style="color:#e2e8f0;font-size:13px;margin-top:6px;">{T('s2_what')}</div>
<div style="color:#94a3b8;font-size:12px;margin-top:8px;">{T('s2_you_get')}</div>
</div>""", unsafe_allow_html=True)
    st.markdown("---")

    idea     = st.session_state.s1_idea
    s1c      = st.session_state.s1_classification
    quadrant = s1c.get("quadrant","RADICAL")

    # Intro
    if st.session_state.s2_step == "intro":
        st.info(f"**Idea:** {idea}")
        st.markdown(f"**Quadrant:** {quadrant} · {s1c.get('technology_novelty','')}")
        if st.button(T("s2_run_btn"), type="primary"):
            st.session_state.s2_step = "running"
            st.rerun()

    # Running
    elif st.session_state.s2_step == "running":
        progress = st.progress(0)
        status   = st.empty()

        status.markdown("🔍 Gathering market data...")
        progress.progress(15)
        search_results = []
        if TAVILY_KEY:
            for q in [f"{idea} market size 2024", f"{idea} competitors", f"{idea} market growth"]:
                search_results.extend(tavily_search(q))
        web_ctx = ""
        if search_results:
            web_ctx = "\n\nWeb results:\n" + "\n".join(
                f"- [{r['title']}]({r['url']}): {r['content'][:300]}" for r in search_results[:6])

        status.markdown("📊 Analysing market size and growth...")
        progress.progress(35)
        system_market = """You are a senior market analyst. Analyse the market for this innovation idea.
RULES:
- Use only the most credible sources: McKinsey Global Institute, Gartner, Frost & Sullivan, BloombergNEF, IEA, Roland Berger, Statista, MarketsandMarkets, Grand View Research, Allied Market Research, Mordor Intelligence, Fortune Business Insights, IDC, Wood Mackenzie, S&P Global.
- For each market figure, provide structured source objects. Each source MUST have its own entry in the sources array — never combine two sources into one object.
- For URLs: provide the EXACT deep-link URL to the specific report page. If you do not know the exact URL, omit the url field entirely.
- market_size_current = most recent available year (2024 or 2025). market_size_forecast = 5-7 year projection. cagr = compound annual growth rate for that period.
- market_score: integer 1-10. Rubric: 9-10=large fast-growing (>$10bn, >15% CAGR); 7-8=strong ($2-10bn, 8-15%); 5-6=moderate; 3-4=niche; 1-2=tiny or declining.
Return ONLY valid JSON with NO extra text, NO comments inside the JSON:
{"market_name":"string",
"market_size_current":{"value":"$X.XB","year":"2024","sources":[{"org":"OrgName","title":"Report Title Year","year":"2024","url":"https://url-or-omit-field"}]},
"market_size_forecast":{"value":"$X.XB","year":"2030","sources":[{"org":"OrgName","title":"Report Title Year","year":"2024","url":"https://url-or-omit-field"}]},
"cagr":{"value":"X%","period":"2024-2030","sources":[{"org":"OrgName","title":"Report Title Year","year":"2024","url":"https://url-or-omit-field"}]},
"growth_drivers":["driver 1","driver 2","driver 3"],"market_maturity":"Emerging/Growing/Mature/Declining",
"geographic_focus":"string","market_score":7,"market_score_rationale":"2 sentences"}"""
        try:
            raw = call_claude(system_market, f"Idea: {idea}\nQuadrant: {quadrant}\nTech: {s1c.get('technology_novelty','')}{web_ctx}", max_tokens=1400)
            market = _parse_json(raw)
        except Exception:
            market = {"market_name":"N/A","market_size_current":{"value":"N/A","year":"2024","sources":[]},"market_size_forecast":{"value":"N/A","year":"2030","sources":[]},"cagr":{"value":"N/A","period":"","sources":[]},"growth_drivers":[],"market_maturity":"N/A","geographic_focus":"N/A","market_score":5,"market_score_rationale":""}

        status.markdown("🏢 Mapping competitive landscape...")
        progress.progress(55)
        system_comp = """You are a competitive intelligence analyst. Identify key competitors for this idea.
Every company must have a source. Return ONLY valid JSON with NO inline comments or annotations:
{"competitors":[{"name":"string","type":"Incumbent/Startup/Research","relevance":"one sentence","source":"Source: X, Y"}],
"competitive_intensity":"Low/Medium/High/Very High","white_space":"one sentence","schaeffler_advantage":"one sentence",
"competition_score":7,"competition_score_rationale":"2 sentences"}
Scoring guide (do NOT include in the JSON): 9-10=very open/few players; 7-8=some room; 5-6=moderate; 3-4=crowded; 1-2=saturated."""
        try:
            raw = call_claude(system_comp, f"Idea: {idea}\nMarket: {market.get('market_name','')}\nQuadrant: {quadrant}{web_ctx}")
            comp = _parse_json(raw)
        except Exception:
            comp = {"competitors":[],"competitive_intensity":"N/A","white_space":"N/A",
                    "schaeffler_advantage":"N/A","competition_score":5,"competition_score_rationale":""}

        status.markdown("🎯 Scoring Schaeffler sector clusters...")
        progress.progress(75)
        system_sectors = """You are a Schaeffler strategist. Score fit against Schaeffler's 10 sector clusters (0-10 each).
Clusters: Passenger Cars, Commercial Vehicles, Industrial Machinery, Rail, Aerospace, Two-Wheelers, Construction & Agriculture, Medical Equipment, Conventional Energy, Renewable Energy.
0-2=No relevance, 3-4=Low, 5-6=Moderate, 7-8=High, 9-10=Primary target.
Return ONLY valid JSON with NO inline comments:
{"sector_scores":{"Passenger Cars":{"score":5,"rationale":"one sentence"},"Commercial Vehicles":{"score":5,"rationale":"one sentence"},"Industrial Machinery":{"score":5,"rationale":"one sentence"},"Rail":{"score":5,"rationale":"one sentence"},"Aerospace":{"score":5,"rationale":"one sentence"},"Two-Wheelers":{"score":5,"rationale":"one sentence"},"Construction & Agriculture":{"score":5,"rationale":"one sentence"},"Medical Equipment":{"score":5,"rationale":"one sentence"},"Conventional Energy":{"score":5,"rationale":"one sentence"},"Renewable Energy":{"score":5,"rationale":"one sentence"}},
"primary_sectors":["sector name 1","sector name 2"],"sector_fit_score":7,"sector_fit_rationale":"2 sentences"}"""
        try:
            raw = call_claude(system_sectors, f"Idea: {idea}\nQuadrant: {quadrant}")
            sectors = _parse_json(raw)
        except Exception:
            sectors = {"sector_scores":{},"primary_sectors":[],"sector_fit_score":5,"sector_fit_rationale":""}

        # Final weighted score
        ms  = float(market.get("market_score",5))
        ss  = float(sectors.get("sector_fit_score",5))
        cs  = float(comp.get("competition_score",5))
        final = round(ms*0.40 + ss*0.35 + cs*0.25, 1)

        st.session_state.s2_data = {
            "market": market, "comp": comp, "sectors": sectors,
            "weights": {"Market Attractiveness":(ms,0.40),"Sector Fit":(ss,0.35),"Competition Opportunity":(cs,0.25)},
            "final_score": final, "web_results": search_results
        }
        progress.progress(100)
        status.markdown("✓ Complete.")
        time.sleep(0.5)
        st.session_state.s2_step = "done"
        st.rerun()

    # Results
    elif st.session_state.s2_step == "done":
        d       = st.session_state.s2_data
        market  = d["market"]; comp = d["comp"]; sectors = d["sectors"]
        weights = d["weights"]; final = d["final_score"]

        # ── Score banner ──────────────────────────────────────
        score_col = "#22c55e" if final>=7 else "#f59e0b" if final>=4 else "#ef4444"
        banner = f"""<div style="background:#0f1e35;border-radius:8px;padding:20px 24px;margin-bottom:20px;border:1px solid #2a4a70;"><div style="color:{WHITE};font-size:11px;letter-spacing:1.5px;opacity:0.5;margin-bottom:4px;">MARKET INTELLIGENCE SCORE</div><div style="color:{score_col};font-size:44px;font-weight:700;line-height:1;">{final}<span style="font-size:18px;color:#94a3b8;"> / 10</span></div><div style="color:{WHITE};font-size:13px;margin-top:6px;opacity:0.8;">{market.get('market_name','')}</div></div>"""
        st.markdown(banner, unsafe_allow_html=True)

        # ── Score breakdown ───────────────────────────────────
        cols = st.columns(3)
        for i,(label,(score,weight)) in enumerate(weights.items()):
            cols[i].metric(label, f"{score:.1f}/10", f"{int(weight*100)}% weight")
        st.markdown("---")

        # ── Market size ───────────────────────────────────────
        st.markdown("#### 📊 Market")

        import re as _re, urllib.parse as _up

        # Known org → stable publications page (fallback when no deep URL provided)
        _ORG_PAGES = {
            "mckinsey":             "https://www.mckinsey.com/mgi/research",
            "mckinsey global":      "https://www.mckinsey.com/mgi/research",
            "gartner":              "https://www.gartner.com/en/research/publications",
            "bloomberg":            "https://www.bloomberg.com/professional/insights/",
            "bloombergnef":         "https://about.bnef.com/insights/",
            "bnef":                 "https://about.bnef.com/insights/",
            "iea":                  "https://www.iea.org/reports",
            "statista":             "https://www.statista.com/markets/",
            "roland berger":        "https://www.rolandberger.com/en/Insights/Publications/",
            "frost & sullivan":     "https://store.frost.com/reports.html",
            "frost":                "https://store.frost.com/reports.html",
            "deloitte":             "https://www2.deloitte.com/global/en/insights.html",
            "pwc":                  "https://www.pwc.com/gx/en/industries/",
            "ihs markit":           "https://ihsmarkit.com/research-analysis/",
            "s&p global":           "https://www.spglobal.com/marketintelligence/en/news-insights/research",
            "wood mackenzie":       "https://www.woodmac.com/reports/",
            "mordor":               "https://www.mordorintelligence.com/industry-reports",
            "mordor intelligence":  "https://www.mordorintelligence.com/industry-reports",
            "grand view":           "https://www.grandviewresearch.com/industry-analysis",
            "grand view research":  "https://www.grandviewresearch.com/industry-analysis",
            "allied market":        "https://www.alliedmarketresearch.com/market-research-report",
            "allied market research": "https://www.alliedmarketresearch.com/market-research-report",
            "marketsandmarkets":    "https://www.marketsandmarkets.com/Market-Reports/",
            "fortune business":     "https://www.fortunebusinessinsights.com/reports",
            "fortune business insights": "https://www.fortunebusinessinsights.com/reports",
            "idc":                  "https://www.idc.com/research/viewtoc",
            "precedence":           "https://www.precedenceresearch.com/",
            "technavio":            "https://www.technavio.com/report-store",
            "ibisworld":            "https://www.ibisworld.com/global/",
            "euromonitor":          "https://www.euromonitor.com/reports",
        }

        def _resolve_url(src_obj):
            """
            Test Claude's URL with a HEAD request.
            Returns (url, is_direct) where is_direct=True means the link lands on the actual page.
            Falls back to a targeted Google Search if the URL fails or is missing.
            """
            import urllib.parse as _up2
            raw_url = src_obj.get("url", "").strip()

            # Build the Google Search fallback first — always works
            q_parts = [src_obj.get("org",""), src_obj.get("title",""), src_obj.get("year","")]
            q = " ".join(p for p in q_parts if p).strip()
            google_url = f"https://www.google.com/search?q={_up2.quote(q)}"

            if not raw_url or not raw_url.startswith("http"):
                # No URL provided — try org page first, then Google
                org_lower = src_obj.get("org", "").lower()
                for org_key, org_page in _ORG_PAGES.items():
                    if org_key in org_lower:
                        return org_page, False
                return google_url, False

            # Check the URL has a real path (not just homepage)
            try:
                from urllib.parse import urlparse as _ulp
                _parsed = _ulp(raw_url)
                path = _parsed.path.strip("/")
                if not path or len(path) <= 3:
                    return google_url, False
            except Exception:
                return google_url, False

            # Live HEAD check — 4 second timeout
            try:
                resp = requests.head(raw_url, timeout=4, allow_redirects=True,
                                     headers={"User-Agent": "Mozilla/5.0"})
                if resp.status_code < 400:
                    return raw_url, True
            except Exception:
                pass

            return google_url, False

        def _pill_label(src_obj):
            """Short display label for a source pill."""
            org = src_obj.get("org", "")
            year = src_obj.get("year", "")
            title = src_obj.get("title", "")
            if org and year:
                return f"{org}, {year}"
            elif org:
                return org
            elif title:
                return title[:35] + ("…" if len(title) > 35 else "")
            return "Source"

        def render_pills_structured(sources_list):
            """Return HTML for individual pill <a> tags from a list of source dicts.
            🔗 = direct link verified live. 🔍 = Google Search (guaranteed to open)."""
            if not sources_list:
                return '<span style="color:#4a6fa5;font-size:10px;">no source available</span>'
            html = ""
            for src in sources_list:
                url, is_direct = _resolve_url(src)
                lbl = _pill_label(src)
                icon = "🔗" if is_direct else "🔍"
                html += (
                    f'<a href="{url}" target="_blank" '
                    f'style="display:inline-block;background:#1e3a5f;color:#93c5fd;'
                    f'font-size:10px;padding:3px 10px;border-radius:12px;margin:3px 2px 0;'
                    f'text-decoration:none;border:1px solid #2a4a70;'
                    f'white-space:nowrap;line-height:1.5;">{icon} {lbl}</a>'
                )
            return html

        def _compat_field(field_val, fallback_key=None):
            """Handle both new structured format and legacy string format gracefully."""
            if isinstance(field_val, dict):
                return field_val
            # Legacy string format — wrap it
            raw = str(field_val) if field_val else "N/A"
            val = raw.split("[Source:")[0].strip().split("(")[0].split(";")[0].strip()
            return {"value": val, "year": "", "period": "", "sources": []}

        cur_obj  = _compat_field(market.get("market_size_current") or market.get("market_size_2024"))
        fore_obj = _compat_field(market.get("market_size_forecast") or market.get("market_size_2030"))
        cagr_obj = _compat_field(market.get("cagr"))

        # ── Big-number cards ──────────────────────────────────
        mc1, mc2, mc3 = st.columns(3)
        cards = [
            (mc1, "CURRENT MARKET SIZE",  cur_obj["value"],  cur_obj.get("year",""),         cur_obj.get("sources",[])),
            (mc2, "FORECAST MARKET SIZE", fore_obj["value"], fore_obj.get("year",""),         fore_obj.get("sources",[])),
            (mc3, "CAGR",                 cagr_obj["value"], cagr_obj.get("period",""),       cagr_obj.get("sources",[])),
        ]
        for col, top_label, num_val, sub_label, srcs in cards:
            pills_html = render_pills_structured(srcs)
            col.markdown(f"""
<div style="background:#1a2d45;border-radius:8px;padding:18px 14px 14px;text-align:center;min-height:140px;">
  <div style="color:#94a3b8;font-size:10px;letter-spacing:1.5px;font-weight:600;margin-bottom:6px;">{top_label}</div>
  <div style="color:#60a5fa;font-size:36px;font-weight:700;line-height:1.1;">{num_val}</div>
  <div style="color:#64748b;font-size:11px;margin-top:4px;margin-bottom:8px;">{sub_label}</div>
  <div style="margin-top:8px;line-height:2.0;">{pills_html}</div>
</div>""", unsafe_allow_html=True)

        st.markdown("<div style='margin-top:4px'></div>", unsafe_allow_html=True)
        st.caption("🔗 = direct link verified · 🔍 = Google Search to find the report (both always open)")

        col_l, col_r = st.columns(2)
        col_l.markdown(f"**Maturity** · {market.get('market_maturity','')}")
        col_r.markdown(f"**Geography** · {market.get('geographic_focus','')}")
        if market.get("growth_drivers"):
            drivers = market["growth_drivers"]
            show_n   = min(3, len(drivers))
            total_n  = len(drivers)
            drivers_html = "".join(
                f'<div style="display:flex;gap:10px;margin:6px 0;color:#e2e8f0;font-size:13px;'
                f'"><span style="color:#60a5fa;flex-shrink:0;">›</span><span>{drv}</span></div>'
                for drv in drivers[:3]
            )
            extra_html = '<div style="color:#64748b;font-size:11px;margin-top:6px;">Full list in the downloaded report.</div>' if total_n > 3 else ""
            st.markdown(
                '<div style="background:#1a2d45;border-radius:6px;padding:12px 16px;margin:8px 0 4px 0;">' +
                f'<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;margin-bottom:8px;">GROWTH DRIVERS &middot; TOP {show_n} OF {total_n}</div>' +
                drivers_html + extra_html +
                '</div>',
                unsafe_allow_html=True
            )
        st.markdown("---")

        # ── Sector cluster chart ──────────────────────────────
        st.markdown("#### 🎯 Schaeffler Sector Fit")
        sector_scores = sectors.get("sector_scores",{})
        primary = sectors.get("primary_sectors",[])
        if sector_scores:
            names  = list(sector_scores.keys())
            vals   = [sector_scores[s]["score"] for s in names]
            colors = [BLUE if s in primary else DIM for s in names]
            fig_bar = go.Figure(go.Bar(
                x=vals, y=names, orientation="h", marker_color=colors,
                text=[f"{v}/10" for v in vals], textposition="outside",
                textfont=dict(color=WHITE,size=11),
                hovertemplate="<b>%{y}</b><br>%{x}/10<extra></extra>"
            ))
            fig_bar.update_layout(
                plot_bgcolor=BG, paper_bgcolor=BG, height=340,
                xaxis=dict(range=[0,12],showgrid=False,zeroline=False,
                           tickfont=dict(color=WHITE),title_font=dict(color=DIM)),
                yaxis=dict(showgrid=False,zeroline=False,tickfont=dict(color=WHITE,size=11)),
                margin=dict(l=10,r=60,t=10,b=30), font=dict(color=WHITE)
            )
            st.plotly_chart(fig_bar, use_container_width=True)
            st.caption(f"Primary sectors: {', '.join(primary)}")
        st.markdown("---")

        # ── Competitive landscape — summary only ──────────────
        st.markdown("#### 🏢 Competitive Landscape")
        cc1,cc2,cc3 = st.columns(3)
        cc1.metric("Intensity", comp.get("competitive_intensity",""))
        cc2.metric("Competition Score", f"{comp.get('competition_score',0)}/10")
        cc3.metric("Key Players", len(comp.get("competitors",[])))
        st.markdown(f"**White space** · {comp.get('white_space','')}")
        st.markdown(f"**Schaeffler edge** · {comp.get('schaeffler_advantage','')}")
        competitors = comp.get("competitors", [])
        if competitors:
            with st.expander(f"View all {len(competitors)} key players"):
                type_cols = {"Incumbent":"#60a5fa","Startup":"#22c55e","Research":"#a78bfa"}
                for ci in competitors:
                    tc = type_cols.get(ci.get("type","Incumbent"), "#60a5fa")
                    st.markdown(f"""
<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:4px 0;display:flex;align-items:center;gap:12px;">
  <div style="color:#e2e8f0;font-weight:600;font-size:13px;min-width:150px;">{ci.get("name","")}</div>
  <div style="background:{tc}22;color:{tc};font-size:11px;padding:2px 8px;border-radius:10px;min-width:90px;text-align:center;">{ci.get("type","")}</div>
  <div style="color:#94a3b8;font-size:12px;flex:1;">{ci.get("relevance","")} <span style="color:#4a6fa5;font-size:11px;">{ci.get("source","")}</span></div>
</div>
""", unsafe_allow_html=True)


        # ── Download report ──────────────────────────────────
        st.markdown("---")
        _one_click_dl(
            T("dl_market"),
            lambda: generate_market_report(idea, quadrant, s1c, market, comp, sectors, weights, final),
            f"Schaeffler_Market_Intelligence_{datetime.now().strftime('%Y%m%d')}.docx"
        )

        # ── Chat ──────────────────────────────────────────────
        st.markdown("---")
        st.subheader(T("s2_chat_header"))
        for msg in st.session_state.s2_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        user_q = st.chat_input(T("s2_chat_ph"))
        if user_q:
            st.session_state.s2_chat.append({"role":"user","content":user_q})
            with st.chat_message("user"):
                st.markdown(user_q)
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    ctx = f"""You are a senior market analyst discussing market intelligence results for a Schaeffler innovation idea.
Idea: {idea} | Quadrant: {quadrant}
Market: {market.get('market_name','')} | Size 2024: {_mval(market.get('market_size_current') or market.get('market_size_2024',''))} | CAGR: {_mval(market.get('cagr',''))}
Competitive intensity: {comp.get('competitive_intensity','')} | White space: {comp.get('white_space','')}
Primary sectors (Schaeffler's 10 clusters): {', '.join(primary)} | Final score: {final}/10
Be specific, cite sources where possible, 3-4 sentences max."""
                    history = [{"role":m["role"],"content":m["content"]} for m in st.session_state.s2_chat]
                    reply = call_claude_chat(ctx, history)
                    st.markdown(reply)
                    st.session_state.s2_chat.append({"role":"assistant","content":reply})

        # ── Continue ──────────────────────────────────────────
        st.markdown("---")
        st.success(T("s2_success").format(score=final))
        if st.button(T("s2_continue"), type="primary", key="s2_continue"):
            st.session_state.active_stage = 3
            st.rerun()

        if st.button(T("s2_rerun"), key="s2_rerun"):
            st.session_state.s2_step = "intro"
            st.session_state.s2_data = {}
            st.session_state.s2_chat = []
            st.session_state.s2_report_buf = None
            st.rerun()





# ════════════════════════════════════════════════════════════
# STAGE 03 — PATENT INTELLIGENCE
# ════════════════════════════════════════════════════════════
elif st.session_state.active_stage == 3:
    st.markdown(f"## {T('s3_title')}")
    st.markdown(f"""<div style="background:#1a2d45;border-radius:6px;padding:14px 18px;margin-bottom:16px;border-left:4px solid #2E75B6;">
<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;">{T('stage_what_label')}</div>
<div style="color:#e2e8f0;font-size:13px;margin-top:6px;">{T('s3_what')}</div>
<div style="color:#94a3b8;font-size:12px;margin-top:8px;">{T('s3_you_get')}</div>
</div>""", unsafe_allow_html=True)
    st.markdown("---")

    idea     = st.session_state.s1_idea
    s1c      = st.session_state.s1_classification
    quadrant = s1c.get("quadrant", "RADICAL")

    # ── Session state for stage 03 ────────────────────────────
    for k, v in {"s3_step":"intro","s3_data":{},"s3_chat":[]}.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ── Intro ─────────────────────────────────────────────────
    if st.session_state.s3_step == "intro":
        st.info(f"**Idea:** {idea}")
        st.markdown(f"**Quadrant:** {quadrant}")
        if st.button(T("s3_run_btn"), type="primary"):
            st.session_state.s3_step = "running"
            st.rerun()

    # ── Running ───────────────────────────────────────────────
    elif st.session_state.s3_step == "running":
        progress = st.progress(0)
        status   = st.empty()

        status.markdown("🔍 Analysing external patent landscape...")
        progress.progress(20)

        system_landscape = """You are a patent intelligence analyst specialising in industrial technology.
Analyse the external patent landscape for this innovation idea.

RULES:
- Focus on COMPANIES filing patents, not individual patents
- Classify each company as: Competitor / Customer / Research Institution / Patent Troll / Adjacent Player
- Every claim must have [Source: org, year] where possible
- Be specific about technology sub-areas

Return ONLY valid JSON:
{
  "technology_keywords": ["3-5 key patent search terms for this idea"],
  "landscape_summary": "2-3 sentences on overall patent activity in this space",
  "activity_level": "Low / Moderate / High / Very High",
  "filing_trend": "Increasing / Stable / Decreasing",
  "filing_trend_rationale": "one sentence",
  "patent_landscape_score": integer 1-10 for landscape openness. Use: 9-10=very few filings, open territory; 7-8=some activity but clear gaps; 5-6=moderate filing density; 3-4=dense filing landscape; 1-2=saturated with patents by large incumbents,
  "key_filers": [
    {
      "company": "company name",
      "type": "Competitor / Customer / Research Institution / Adjacent Player",
      "focus": "one sentence on what they are patenting",
      "threat_level": "Low / Medium / High",
      "schaeffler_relationship": "Direct competitor / Potential customer / Partner / Unknown",
      "source": "Source: X, Year"
    }
  ],
  "white_spaces": ["white space opportunity 1", "white space opportunity 2", "white space opportunity 3"]
}"""

        try:
            raw = call_claude(system_landscape,
                f"Idea: {idea}\nQuadrant: {quadrant}\nTech novelty: {s1c.get('technology_novelty','')}", max_tokens=2000)
            landscape = _parse_json(raw)
        except Exception as e:
            landscape = {"technology_keywords":[],"landscape_summary":"Analysis unavailable.",
                        "activity_level":"N/A","filing_trend":"N/A","filing_trend_rationale":"",
                        "key_filers":[],"white_spaces":[],"patent_landscape_score":5}

        progress.progress(50)
        status.markdown("🏢 Mapping filers onto Schaeffler's Ansoff matrix...")

        system_ansoff = """You are a Schaeffler Group patent strategist.
Map patent filing companies onto Schaeffler's modified Ansoff matrix based on where their patents sit.

The matrix axes (MUST match Schaeffler's Stage 1 Innovation Framework):
- X axis: Technology Dimension (0=Existing/Established Technology → 10=New to the World)
- Y axis: Market Dimension (0=Existing/Established Market → 10=New to the World)

Quadrant positions (identical to Schaeffler's quadrant classifier):
- EXPLOIT  (bottom-left):  existing tech  + existing market  — x_score 0–5,  y_score 0–5
- EXTEND   (top-left):     existing tech  + new market       — x_score 0–5,  y_score 5–10
- RADICAL  (top-right):    new tech       + new market       — x_score 5–10, y_score 5–10
- DISRUPT  (bottom-right): new tech       + existing market  — x_score 5–10, y_score 0–5

For each company, assign:
- matrix_position: which quadrant their patent activity sits in
- x_score: 0-10 (0=existing technology, 10=new to the world technology)
- y_score: 0-10 (0=existing market, 10=new to the world market)

Also map where SCHAEFFLER's own known IP sits relative to this idea.

Return ONLY valid JSON:
{
  "filer_positions": [
    {
      "company": "company name",
      "matrix_position": "EXPLOIT/EXTEND/RADICAL/DISRUPT",
      "x_score": 0-10,
      "y_score": 0-10,
      "rationale": "one sentence"
    }
  ],
  "schaeffler_position": {
    "matrix_position": "EXPLOIT/EXTEND/RADICAL/DISRUPT",
    "x_score": 0-10,
    "y_score": 0-10,
    "existing_ip": "one sentence on what Schaeffler already has in this space",
    "gap": "one sentence on the IP gap this idea addresses"
  },
  "idea_position": {
    "x_score": 0-10,
    "y_score": 0-10
  },
  "novelty_signal": "Strong / Moderate / Weak",
  "novelty_rationale": "one sentence",
  "ip_risk": "Low / Medium / High",
  "ip_risk_rationale": "one sentence"
}"""

        # Pass FULL filer objects (not just names) so Claude can map every one
        key_filers = landscape.get("key_filers", [])
        filers_full_context = json.dumps([
            {"company": f.get("company",""), "type": f.get("type",""), "focus": f.get("focus",""), "threat_level": f.get("threat_level","")}
            for f in key_filers
        ])
        try:
            raw2 = call_claude(system_ansoff,
                f"Idea: {idea}\nQuadrant: {quadrant}\nTech keywords: {landscape.get('technology_keywords','')}\n"
                f"IMPORTANT: You MUST map ALL {len(key_filers)} filers listed below. Do not skip any.\n"
                f"Key filers (map every single one): {filers_full_context}",
                max_tokens=max(2500, len(key_filers) * 250 + 1200))
            ansoff_data = _parse_json(raw2)
        except Exception as e:
            ansoff_data = {"filer_positions":[],"schaeffler_position":{"matrix_position":"EXPLOIT","x_score":2,"y_score":2,"existing_ip":"N/A","gap":"N/A"},
                          "idea_position":{"x_score":7,"y_score":7},"novelty_signal":"Moderate","novelty_rationale":"","ip_risk":"Medium","ip_risk_rationale":""}

        # ── Guarantee every key_filer appears in filer_positions ──────────
        # Build a lookup of which companies already have positions
        positioned_companies = {fp.get("company","").lower() for fp in ansoff_data.get("filer_positions", [])}
        # Quadrant → default score ranges for auto-placement fallback
        # Quadrant → default score ranges (X=Technology, Y=Market — matches Stage 1 convention)
        # EXPLOIT=bottom-left(low x,low y), EXTEND=top-left(low x,high y),
        # RADICAL=top-right(high x,high y), DISRUPT=bottom-right(high x,low y)
        type_defaults = {
            "Competitor":          ("EXPLOIT", 3.0, 3.0),   # established tech + established market
            "Customer":            ("EXTEND",  2.5, 6.5),   # established tech + new market
            "Research Institution":("RADICAL", 7.0, 7.5),  # new tech + new market
            "Adjacent Player":     ("DISRUPT", 6.5, 3.5),  # new tech + established market
            "Patent Troll":        ("EXPLOIT", 2.0, 2.0),  # established tech + established market
        }
        import random
        for i, f in enumerate(key_filers):
            name = f.get("company","")
            if name.lower() not in positioned_companies and name:
                quad, bx, by = type_defaults.get(f.get("type","Adjacent Player"), ("EXPLOIT", 4.0, 4.0))
                # Small deterministic nudge so overlapping filers spread out
                nudge_x = ((i * 0.7) % 2.0) - 1.0
                nudge_y = ((i * 1.1) % 2.0) - 1.0
                ansoff_data.setdefault("filer_positions", []).append({
                    "company": name,
                    "type": f.get("type","Adjacent Player"),
                    "matrix_position": quad,
                    "x_score": round(min(9.5, max(0.5, bx + nudge_x)), 1),
                    "y_score": round(min(9.5, max(0.5, by + nudge_y)), 1),
                    "rationale": f.get("focus","Auto-placed based on filer type")
                })

        progress.progress(70)
        status.markdown("🔬 Assessing novelty signal...")

        # ── Novelty: separate domain expert Claude call ─────────────────────────
        filer_summary = ", ".join(
            f"{f.get('company','')} ({f.get('type','')})"
            for f in key_filers[:8]
        )
        system_novelty = """You are a domain expert academic assessing patent novelty for an innovation idea.
Evaluate the novelty signal based on the patent landscape and filer data provided.

Novelty Signal definitions (apply these strictly):
- High:   Blue ocean. Very few patent filers in this specific technology/market combination.
          The idea occupies largely uncontested IP territory.
- Medium: Oligopoly. A handful of established players dominate filings but the specific
          combination or application has meaningful differentiation potential.
- Low:    Red ocean. Many patent filers are active. Technology is well-covered by existing patents.

Return ONLY valid JSON:
{
  "novelty_signal": "High / Medium / Low",
  "novelty_rationale": "2-3 sentences from a domain expert perspective",
  "novelty_score": <integer 1-10>
}
Scoring guide:
9-10 = High — genuinely uncontested territory
7-8  = High-leaning — few filers, clear differentiation opportunity
5-6  = Medium — some prior art but meaningful gaps remain
3-4  = Medium-low — significant prior art, differentiation is challenging
1-2  = Low — red ocean, technology space is saturated"""

        novelty_ctx = (
            f"Idea: {idea}\nQuadrant: {quadrant}\n"
            f"Landscape summary: {landscape.get('landscape_summary','')}\n"
            f"Activity level: {landscape.get('activity_level','')}\n"
            f"Filing trend: {landscape.get('filing_trend','')}\n"
            f"Key filers: {filer_summary}"
        )
        try:
            raw_nov = call_claude(system_novelty, novelty_ctx, max_tokens=500)
            novelty_data = _parse_json_robust(raw_nov)
            if not novelty_data or not isinstance(novelty_data, dict):
                novelty_data = {"novelty_signal": "Medium", "novelty_rationale": "Assessment unavailable.", "novelty_score": 6}
        except Exception:
            novelty_data = {"novelty_signal": "Medium", "novelty_rationale": "Assessment unavailable.", "novelty_score": 6}
        novelty_score = float(novelty_data.get("novelty_score", 6))

        progress.progress(85)
        status.markdown("📊 Calculating patent intelligence score...")

        # ── Landscape score from LLM ────────────────────────────────────────────
        landscape_score = float(landscape.get("patent_landscape_score", 5))

        # ── IP Risk scores: deterministic from Ansoff matrix proximity ────────────
        filer_positions_r3 = ansoff_data.get("filer_positions", [])
        idea_x  = float(ansoff_data.get("idea_position", {}).get("x_score", 7.0))
        idea_y  = float(ansoff_data.get("idea_position", {}).get("y_score", 7.0))
        sch_x   = float(ansoff_data.get("schaeffler_position", {}).get("x_score", 3.0))
        sch_y   = float(ansoff_data.get("schaeffler_position", {}).get("y_score", 3.0))

        ip_idea_label, ip_idea_score, ip_idea_nearby           = _compute_ip_proximity_risk(filer_positions_r3, idea_x, idea_y)
        ip_sch_label,  ip_sch_score,  ip_sch_nearby            = _compute_ip_proximity_risk(filer_positions_r3, sch_x, sch_y)

        # ── Final score: 4 components, 25% each ────────────────────────────────
        final_patent = round((landscape_score + novelty_score + ip_idea_score + ip_sch_score) / 4, 1)

        st.session_state.s3_data = {
            "landscape":           landscape,
            "ansoff_data":         ansoff_data,
            "novelty_data":        novelty_data,
            "landscape_score":     landscape_score,
            "novelty_score":       novelty_score,
            "ip_idea_label":       ip_idea_label,
            "ip_idea_score":       ip_idea_score,
            "ip_idea_nearby":      ip_idea_nearby,
            "ip_schaeffler_label": ip_sch_label,
            "ip_schaeffler_score": ip_sch_score,
            "ip_sch_nearby":       ip_sch_nearby,
            "final_score":         final_patent,
        }

        progress.progress(100)
        status.markdown("✓ Complete.")
        time.sleep(0.5)
        st.session_state.s3_step = "done"
        st.rerun()

    # ── Results ───────────────────────────────────────────────
    elif st.session_state.s3_step == "done":
        d            = st.session_state.s3_data
        # Guard: if cached data is old 3-component format, reset and re-run
        if "ip_idea_score" not in d:
            st.session_state.s3_step = "intro"
            st.session_state.s3_data = {}
            st.rerun()
        landscape    = d["landscape"]
        ansoff_data  = d["ansoff_data"]
        final        = d["final_score"]

        # Score banner
        score_col = "#22c55e" if final>=7 else "#f59e0b" if final>=4 else "#ef4444"
        st.markdown(f"""
<div style="background:#0f1e35;border-radius:8px;padding:20px 24px;margin-bottom:20px;border:1px solid #2a4a70;">
  <div style="color:#94a3b8;font-size:9px;letter-spacing:2px;font-weight:700;font-family:Arial,sans-serif;margin-bottom:4px;">PATENT INTELLIGENCE SCORE</div>
  <div style="color:{score_col};font-size:42px;font-weight:700;line-height:1;">{final}<span style="font-size:18px;color:#94a3b8;"> / 10</span></div>
  <div style="color:{WHITE};font-size:13px;margin-top:6px;opacity:0.8;">{landscape.get("landscape_summary","")}</div>
</div>
""", unsafe_allow_html=True)

        # Score breakdown — 4 components 25% each
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Landscape Openness",  f"{d['landscape_score']:.1f} / 10", "25% weight")
        col2.metric("Novelty Signal",       f"{d['novelty_score']:.1f} / 10",  "25% weight")
        col3.metric("IP Risk — Idea",       f"{d['ip_idea_score']:.1f} / 10",
                    f"25% · {d.get('ip_idea_label','?')} ({d.get('ip_idea_nearby',0)} nearby filers)")
        col4.metric("IP Risk — Schaeffler", f"{d['ip_schaeffler_score']:.1f} / 10",
                    f"25% · {d.get('ip_schaeffler_label','?')} ({d.get('ip_sch_nearby',0)} nearby filers)")
        st.markdown("---")

        # ── Patent activity overview ──────────────────────────
        st.markdown("#### 📋 Patent Activity")
        c1, c2 = st.columns(2)
        c1.metric("Filing Activity", landscape.get("activity_level",""))
        c2.metric("Trend", landscape.get("filing_trend",""))
        st.caption(landscape.get("filing_trend_rationale",""))

        if landscape.get("technology_keywords"):
            kws = "  ·  ".join([f"`{k}`" for k in landscape["technology_keywords"]])
            st.markdown(f"**Search terms:** {kws}")
        st.markdown("---")

        # ── Ansoff matrix with all filers plotted ─────────────
        st.markdown("#### 🗺️ Patent Position Map — Schaeffler Ansoff Matrix")

        # ── Pattern commentary — generated from the filer data ──
        filer_positions = ansoff_data.get("filer_positions", [])
        schaeffler_pos  = ansoff_data.get("schaeffler_position", {})
        idea_pos        = ansoff_data.get("idea_position", {"x_score":7,"y_score":7})

        # Enrich filer_positions with 'type' from key_filers where missing
        key_filers_lookup = {f.get("company","").lower(): f for f in landscape.get("key_filers",[])}
        for fp in filer_positions:
            if not fp.get("type"):
                kf = key_filers_lookup.get(fp.get("company","").lower(), {})
                fp["type"] = kf.get("type", "Adjacent Player")

        # Nudge overlapping points apart so labels don't stack
        used_positions = []
        for fp in filer_positions:
            x, y = fp.get("x_score", 5), fp.get("y_score", 5)
            attempts = 0
            while any(abs(x-ux) < 0.6 and abs(y-uy) < 0.6 for ux,uy in used_positions) and attempts < 8:
                x = round(min(9.5, max(0.5, x + 0.4)), 1)
                y = round(min(9.5, max(0.5, y + 0.3)), 1)
                attempts += 1
            fp["x_score"] = x
            fp["y_score"] = y
            used_positions.append((x, y))

        # Compute pattern stats for commentary
        if filer_positions:
            idea_x = idea_pos.get("x_score", 7)
            idea_y = idea_pos.get("y_score", 7)
            # Quadrant distribution
            q_counts = {"EXPLOIT":0,"EXTEND":0,"RADICAL":0,"DISRUPT":0}
            close_filers = []
            for fp in filer_positions:
                q_counts[fp.get("matrix_position","EXPLOIT")] = q_counts.get(fp.get("matrix_position","EXPLOIT"),0) + 1
                # Close = within 2 units on both axes
                dx = abs(fp.get("x_score",5) - idea_x)
                dy = abs(fp.get("y_score",5) - idea_y)
                if dx <= 2.5 and dy <= 2.5:
                    close_filers.append(fp.get("company",""))

            dominant_q  = max(q_counts, key=q_counts.get)
            dominant_n  = q_counts[dominant_q]
            total_filers = len(filer_positions)
            competitors_count = sum(1 for fp in filer_positions if fp.get("type")=="Competitor")
            research_count    = sum(1 for fp in filer_positions if fp.get("type")=="Research Institution")

            # Build pointer bullets
            pointers = []
            if close_filers:
                pointers.append(f"**{len(close_filers)} filer{'s' if len(close_filers)>1 else ''} sitting close to your idea** — {', '.join(close_filers[:3])}{'...' if len(close_filers)>3 else ''}. Check IP risk before external disclosure.")
            else:
                pointers.append(f"**No filers plotted close to your idea's position** — the immediate IP zone appears uncontested.")
            if dominant_n > total_filers * 0.5 and dominant_q != quadrant:
                pointers.append(f"**Most activity is in the {dominant_q} quadrant** ({dominant_n} of {total_filers} filers) — your idea targets a less contested zone.")
            if competitors_count >= 3:
                pointers.append(f"**{competitors_count} direct competitors** identified — conduct freedom-to-operate analysis before committing R&D budget.")
            elif competitors_count == 0:
                pointers.append("**No direct competitors mapped** — validate this through a formal patent search before relying on it.")
            if research_count >= 2:
                pointers.append(f"**{research_count} research institutions** active in this space — potential co-development or licensing partners.")
            sch_x = schaeffler_pos.get("x_score", 2)
            sch_y = schaeffler_pos.get("y_score", 2)
            gap_x = abs(idea_x - sch_x)
            gap_y = abs(idea_y - sch_y)
            if gap_x > 3 or gap_y > 3:
                pointers.append(f"**Schaeffler's existing IP sits far from this idea** — significant new IP investment likely required to protect the concept.")
            else:
                pointers.append(f"**Schaeffler's existing IP is adjacent to this idea** — existing patents may provide partial protection or a foundation to build from.")

            # Render commentary box — convert **markdown bold** to <b> for HTML context
            import re as _re
            def md_bold_to_html(text):
                return _re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            bullets_html = "".join(f'<div style="margin:5px 0;color:#e2e8f0;font-size:13px;">› {md_bold_to_html(p)}</div>' for p in pointers)
            st.markdown(f"""
<div style="background:#1a2d45;border-left:3px solid #2E75B6;border-radius:4px;padding:12px 16px;margin-bottom:14px;">
<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;margin-bottom:8px;">PATTERN NOTES</div>
{bullets_html}
</div>
""", unsafe_allow_html=True)

        st.caption("Each point = a company's patent filing position. Schaeffler's existing IP shown in green. Your idea shown in orange.")

        fig = go.Figure()

        # Quadrant shading — X=Technology, Y=Market (matches Stage 1 Ansoff convention)
        # EXPLOIT=bottom-left, EXTEND=top-left, RADICAL=top-right, DISRUPT=bottom-right
        q_fills = [
            dict(x=[0,5,5,0],   y=[0,0,5,5],   name="EXPLOIT", fill="#1a2d45", lx=2.5,ly=2.5),
            dict(x=[0,5,5,0],   y=[5,5,10,10], name="EXTEND",  fill="#1e3a5f", lx=2.5,ly=7.5),
            dict(x=[5,10,10,5], y=[5,5,10,10], name="RADICAL", fill="#1F3864", lx=7.5,ly=7.5),
            dict(x=[5,10,10,5], y=[0,0,5,5],   name="DISRUPT", fill="#0d2137", lx=7.5,ly=2.5),
        ]
        for q in q_fills:
            fig.add_trace(go.Scatter(
                x=q["x"]+[q["x"][0]], y=q["y"]+[q["y"][0]],
                fill="toself", fillcolor=q["fill"],
                line=dict(color="#2a4a70",width=1),
                mode="lines", showlegend=False, hoverinfo="skip"
            ))
            fig.add_annotation(x=q["lx"],y=q["ly"],text=f"<b>{q['name']}</b>",
                showarrow=False, font=dict(size=12,color="#4a6fa5"))

        # Grid lines
        fig.add_shape(type="line",x0=5,x1=5,y0=0,y1=10,line=dict(color="#2a4a70",width=1.5,dash="dot"))
        fig.add_shape(type="line",x0=0,x1=10,y0=5,y1=5,line=dict(color="#2a4a70",width=1.5,dash="dot"))

        # Competitor/filer points
        type_colours = {
            "Competitor":         "#ef4444",
            "Customer":           "#60a5fa",
            "Research Institution":"#a78bfa",
            "Adjacent Player":    "#f59e0b",
            "Patent Troll":       "#6b7280",
        }
        for fp in filer_positions:
            col = type_colours.get(fp.get("type","Adjacent Player"), "#6b7280")
            fig.add_trace(go.Scatter(
                x=[fp.get("x_score",5)], y=[fp.get("y_score",5)],
                mode="markers+text",
                marker=dict(size=12, color=col, line=dict(color="white",width=1.5)),
                text=[f"  {fp.get('company','')}"],
                textposition="middle right",
                textfont=dict(size=10, color="#e2e8f0"),
                name=fp.get("type",""),
                hovertemplate=f"<b>{fp.get('company','')}</b><br>{fp.get('type','')}<br>{fp.get('rationale','')}<extra></extra>",
                showlegend=False
            ))

        # Schaeffler existing IP — GREEN diamond
        if schaeffler_pos:
            fig.add_trace(go.Scatter(
                x=[schaeffler_pos.get("x_score",2)], y=[schaeffler_pos.get("y_score",2)],
                mode="markers+text",
                marker=dict(size=16, color="#22c55e", symbol="diamond",
                           line=dict(color="white",width=2)),
                text=["  Schaeffler IP"],
                textposition="middle right",
                textfont=dict(size=11,color="#22c55e",family="Arial Bold"),
                showlegend=False,
                hovertemplate=f"<b>Schaeffler existing IP</b><br>{schaeffler_pos.get('existing_ip','')}<extra></extra>"
            ))

        # This idea — ORANGE star
        fig.add_trace(go.Scatter(
            x=[idea_pos.get("x_score",7)], y=[idea_pos.get("y_score",7)],
            mode="markers+text",
            marker=dict(size=18, color="#f97316", symbol="star",
                       line=dict(color="white",width=2)),
            text=["  Your idea"],
            textposition="middle right",
            textfont=dict(size=12,color="#f97316",family="Arial Bold"),
            showlegend=False,
            hovertemplate="<b>Your innovation idea</b><extra></extra>"
        ))

        # Axis labels
        for ann in [
            dict(x=2.5,y=-0.9,text="Established Tech",angle=0),
            dict(x=7.5,y=-0.9,text="New Tech",angle=0),
            dict(x=-1.1,y=2.5,text="Established Market",angle=-90),
            dict(x=-1.1,y=7.5,text="New Market",angle=-90),
        ]:
            fig.add_annotation(x=ann["x"],y=ann["y"],text=ann["text"],
                showarrow=False,font=dict(size=10,color="#e2e8f0"),textangle=ann["angle"])

        fig.update_layout(
            plot_bgcolor=BG, paper_bgcolor=BG, height=460,
            xaxis=dict(range=[-1.5,12],showticklabels=False,showgrid=False,zeroline=False,
                      title="Technology Dimension →",title_font=dict(size=11,color=DIM)),
            yaxis=dict(range=[-1.5,11],showticklabels=False,showgrid=False,zeroline=False,
                      title="Market Dimension →",title_font=dict(size=11,color=DIM)),
            margin=dict(l=70,r=20,t=20,b=55), font=dict(color=WHITE)
        )
        st.plotly_chart(fig, use_container_width=True)

        # Legend
        legend_items = [("🔴","Competitor"),("🔵","Customer"),("🟣","Research Institution"),("🟡","Adjacent Player"),("🟢","Schaeffler existing IP"),("🟠","Your idea")]
        st.markdown("  ".join([f"{e} {l}" for e,l in legend_items]))
        st.markdown("---")

        # ── Key filers summary ────────────────────────────────
        st.markdown("#### 🏢 Key Patent Filers")
        filers = landscape.get("key_filers",[])
        if filers:
            top_filers = sorted(filers, key=lambda x: {"High":0,"Medium":1,"Low":2}.get(x.get("threat_level","Medium"),1))[:5]
            if len(filers) > 5:
                st.caption(f"Showing top 5 of {len(filers)} filers by threat level. Full list in the downloaded report.")
            for fi in top_filers:
                threat_col = {"High":"#ef4444","Medium":"#f59e0b","Low":"#22c55e"}.get(fi.get("threat_level","Medium"),"#6b7280")
                type_col   = type_colours.get(fi.get("type","Adjacent Player"),"#6b7280")
                st.markdown(f"""
<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:5px 0;display:flex;align-items:center;gap:12px;">
  <div style="min-width:130px;color:{WHITE};font-weight:600;font-size:13px;">{fi.get('company','')}</div>
  <div style="background:{type_col}22;color:{type_col};font-size:11px;padding:2px 8px;border-radius:10px;min-width:120px;text-align:center;">{fi.get('type','')}</div>
  <div style="background:{threat_col}22;color:{threat_col};font-size:11px;padding:2px 8px;border-radius:10px;min-width:80px;text-align:center;">⚡ {fi.get('threat_level','')} threat</div>
  <div style="color:#94a3b8;font-size:12px;flex:1;">{fi.get('focus','')}</div>
</div>
""", unsafe_allow_html=True)

        st.markdown("---")

        # ── Schaeffler IP position ────────────────────────────
        st.markdown("#### 🏭 Schaeffler IP Position")
        sc1, sc2, sc3 = st.columns(3)
        sc1.metric("Novelty Signal",
                   d.get("novelty_data", {}).get("novelty_signal", "—"),
                   f"Score: {d.get('novelty_score', 0):.1f}/10")
        sc2.metric("IP Risk — Idea",
                   d.get("ip_idea_label", "—"),
                   f"{d.get('ip_idea_nearby', 0)} filers within 2 units")
        sc3.metric("IP Risk — Schaeffler",
                   d.get("ip_schaeffler_label", "—"),
                   f"{d.get('ip_sch_nearby', 0)} filers within 2 units")
        st.caption(f"Novelty rationale: {d.get('novelty_data', {}).get('novelty_rationale', '')}")
        st.caption(f"Existing Schaeffler IP: {schaeffler_pos.get('existing_ip', '')}")
        st.caption(f"IP gap this idea addresses: {schaeffler_pos.get('gap', '')}")

        # White spaces
        white_spaces = landscape.get("white_spaces",[])
        if white_spaces:
            st.markdown("**IP white spaces identified:**")
            for ws in white_spaces[:3]:
                st.markdown(f"- {ws}")
            if len(white_spaces) > 3:
                st.caption(f"+{len(white_spaces)-3} more in the downloaded report.")

        # ── Download report ───────────────────────────────────
        st.markdown("---")
        _one_click_dl(
            T("dl_patent"),
            lambda: generate_patent_report(idea, quadrant, s1c, landscape, ansoff_data, d),
            f"Schaeffler_Patent_Intelligence_{datetime.now().strftime('%Y%m%d')}.docx"
        )

        # ── Chat ──────────────────────────────────────────────
        st.markdown("---")
        st.subheader(T("s3_chat_header"))
        for msg in st.session_state.s3_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        user_q = st.chat_input(T("s3_chat_ph"))
        if user_q:
            st.session_state.s3_chat.append({"role":"user","content":user_q})
            with st.chat_message("user"):
                st.markdown(user_q)
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    ctx = f"""You are a Schaeffler patent intelligence expert.
Idea: {idea} | Quadrant: {quadrant}
Patent activity: {landscape.get('activity_level','')} | Trend: {landscape.get('filing_trend','')}
Novelty signal: {ansoff_data.get('novelty_signal','')} | IP risk: {ansoff_data.get('ip_risk','')}
White spaces: {', '.join(landscape.get('white_spaces',[]))}
Schaeffler IP gap: {schaeffler_pos.get('gap','')}
Be specific and concise — 2-4 sentences."""
                    history = [{"role":m["role"],"content":m["content"]} for m in st.session_state.s3_chat]
                    reply = call_claude_chat(ctx, history)
                    st.markdown(reply)
                    st.session_state.s3_chat.append({"role":"assistant","content":reply})

        # ── Continue ──────────────────────────────────────────
        st.markdown("---")
        st.success(T("s3_success").format(score=final))
        if st.button(T("s3_continue"), type="primary", key="s3_continue"):
            st.session_state.active_stage = 4
            st.rerun()

        if st.button(T("s3_rerun"), key="s3_rerun"):
            st.session_state.s3_step = "intro"
            st.session_state.s3_data = {}
            st.session_state.s3_chat = []
            st.session_state.s3_report_buf = None
            st.rerun()



# ════════════════════════════════════════════════════════════
# STAGE 04 — TECHNICAL FEASIBILITY
# ════════════════════════════════════════════════════════════
elif st.session_state.active_stage == 4:
    st.markdown(f"## {T('s4_title')}")
    st.markdown(f"""<div style="background:#1a2d45;border-radius:6px;padding:14px 18px;margin-bottom:16px;border-left:4px solid #2E75B6;">
<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;">{T('stage_what_label')}</div>
<div style="color:#e2e8f0;font-size:13px;margin-top:6px;">{T('s4_what')}</div>
<div style="color:#94a3b8;font-size:12px;margin-top:8px;">{T('s4_you_get')}</div>
</div>""", unsafe_allow_html=True)
    st.markdown("---")

    idea     = st.session_state.s1_idea
    s1c      = st.session_state.s1_classification
    quadrant = s1c.get("quadrant", "RADICAL")

    for k, v in {"s4_step":"intro","s4_data":{},"s4_chat":[]}.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ── Intro ─────────────────────────────────────────────────
    if st.session_state.s4_step == "intro":
        st.info(f"**Idea:** {idea}")
        st.markdown(f"**Quadrant:** {quadrant}")
        if st.button(T("s4_run_btn"), type="primary"):
            st.session_state.s4_step = "running"
            st.rerun()

    # ── Running ───────────────────────────────────────────────
    elif st.session_state.s4_step == "running":
        progress = st.progress(0)
        status   = st.empty()

        # ── Step A: Existence check & evidence ───────────────
        status.markdown("🔬 Checking for technology demonstrations...")
        progress.progress(20)

        system_existence = """You are a technology intelligence analyst specialising in industrial and automotive R&D.
Check whether the core technology behind this innovation idea has been demonstrated anywhere.

Look for evidence in: academic research, university labs, startup products, government programmes, 
industry pilots, defence/aerospace applications (often ahead of industrial), and adjacent industries.

RULES:
- Every evidence item MUST include a source citation [Source: org/journal/publication, year]
- Only cite sources you are confident exist — do not fabricate
- Be specific about what was demonstrated vs what is still theoretical
- Rate confidence in each evidence item: High / Medium / Low

Return ONLY valid JSON:
{
  "technology_core": "one sentence describing the core technology mechanism",
  "existence_verdict": "Demonstrated / Partially Demonstrated / Research Stage / Theoretical",
  "existence_summary": "2-3 sentences on the state of the art",
  "evidence": [
    {
      "type": "Academic Paper / Startup / Pilot / Government Programme / Industry Deployment",
      "title": "name of the paper, company, or programme",
      "description": "one sentence on what was demonstrated",
      "source": "Source: org/journal, year",
      "confidence": "High / Medium / Low",
      "relevance": "Direct / Adjacent / Analogous"
    }
  ],
  "technology_gaps": ["gap 1 between current state and full deployment", "gap 2", "gap 3"],
  "time_to_readiness": "estimated years to production readiness",
  "keywords": ["6-10 key technical terms from this domain for keyword map"]
}"""

        try:
            raw = call_claude(system_existence,
                f"Idea: {idea}\nQuadrant: {quadrant}\nTech: {s1c.get('technology_novelty','')}", max_tokens=3000)
            existence = _parse_json_robust(raw)
            if not existence or not isinstance(existence, dict):
                raise ValueError("Parse failed")
        except Exception as e:
            st.warning(f"Evidence parsing issue: {e} — using fallback")
            existence = {"technology_core":"N/A","existence_verdict":"Research Stage","existence_summary":"",
                        "evidence":[],"technology_gaps":[],"time_to_readiness":"Estimated 5–8 years (fallback — re-run for accurate estimate)","keywords":[]}

        progress.progress(50)

        # ── Step B: Schaeffler-adapted TRL assessment ─────────
        status.markdown("📊 Assessing technology readiness level...")
        progress.progress(65)

        system_trl = """You are a Schaeffler R&D director assessing technology maturity.
Use the Schaeffler-adapted TRL framework (modified from NASA TRL for industrial/automotive context):

TRL 1 — Basic principles observed (theoretical concept only)
TRL 2 — Technology concept formulated (application identified, no testing)
TRL 3 — Experimental proof of concept (lab demonstration, key functions validated)
TRL 4 — Technology validated in lab (component/subsystem tested in controlled environment)
TRL 5 — Technology validated in relevant environment (prototype tested in industrial-like conditions)
TRL 6 — Technology demonstrated in relevant environment (system prototype demonstrated)
TRL 7 — System prototype demonstrated in operational environment (field trial or pilot)
TRL 8 — System complete and qualified (full production design, limited production run)
TRL 9 — Actual system proven in operational environment (commercial deployment at scale)

For Schaeffler's innovation pipeline, ideas typically enter at TRL 3-5.
TRL 1-2 = too early for Innovation; TRL 6+ = should be in Product Development.

Return ONLY valid JSON:
{
  "trl_level": 1-9,
  "trl_label": "TRL X — label from framework above",
  "trl_rationale": "2-3 sentences justifying this TRL rating based on evidence",
  "schaeffler_entry_readiness": "Too Early / Ready for Innovation / Ready for Product Development",
  "entry_rationale": "one sentence",
  "key_technical_risks": [
    {"risk": "technical risk description", "severity": "High/Medium/Low", "mitigation": "one sentence"},
    {"risk": "technical risk description", "severity": "High/Medium/Low", "mitigation": "one sentence"},
    {"risk": "technical risk description", "severity": "High/Medium/Low", "mitigation": "one sentence"}
  ],
  "analogous_schaeffler_technologies": "one sentence on which of Schaeffler's 8 Motion Product Families (Guide Motion/Transmit Motion/Control Motion/Generate Motion/Power Motion/Drive Motion/Energize Motion/Sustain Motion) this technology is closest to",
  "trl_score": integer 1-10 mapped directly from TRL level. TRL1=1, TRL2=2, TRL3=3.5, TRL4=5, TRL5=6, TRL6=7, TRL7=8, TRL8=9, TRL9=10
}"""

        try:
            raw2 = call_claude(system_trl,
                f"Idea: {idea}\nExistence verdict: {existence.get('existence_verdict','')}\nEvidence: {json.dumps(existence.get('evidence',[])[:3])}\nGaps: {existence.get('technology_gaps',[])}",
                max_tokens=1500)
            raw2_clean = raw2.strip().replace("```json","").replace("```","").strip()
            first_brace = raw2_clean.find("{")
            last_brace  = raw2_clean.rfind("}") + 1
            if first_brace >= 0:
                raw2_clean = raw2_clean[first_brace:last_brace]
            trl = json.loads(raw2_clean)
        except Exception as e:
            st.warning(f"TRL parsing issue: {e} — using fallback")
            trl = {"trl_level":3,"trl_label":"TRL 3 — Experimental proof of concept",
                  "trl_rationale":"","schaeffler_entry_readiness":"Ready for Innovation",
                  "entry_rationale":"","key_technical_risks":[],"analogous_schaeffler_technologies":"","trl_score":5}

        progress.progress(85)
        status.markdown("📐 Calculating feasibility score...")

        # Feasibility score: TRL score (50%) + existence quality (30%) + risk profile (20%)
        trl_score = float(trl.get("trl_score", 5))
        existence_map = {"Demonstrated":9,"Partially Demonstrated":6,"Research Stage":3,"Theoretical":1}
        existence_score = existence_map.get(existence.get("existence_verdict","Research Stage"), 5)
        risk_scores = [{"High":2,"Medium":5,"Low":8}.get(r.get("severity","Medium"),5)
                      for r in trl.get("key_technical_risks",[])]
        risk_score = sum(risk_scores)/len(risk_scores) if risk_scores else 5.0
        final_feasibility = round(trl_score*0.50 + existence_score*0.30 + risk_score*0.20, 1)

        st.session_state.s4_data = {
            "existence": existence,
            "trl": trl,
            "trl_score": trl_score,
            "existence_score": existence_score,
            "risk_score": round(risk_score,1),
            "final_score": final_feasibility
        }

        progress.progress(100)
        status.markdown("✓ Complete.")
        time.sleep(0.5)
        st.session_state.s4_step = "done"
        st.rerun()

    # ── Results ───────────────────────────────────────────────
    elif st.session_state.s4_step == "done":
        d         = st.session_state.s4_data
        existence = d["existence"]
        trl       = d["trl"]
        final     = d["final_score"]

        # ── Score banner ──────────────────────────────────────
        score_col = "#22c55e" if final>=7 else "#f59e0b" if final>=4 else "#ef4444"
        trl_level = trl.get("trl_level", 3)
        trl_pct   = int((trl_level / 9) * 100)
        entry_col = {"Too Early":"#ef4444","Ready for Innovation":"#22c55e",
                     "Ready for Product Development":"#60a5fa"}.get(
                     trl.get("schaeffler_entry_readiness",""), "#f59e0b")

        st.markdown(f"""
<div style="background:#0f1e35;border-radius:8px;padding:20px 24px;margin-bottom:20px;border:1px solid #2a4a70;">
  <div style="color:#94a3b8;font-size:9px;letter-spacing:2px;font-weight:700;font-family:Arial,sans-serif;margin-bottom:4px;">TECHNICAL FEASIBILITY SCORE</div>
  <div style="display:flex;align-items:flex-end;gap:24px;">
    <div style="color:{score_col};font-size:42px;font-weight:700;line-height:1;">{final}<span style="font-size:18px;color:#94a3b8;"> / 10</span></div>
    <div>
      <div style="color:{entry_col};font-size:14px;font-weight:600;">{trl.get("schaeffler_entry_readiness","")}</div>
      <div style="color:{WHITE};font-size:12px;opacity:0.7;">{existence.get("existence_verdict","")}</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

        # Score breakdown
        col1, col2, col3 = st.columns(3)
        col1.metric("TRL Score",        f"{d['trl_score']:.1f} / 10", "33% weight")
        col2.metric("Existence Quality", f"{d['existence_score']:.1f} / 10", "33% weight")
        col3.metric("Risk Profile",      f"{d['risk_score']:.1f} / 10", "33% weight")
        st.markdown("---")

        # ── TRL gauge ─────────────────────────────────────────
        st.markdown("#### 🎯 Technology Readiness Level")
        trl_colours = {
            1:"#ef4444",2:"#ef4444",3:"#f97316",
            4:"#f59e0b",5:"#eab308",6:"#84cc16",
            7:"#22c55e",8:"#10b981",9:"#06b6d4"
        }
        trl_col = trl_colours.get(trl_level,"#f59e0b")

        # TRL bar
        st.markdown(f"""
<div style="margin:12px 0;">
  <div style="display:flex;justify-content:space-between;margin-bottom:4px;">
    <span style="color:{WHITE};font-weight:600;font-size:15px;">{trl.get("trl_label","")}</span>
    <span style="color:{trl_col};font-weight:700;font-size:18px;">TRL {trl_level}/9</span>
  </div>
  <div style="background:#1a2d45;border-radius:6px;height:14px;overflow:hidden;">
    <div style="background:{trl_col};height:100%;width:{trl_pct}%;border-radius:6px;transition:width 0.3s;"></div>
  </div>
  <div style="display:flex;justify-content:space-between;margin-top:4px;">
    <span style="color:#4a6fa5;font-size:10px;">TRL 1 — Theoretical</span>
    <span style="color:#4a6fa5;font-size:10px;">TRL 9 — Commercial</span>
  </div>
</div>
""", unsafe_allow_html=True)

        st.caption(trl.get("trl_rationale",""))
        st.caption(f"Entry rationale: {trl.get('entry_rationale','')}")
        if trl.get("analogous_schaeffler_technologies") or trl.get("analogous_schaeffler_tech"):
            st.caption(f"Schaeffler analogous experience: {trl.get('analogous_schaeffler_technologies') or trl.get('analogous_schaeffler_tech','')}")
        st.markdown("---")

        # ── TRL scale reference ───────────────────────────────
        trl_descriptions = [
            (1,"Basic principles observed","Theoretical concept only, no testing"),
            (2,"Technology concept formulated","Application identified, no experimental testing"),
            (3,"Experimental proof of concept","Lab demonstration, key functions validated"),
            (4,"Technology validated in lab","Component tested in controlled environment"),
            (5,"Validated in relevant environment","Prototype tested in industrial-like conditions"),
            (6,"Demonstrated in relevant environment","System prototype demonstrated"),
            (7,"Prototype in operational environment","Field trial or industrial pilot"),
            (8,"System complete and qualified","Full production design, limited production run"),
            (9,"Proven in operational environment","Commercial deployment at scale"),
        ]
        current_entry = next(((lvl,lbl,dsc) for lvl,lbl,dsc in trl_descriptions if lvl == trl_level), None)
        if current_entry:
            lvl_c, lbl_c, dsc_c = current_entry
            st.markdown(f"""
<div style="background:#1F3864;border-radius:4px;padding:8px 14px;margin:4px 0;display:flex;gap:12px;align-items:center;border:1px solid #2E75B6;">
  <div style="color:{trl_colours.get(lvl_c,'#f59e0b')};font-weight:700;min-width:40px;">TRL {lvl_c}</div>
  <div style="color:#e2e8f0;font-size:13px;font-weight:600;min-width:220px;">▶ {lbl_c}</div>
  <div style="color:#60a5fa;font-size:12px;">{dsc_c}</div>
</div>
""", unsafe_allow_html=True)
        with st.expander("View full TRL scale reference"):
            for lvl, label, desc in trl_descriptions:
                is_current = (lvl == trl_level)
                bg_col = "#1F3864" if is_current else "#1a2d45"
                text_col_inner = "#60a5fa" if is_current else "#94a3b8"
                st.markdown(f"""
<div style="background:{bg_col};border-radius:4px;padding:6px 12px;margin:3px 0;display:flex;gap:12px;align-items:center;">
  <div style="color:{trl_colours.get(lvl,'#f59e0b')};font-weight:700;min-width:40px;">TRL {lvl}</div>
  <div style="color:{'#e2e8f0' if is_current else '#e2e8f0'};font-size:13px;min-width:220px;">{"▶ " if is_current else ""}{label}</div>
  <div style="color:{text_col_inner};font-size:12px;">{desc}</div>
</div>
""", unsafe_allow_html=True)

        st.markdown("---")

        # ── Evidence ──────────────────────────────────────────
        st.markdown("#### 📚 Existing Evidence")
        st.caption(existence.get("existence_summary",""))

        evidence_items = existence.get("evidence",[])
        if evidence_items:
            type_icons = {
                "Academic Paper":"📄","Startup":"🚀","Pilot":"🔧",
                "Government Programme":"🏛️","Industry Deployment":"🏭"
            }
            rel_cols = {"Direct":"#22c55e","Adjacent":"#60a5fa","Analogous":"#f59e0b"}
            conf_cols = {"High":"#22c55e","Medium":"#f59e0b","Low":"#ef4444"}
            # Sort: Direct > Adjacent > Analogous, then High > Medium > Low confidence
            rel_order  = {"Direct":0,"Adjacent":1,"Analogous":2}
            conf_order = {"High":0,"Medium":1,"Low":2}
            sorted_ev = sorted(evidence_items,
                key=lambda x: (rel_order.get(x.get("relevance","Adjacent"),1),
                               conf_order.get(x.get("confidence","Medium"),1)))
            top_ev = sorted_ev[:5]
            if len(evidence_items) > 5:
                st.caption(f"Showing top 5 of {len(evidence_items)} evidence items by relevance. Full list in the downloaded report.")
            for ev in top_ev:
                icon = type_icons.get(ev.get("type",""), "📌")
                rel_col  = rel_cols.get(ev.get("relevance","Adjacent"), "#60a5fa")
                conf_col = conf_cols.get(ev.get("confidence","Medium"), "#f59e0b")
                st.markdown(f"""
<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:5px 0;">
  <div style="display:flex;align-items:center;gap:8px;margin-bottom:4px;">
    <span style="font-size:14px;">{icon}</span>
    <span style="color:{WHITE};font-weight:600;font-size:13px;">{ev.get("title","")}</span>
    <span style="background:{rel_col}22;color:{rel_col};font-size:10px;padding:2px 7px;border-radius:8px;">{ev.get("relevance","")}</span>
    <span style="background:{conf_col}22;color:{conf_col};font-size:10px;padding:2px 7px;border-radius:8px;">{ev.get("confidence","")} confidence</span>
  </div>
  <div style="color:#cbd5e1;font-size:12px;">{ev.get("description","")} <span style="color:#4a6fa5;">{ev.get("source","")}</span></div>
</div>
""", unsafe_allow_html=True)

            # ── Research paper search links ────────────────────────────────────
            import urllib.parse as _up4
            with st.expander(f"🔗 Search links for {len(top_ev)} evidence items"):
                st.caption("Links open targeted searches — paper titles from LLM training memory, verify before citing.")
                for ev in top_ev:
                    title  = ev.get("title", "")
                    source = ev.get("source", "")
                    etype  = ev.get("type", "")
                    if not title:
                        continue
                    import re as _re4
                    source_clean = _re4.sub(r',?\s*\(?\d{4}(?:[–\-]\d{4})?\)?', '', source).strip().rstrip(',').strip()
                    q_gs   = _up4.quote(f'{title} {source_clean}')
                    q_ss   = _up4.quote(f'{title} {source_clean}')
                    url_gs = f"https://scholar.google.com/scholar?q={q_gs}"
                    url_ss = f"https://www.semanticscholar.org/search?q={q_ss}&sort=Relevance"
                    url_pub = f"https://pubmed.ncbi.nlm.nih.gov/?term={_up4.quote(title)}" if etype == "Academic Paper" else ""
                    links = f'[Google Scholar]({url_gs})  ·  [Semantic Scholar]({url_ss})'
                    if url_pub:
                        links += f'  ·  [PubMed]({url_pub})'
                    st.markdown(f"**{title}**  \nSource: {source_clean}  \n{links}")
                    st.markdown("---")

        st.markdown("---")

        # ── Keyword map ───────────────────────────────────────
        st.markdown("#### 🗝️ Technology Keyword Map")
        st.caption("Key technical terms from this domain — indicative of academic and research weight.")
        keywords = existence.get("keywords", [])
        if keywords:
            # Size keywords by estimated relevance (first = most relevant)
            sizes = [28, 24, 22, 20, 18, 16, 15, 14, 13, 12]
            colours_kw = ["#60a5fa","#34d399","#a78bfa","#f59e0b","#f472b6",
                          "#60a5fa","#34d399","#a78bfa","#f59e0b","#f472b6"]
            badges = ""
            for i, kw in enumerate(keywords[:10]):
                sz  = sizes[i] if i < len(sizes) else 12
                col = colours_kw[i % len(colours_kw)]
                badges += f'<span style="background:{col}22;color:{col};font-size:{sz}px;padding:6px 14px;border-radius:20px;margin:4px;display:inline-block;font-weight:600;">{kw}</span>'
            st.markdown(f'<div style="background:#0f1e35;border-radius:8px;padding:16px;line-height:2.2;border:1px solid #2a4a70;">{badges}</div>',
                        unsafe_allow_html=True)
        st.markdown("---")

        # ── Technical risks ───────────────────────────────────
        st.markdown("#### ⚠️ Key Technical Risks")
        risks = trl.get("key_technical_risks", [])
        if risks:
            top_risks = sorted(risks, key=lambda x: {"High":0,"Medium":1,"Low":2}.get(x.get("severity","Medium"),1))[:5]
            if len(risks) > 5:
                st.caption(f"Showing top 5 of {len(risks)} risks by severity. Full list in the downloaded report.")
            for risk in top_risks:
                sev_col = {"High":"#ef4444","Medium":"#f59e0b","Low":"#22c55e"}.get(risk.get("severity","Medium"),"#f59e0b")
                st.markdown(f"""
<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:5px 0;border-left:3px solid {sev_col};">
  <div style="display:flex;align-items:center;gap:8px;margin-bottom:3px;">
    <span style="background:{sev_col}22;color:{sev_col};font-size:10px;padding:2px 8px;border-radius:8px;">{risk.get("severity","")} severity</span>
    <span style="color:{WHITE};font-size:13px;">{risk.get("risk","")}</span>
  </div>
  <div style="color:#94a3b8;font-size:12px;">↳ Mitigation: {risk.get("mitigation","")}</div>
</div>
""", unsafe_allow_html=True)

        if existence.get("technology_gaps"):
            gaps = existence["technology_gaps"]
            st.markdown("**Technology gaps to bridge:**")
            for gap in gaps[:5]:
                st.markdown(f"- {gap}")
            if len(gaps) > 5:
                st.caption(f"+{len(gaps)-5} more in the downloaded report.")

        st.caption(f"Estimated time to production readiness: **{existence.get('time_to_readiness','')}**")

        # ── Download report ───────────────────────────────────
        st.markdown("---")
        _one_click_dl(
            T("dl_feasib"),
            lambda: generate_feasibility_report(idea, quadrant, s1c, existence, trl, d),
            f"Schaeffler_Technical_Feasibility_{datetime.now().strftime('%Y%m%d')}.docx"
        )

        # ── Chat ──────────────────────────────────────────────
        st.markdown("---")
        st.subheader(T("s4_chat_header"))
        for msg in st.session_state.s4_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        user_q = st.chat_input(T("s4_chat_ph"))
        if user_q:
            st.session_state.s4_chat.append({"role":"user","content":user_q})
            with st.chat_message("user"):
                st.markdown(user_q)
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    ctx = f"""You are a Schaeffler R&D director discussing technical feasibility results.
Idea: {idea} | Quadrant: {quadrant}
TRL: {trl.get("trl_level","")} — {trl.get("trl_label","")}
Existence: {existence.get("existence_verdict","")}
Entry readiness: {trl.get("schaeffler_entry_readiness","")}
Time to readiness: {existence.get("time_to_readiness","")}
Key gaps: {existence.get("technology_gaps",[])}
Be specific, reference evidence where relevant, keep to 3-4 sentences."""
                    history = [{"role":m["role"],"content":m["content"]} for m in st.session_state.s4_chat]
                    reply = call_claude_chat(ctx, history)
                    st.markdown(reply)
                    st.session_state.s4_chat.append({"role":"assistant","content":reply})

        # ── Continue ──────────────────────────────────────────
        st.markdown("---")
        st.success(T("s4_success").format(score=final))
        if st.button(T("s4_continue"), type="primary", key="s4_continue"):
            st.session_state.active_stage = 5
            st.rerun()

        if st.button(T("s4_rerun"), key="s4_rerun"):
            st.session_state.s4_step = "intro"
            st.session_state.s4_data = {}
            st.session_state.s4_chat = []
            st.session_state.s4_report_buf = None
            st.rerun()



# ════════════════════════════════════════════════════════════
# STAGE 05 — ORGANISATIONAL READINESS
# ════════════════════════════════════════════════════════════
elif st.session_state.active_stage == 5:
    st.markdown(f"## {T('s5_title')}")
    st.markdown(f"""<div style="background:#1a2d45;border-radius:6px;padding:14px 18px;margin-bottom:16px;border-left:4px solid #2E75B6;">
<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;">{T('stage_what_label')}</div>
<div style="color:#e2e8f0;font-size:13px;margin-top:6px;">{T('s5_what')}</div>
<div style="color:#94a3b8;font-size:12px;margin-top:8px;">{T('s5_you_get')}</div>
</div>""", unsafe_allow_html=True)
    st.markdown("---")

    idea     = st.session_state.s1_idea
    s1c      = st.session_state.s1_classification
    quadrant = s1c.get("quadrant", "RADICAL")

    for k, v in {"s5_step":"intro","s5_data":{},"s5_chat":[]}.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ── Intro ─────────────────────────────────────────────────
    if st.session_state.s5_step == "intro":
        st.info(f"**Idea:** {idea}")
        st.markdown(f"**Quadrant:** {quadrant}")
        # Pull context from previous stages for enriched prompt
        s3_landscape = st.session_state.get("s3_data",{}).get("landscape",{})
        s4_existence = st.session_state.get("s4_data",{}).get("existence",{})
        if st.button(T("s5_run_btn"), type="primary"):
            st.session_state.s5_step = "running"
            st.rerun()

    # ── Running ───────────────────────────────────────────────
    elif st.session_state.s5_step == "running":
        progress = st.progress(0)
        status   = st.empty()

        # ── Safe variable preparation — guard against None stored by Claude ──
        s3_data_i    = st.session_state.get("s3_data") or {}
        s4_data_i    = st.session_state.get("s4_data") or {}
        s3_land_i    = s3_data_i.get("landscape") or {}
        s4_exist_i   = s4_data_i.get("existence") or {}
        s4_trl_i     = s4_data_i.get("trl") or {}

        raw_filers_i  = s3_land_i.get("key_filers") or []
        raw_sources_i = s4_exist_i.get("evidence") or []
        prior_filers_i          = [str(f.get("company","")) for f in raw_filers_i  if isinstance(f, dict)]
        prior_evidence_sources_i= [str(e.get("source","")) for e in raw_sources_i if isinstance(e, dict)]

        trl_level_i          = int(s4_trl_i.get("trl_level") or 3)
        innovation_cluster_i = str(s1c.get("innovation_cluster") or "")
        product_family_i     = str(s1c.get("product_family") or "")
        raw_trends_i         = s1c.get("trend_alignment")
        trend_alignment_i    = list(raw_trends_i) if isinstance(raw_trends_i, (list, tuple)) else []
        innovation_model_i   = str(s1c.get("innovation_model") or "Integrated")
        pipeline_route_i     = str(s1c.get("pipeline_route") or innovation_model_i)

        filers_str_i  = ", ".join(prior_filers_i[:6])            or "none identified"
        sources_str_i = ", ".join(prior_evidence_sources_i[:5])  or "none identified"
        trends_str_i  = ", ".join(str(t) for t in trend_alignment_i) or "none"

        status.markdown("🏭 Assessing Schaeffler P³ readiness...")
        progress.progress(25)

        system_readiness_i = (
            "You are a senior Schaeffler Group innovation strategist assessing internal P³ Perspective.\n"
            "Schaeffler P³ formula: Performance = Portfolio x People x Process.\n"
            "- Portfolio: Does this idea fit Schaeffler strategic portfolio and innovation clusters?\n"
            "- People: Does Schaeffler have the skills and teams to develop this?\n"
            "- Process: Does Schaeffler have the processes, infrastructure, and assets to execute?\n\n"
            f"Innovation cluster: {innovation_cluster_i} | Product family: {product_family_i} | "
            f"Trends: {trends_str_i} | Current TRL: {trl_level_i}\n"
            f"Patent filers (Stage 03): {filers_str_i}\n"
            f"Evidence sources (Stage 04): {sources_str_i}\n"
            f"Innovation model: {innovation_model_i} | Pipeline route: {pipeline_route_i}\n"
            "Schaeffler competencies: precision bearings, mechatronics, power electronics (Vitesco), "
            "tribology, EV drivetrains, industrial automation, embedded sensors, ASPICE/ISO 26262, "
            "OEM Tier 1 supply chain, 41 R&D centres globally.\n\n"
            "- RADICAL/Integrated: assess readiness for FIP-VEP-PEP process, OEM partnerships, internal R&D.\n"
            "- DISRUPTIVE/Accelerator: assess readiness for VC/startup track; internal gaps are expected.\n\n"
            "Return ONLY valid JSON. Use real integer scores 1-10 and real descriptive strings:\n"
            '{"p3_portfolio":{"score":7,"rationale":"two sentences on portfolio fit","cluster_fit":"one sentence",'
            '"strengths":["strength 1","strength 2"],"gaps":["gap 1"]},'
            '"p3_people":{"score":6,"rationale":"two sentences on competency readiness",'
            '"matched_competencies":["comp 1","comp 2","comp 3"],'
            '"competency_gap":"the single critical missing competency","sourcing_route":"how to close the gap"},'
            '"p3_process":{"score":6,"rationale":"two sentences on process readiness",'
            '"applicable_assets":["asset 1","asset 2"],"investment_required":"what needs building","time_to_close":"12 months"},'
            '"partnership_candidates":[{"name":"org name","type":"Startup","rationale":"why them","route":"Co-develop"}],'
            '"org_gaps":[{"gap":"gap name","severity":"High","closure_route":"how to close","timeline":"6 months"}],'
            '"build_or_partner":{"recommendation":"Co-develop","rationale":"two to three sentences",'
            '"time_to_trl6_internal":"30 months","time_to_trl6_partner":"18 months"},'
            '"p3_perspective_score":6}'
        )

        _p3_fallback_i = {
            "p3_portfolio":{"score":5,"rationale":"Analysis unavailable — re-run Stage 05.","cluster_fit":"N/A","strengths":[],"gaps":[]},
            "p3_people":{"score":5,"rationale":"Analysis unavailable — re-run Stage 05.","matched_competencies":[],"competency_gap":"N/A","sourcing_route":"N/A"},
            "p3_process":{"score":5,"rationale":"Analysis unavailable — re-run Stage 05.","applicable_assets":[],"investment_required":"N/A","time_to_close":"N/A"},
            "partnership_candidates":[],"org_gaps":[],
            "build_or_partner":{"recommendation":"Co-develop","rationale":"Analysis unavailable.","time_to_trl6_internal":"N/A","time_to_trl6_partner":"N/A"},
            "p3_perspective_score":5
        }

        try:
            raw = call_claude(system_readiness_i,
                f"Innovation idea: {idea}\nQuadrant: {quadrant}\nTRL level: {trl_level_i}",
                max_tokens=2500)
            org_data = _parse_json(raw)
            if "p3_portfolio" not in org_data or "p3_people" not in org_data:
                raise ValueError("Missing required P3 keys")
        except Exception as e:
            status.markdown("⚠️ Retrying P³ analysis...")
            try:
                raw_retry = call_claude(
                    "Return ONLY valid JSON. No markdown, no commentary. "
                    "Required keys: p3_portfolio, p3_people, p3_process, "
                    "partnership_candidates, org_gaps, build_or_partner, p3_perspective_score. "
                    "Each score must be an integer 1-10.",
                    f"Idea: {idea}\nQuadrant: {quadrant}\nTRL: {trl_level_i}\n"
                    f"Cluster: {innovation_cluster_i}\nModel: {innovation_model_i}\n"
                    "Assess Schaeffler internal P3 readiness (Portfolio, People, Process).",
                    max_tokens=2500
                )
                org_data = _parse_json(raw_retry)
            except Exception as e2:
                st.error(f"P³ analysis failed ({e2}). Use the individual Stage 05 button to re-run.")
                org_data = _p3_fallback_i

        progress.progress(75)
        status.markdown("🔍 Identifying partnership candidates...")

        # Equal weight: Portfolio 33%, People 33%, Process 34%
        p_portfolio = float(org_data.get("p3_portfolio",{}).get("score",5))
        p_people    = float(org_data.get("p3_people",{}).get("score",5))
        p_process   = float(org_data.get("p3_process",{}).get("score",5))
        final_org   = round((p_portfolio + p_people + p_process) / 3, 1)

        st.session_state.s5_data = {
            "org_data": org_data,
            "p_portfolio": p_portfolio,
            "p_people": p_people,
            "p_process": p_process,
            "final_score": final_org
        }
        progress.progress(100)
        status.markdown("✓ Complete.")
        time.sleep(0.5)
        st.session_state.s5_step = "done"
        st.rerun()

    # ── Results ───────────────────────────────────────────────
    elif st.session_state.s5_step == "done":
        d        = st.session_state.s5_data
        org_data = d["org_data"]
        final    = d["final_score"]

        # ── Score banner ──────────────────────────────────────
        score_col = "#22c55e" if final>=7 else "#f59e0b" if final>=4 else "#ef4444"
        bop_rec = org_data.get("build_or_partner",{}).get("recommendation","Co-develop")
        st.markdown(f"""
<div style="background:#0f1e35;border-radius:8px;padding:20px 24px;margin-bottom:20px;border:1px solid #2a4a70;">
  <div style="color:#94a3b8;font-size:9px;letter-spacing:2px;font-weight:700;font-family:Arial,sans-serif;margin-bottom:4px;">ORGANISATIONAL READINESS SCORE</div>
  <div style="color:{score_col};font-size:44px;font-weight:700;line-height:1;">{final}<span style="font-size:18px;color:#94a3b8;"> / 10</span></div>
  <div style="color:{WHITE};font-size:13px;margin-top:6px;opacity:0.8;">Build strategy: <b>{bop_rec}</b></div>
</div>
""", unsafe_allow_html=True)

        # ── P³ score cards ────────────────────────────────────
        st.markdown("#### P³ Assessment — Portfolio × People × Process")
        st.caption("Schaeffler's own innovation performance formula applied to this idea's P³ Perspective")
        col1, col2, col3 = st.columns(3)
        col1.metric("Portfolio fit",  f"{d['p_portfolio']:.1f}/10", "33% weight")
        col2.metric("People (competency)", f"{d['p_people']:.1f}/10",  "33% weight")
        col3.metric("Process (assets)",   f"{d['p_process']:.1f}/10", "34% weight")
        st.markdown("---")

        # ── Portfolio dimension ───────────────────────────────
        port = org_data.get("p3_portfolio",{})
        st.markdown("#### Portfolio — Strategic fit")
        st.markdown(f"<div style='background:#1a2d45;border-radius:6px;padding:12px 16px;margin-bottom:10px;'><div style='color:#94a3b8;font-size:11px;'>RATIONALE</div><div style='color:#e2e8f0;font-size:13px;margin-top:4px;'>{port.get('rationale','')}</div><div style='color:#94a3b8;font-size:11px;margin-top:8px;'>CLUSTER FIT</div><div style='color:#60a5fa;font-size:13px;margin-top:4px;'>{port.get('cluster_fit','')}</div></div>", unsafe_allow_html=True)
        if port.get("strengths"):
            cols_ps = st.columns(len(port["strengths"][:3]))
            for i, s in enumerate(port["strengths"][:3]):
                cols_ps[i].markdown(f'<div style="background:#0a2010;border:1px solid #60a5fa33;border-radius:6px;padding:8px 12px;font-size:12px;color:#22c55e;">✓ {s}</div>', unsafe_allow_html=True)
        if port.get("gaps"):
            for g in port["gaps"][:2]:
                st.markdown(f'<div style="background:#1a0f0f;border:1px solid #ef444433;border-radius:6px;padding:8px 12px;margin-top:6px;font-size:12px;color:#ef4444;">✗ Gap: {g}</div>', unsafe_allow_html=True)
        st.markdown("---")

        # ── People dimension ──────────────────────────────────
        peop = org_data.get("p3_people",{})
        st.markdown("#### People — Competency & skills")
        st.markdown(f"<div style='background:#1a2d45;border-radius:6px;padding:12px 16px;margin-bottom:10px;'><div style='color:#94a3b8;font-size:11px;'>RATIONALE</div><div style='color:#e2e8f0;font-size:13px;margin-top:4px;'>{peop.get('rationale','')}</div></div>", unsafe_allow_html=True)
        if peop.get("matched_competencies"):
            st.markdown("**Matched Schaeffler competencies**")
            comp_cols = st.columns(min(3,len(peop["matched_competencies"])))
            for i, c in enumerate(peop["matched_competencies"][:3]):
                comp_cols[i].markdown(f'<div style="background:#0a2419;border:1px solid #60a5fa33;border-radius:4px;padding:8px;font-size:12px;color:#60a5fa;text-align:center;">{c}</div>', unsafe_allow_html=True)
        if peop.get("competency_gap"):
            st.markdown(f'<div style="background:#1a1020;border:1px solid #a78bfa44;border-radius:6px;padding:10px 14px;margin-top:8px;font-size:13px;color:#a78bfa;">⚠️ Critical gap: {peop["competency_gap"]}<br><span style="color:#94a3b8;font-size:12px;">Closure route: {peop.get("sourcing_route","")}</span></div>', unsafe_allow_html=True)
        st.markdown("---")

        # ── Process dimension ─────────────────────────────────
        proc = org_data.get("p3_process",{})
        st.markdown("#### Process — Infrastructure & assets")
        st.markdown(f"<div style='background:#1a2d45;border-radius:6px;padding:12px 16px;margin-bottom:10px;'><div style='color:#94a3b8;font-size:11px;'>RATIONALE</div><div style='color:#e2e8f0;font-size:13px;margin-top:4px;'>{proc.get('rationale','')}</div></div>", unsafe_allow_html=True)
        if proc.get("applicable_assets"):
            st.markdown("**Applicable assets / processes**")
            for a in proc["applicable_assets"][:3]:
                st.markdown(f"- {a}")
        if proc.get("investment_required"):
            st.markdown(f'<div style="background:#1a1a0f;border:1px solid #f59e0b44;border-radius:6px;padding:10px 14px;margin-top:8px;font-size:13px;color:#f59e0b;">🔧 Investment required: {proc["investment_required"]}<br><span style="color:#94a3b8;font-size:12px;">Estimated time to close: {proc.get("time_to_close","")}</span></div>', unsafe_allow_html=True)
        st.markdown("---")

        # ── Partnership candidates ────────────────────────────
        partners = org_data.get("partnership_candidates",[])
        if partners:
            st.markdown("#### Partnership candidates")
            route_cols = {"Co-develop":"#60a5fa","Acquire":"#f59e0b","License":"#a78bfa","JDA":"#22c55e"}
            for p in partners[:4]:
                rc = route_cols.get(p.get("route","Co-develop"),"#60a5fa")
                st.markdown(f"""
<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:5px 0;display:flex;align-items:flex-start;gap:12px;">
  <div style="flex:1;">
    <div style="color:#e2e8f0;font-weight:600;font-size:13px;">{p.get('name','')}</div>
    <div style="color:#94a3b8;font-size:12px;margin-top:3px;">{p.get('type','')} · {p.get('rationale','')}</div>
  </div>
  <div style="background:{rc}22;color:{rc};font-size:11px;padding:3px 10px;border-radius:10px;white-space:nowrap;">{p.get('route','')}</div>
</div>""", unsafe_allow_html=True)
            st.markdown("---")

        # ── Org gaps register ─────────────────────────────────
        gaps = org_data.get("org_gaps",[])
        if gaps:
            st.markdown("#### Organisational gaps")
            sev_col = {"High":"#ef4444","Medium":"#f59e0b","Low":"#22c55e"}
            for g in gaps[:4]:
                sc = sev_col.get(g.get("severity","Medium"),"#f59e0b")
                st.markdown(f"""
<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:5px 0;display:flex;align-items:flex-start;gap:12px;">
  <div style="flex:1;">
    <div style="color:#e2e8f0;font-weight:600;font-size:13px;">{g.get('gap','')}</div>
    <div style="color:#94a3b8;font-size:12px;margin-top:3px;">{g.get('closure_route','')} · Est. {g.get('timeline','')}</div>
  </div>
  <div style="background:{sc}22;color:{sc};font-size:11px;padding:3px 10px;border-radius:10px;white-space:nowrap;">{g.get('severity','')} severity</div>
</div>""", unsafe_allow_html=True)
            st.markdown("---")

        # ── Build or partner ──────────────────────────────────
        bop = org_data.get("build_or_partner",{})
        st.markdown("#### Build or partner?")
        bop_c1, bop_c2 = st.columns(2)
        bop_c1.markdown(f'<div style="background:#1a2d45;border-radius:8px;padding:14px 16px;"><div style="color:#94a3b8;font-size:11px;margin-bottom:4px;">RECOMMENDATION</div><div style="color:#60a5fa;font-size:16px;font-weight:600;">{bop.get("recommendation","")}</div><div style="color:#e2e8f0;font-size:12px;margin-top:6px;">{bop.get("rationale","")}</div></div>', unsafe_allow_html=True)
        bop_c2.markdown(f'<div style="background:#1a2d45;border-radius:8px;padding:14px 16px;"><div style="color:#94a3b8;font-size:11px;margin-bottom:4px;">TIME TO TRL 6</div><div style="color:#e2e8f0;font-size:13px;"><b>Internal:</b> {bop.get("time_to_trl6_internal","")}<br><b>With partner:</b> {bop.get("time_to_trl6_partner","")}</div></div>', unsafe_allow_html=True)

        # ── Chat ──────────────────────────────────────────────
        st.markdown("---")
        st.subheader(T("s5_chat_header"))
        for msg in st.session_state.s5_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        user_q = st.chat_input(T("s5_chat_ph"))
        if user_q:
            st.session_state.s5_chat.append({"role":"user","content":user_q})
            with st.chat_message("user"):
                st.markdown(user_q)
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    ctx = f"""You are a Schaeffler innovation readiness expert discussing P³ organisational assessment.
Idea: {idea} | Quadrant: {quadrant}
P³ scores — Portfolio: {d['p_portfolio']}/10 | People: {d['p_people']}/10 | Process: {d['p_process']}/10
Build recommendation: {bop.get('recommendation','')}
Critical gap: {org_data.get('p3_people',{}).get('competency_gap','')}
Be specific to Schaeffler's context (Vitesco integration, E-Mobility shift, OEM relationships). 3-4 sentences."""
                    history = [{"role":m["role"],"content":m["content"]} for m in st.session_state.s5_chat]
                    reply = call_claude_chat(ctx, history)
                    st.markdown(reply)
                    st.session_state.s5_chat.append({"role":"assistant","content":reply})

        # ── Continue ──────────────────────────────────────────
        st.markdown("---")
        _one_click_dl(
            T("dl_org"),
            lambda: generate_org_report(idea, quadrant, s1c, d),
            f"Schaeffler_Org_Readiness_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        st.markdown("---")
        st.success(T("s5_success").format(score=final))
        if st.button(T("s5_continue"), type="primary", key="s5_continue"):
            st.session_state.active_stage = 6
            st.rerun()

        if st.button(T("s5_rerun"), key="s5_rerun"):
            st.session_state.s5_step = "intro"
            st.session_state.s5_data = {}
            st.session_state.s5_chat = []
            st.session_state.s5_report_buf = None
            st.rerun()


# ════════════════════════════════════════════════════════════
# STAGE 06 — SCORING & SYNTHESIS
# ════════════════════════════════════════════════════════════
elif st.session_state.active_stage == 6:
    st.markdown(f"## {T('s6_title')}")
    st.markdown(f"""<div style="background:#1a2d45;border-radius:6px;padding:14px 18px;margin-bottom:16px;border-left:4px solid #2E75B6;">
<div style="color:#60a5fa;font-size:11px;letter-spacing:1px;font-weight:600;">{T('stage_what_label')}</div>
<div style="color:#e2e8f0;font-size:13px;margin-top:6px;">{T('s6_what')}</div>
<div style="color:#94a3b8;font-size:12px;margin-top:8px;">{T('s6_you_get')}</div>
</div>""", unsafe_allow_html=True)
    st.markdown("---")

    idea     = st.session_state.s1_idea
    s1c      = st.session_state.s1_classification
    quadrant = s1c.get("quadrant","RADICAL")

    for k, v in {"s6_step":"intro","s6_data":{},"s6_chat":[]}.items():
        if k not in st.session_state:
            st.session_state[k] = v

    # ── Check all stages are complete ─────────────────────────
    s2_done = bool(st.session_state.get("s2_data"))
    s3_done = bool(st.session_state.get("s3_data"))
    s4_done = bool(st.session_state.get("s4_data"))
    s5_done = bool(st.session_state.get("s5_data"))

    if not (s2_done and s3_done and s4_done and s5_done):
        st.warning("Complete Stages 02, 03, 04, and 05 first before running the final synthesis.")
        missing = []
        if not s2_done: missing.append("Stage 02: Market Intelligence")
        if not s3_done: missing.append("Stage 03: Patent Intelligence")
        if not s4_done: missing.append("Stage 04: Technical Feasibility")
        if not s5_done: missing.append("Stage 05: P³ Perspective")
        for m in missing:
            st.markdown(f"- ⬜ {m}")
        if st.button("← Back", key="s6_back2"):
            st.session_state.active_stage = 5
            st.rerun()
        st.stop()

    # Pull scores from previous stages
    market_score      = st.session_state.s2_data.get("final_score", 5.0)
    patent_score      = st.session_state.s3_data.get("final_score", 5.0)
    feasibility_score = st.session_state.s4_data.get("final_score", 5.0)
    org_score         = st.session_state.s5_data.get("final_score", 5.0)

    # ── Intro: show weights + let user adjust ─────────────────
    if st.session_state.s6_step == "intro":
        st.info(f"**Idea:** {idea}")
        st.markdown(f"**Quadrant:** {quadrant}")
        st.markdown("---")

        st.markdown("#### Scores from previous stages")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Market Intelligence",    f"{market_score} / 10")
        c2.metric("Patent Intelligence",    f"{patent_score} / 10")
        c3.metric("Technical Feasibility",  f"{feasibility_score} / 10")
        c4.metric("P³ Score",          f"{org_score} / 10")

        st.markdown("---")
        st.markdown("#### Innovation Potential Index — Scoring Weights")
        st.caption("Default weights are equal (25% each). Adjust to reflect your strategic priorities.")

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            w_market = st.slider("Market Intelligence", 0, 100, 25, 5, key="w_market")
        with col2:
            w_patent = st.slider("Patent Intelligence", 0, 100, 25, 5, key="w_patent")
        with col3:
            w_feasibility = st.slider("Technical Feasibility", 0, 100, 25, 5, key="w_feasibility")
        with col4:
            w_org = st.slider("P³ Score", 0, 100, 25, 5, key="w_org")

        total_weight = w_market + w_patent + w_feasibility + w_org
        if total_weight != 100:
            st.warning(f"Weights must add up to 100. Current total: {total_weight}. Adjust the sliders.")
        else:
            st.success(f"✓ Weights sum to 100")
            if st.button("Run Final Synthesis →", type="primary"):
                st.session_state.s6_weights = {"market": w_market, "patent": w_patent, "feasibility": w_feasibility, "org": w_org}
                st.session_state.s6_step = "running"
                st.rerun()

    # ── Running ───────────────────────────────────────────────
    elif st.session_state.s6_step == "running":
        progress = st.progress(0)
        status   = st.empty()

        weights = st.session_state.get("s6_weights", {"market":25,"patent":25,"feasibility":25,"org":25})

        status.markdown("📐 Calculating Innovation Potential Index...")
        progress.progress(20)

        wm = weights["market"] / 100
        wp = weights["patent"] / 100
        wf = weights["feasibility"] / 100
        wo = weights.get("org", 15) / 100
        ipi = round(market_score * wm + patent_score * wp + feasibility_score * wf + org_score * wo, 1)

        status.markdown("🧠 Writing narrative synthesis...")
        progress.progress(45)

        s2d = st.session_state.s2_data
        s3d = st.session_state.s3_data
        s4d = st.session_state.s4_data
        s5d = st.session_state.s5_data
        org_d = s5d.get("org_data", {})

        system_synthesis = """You are a senior Schaeffler innovation strategist writing a final investment recommendation.
Synthesise all four stages of analysis into a clear, actionable assessment.

Be direct, specific, and honest — including about weaknesses. 
Reference Schaeffler's strategic context (electrification transition, Vitesco merger, E-Mobility growth).

Return ONLY valid JSON:
{
  "headline": "one punchy sentence summarising the overall verdict",
  "recommendation": "PROCEED / PROCEED WITH CONDITIONS / DEFER / REJECT",
  "recommendation_rationale": "2-3 sentences on why this recommendation",
  "strongest_signals": ["top positive signal 1", "top positive signal 2", "top positive signal 3"],
  "key_concerns": ["main concern 1", "main concern 2", "main concern 3"],
  "conditions": ["condition 1 if applicable", "condition 2"],
  "strategic_fit": "2-3 sentences on how this fits Schaeffler's current strategic priorities",
  "next_steps": ["concrete next step 1", "concrete next step 2", "concrete next step 3"],
  "narrative": "5-6 paragraph narrative synthesis covering the full picture — market opportunity, IP position, technical readiness, and strategic recommendation"
}"""

        synthesis_context = f"""
Idea: {idea}
Quadrant: {quadrant}
Innovation Potential Index: {ipi}/10

Innovation Cluster: {st.session_state.s1_classification.get('innovation_cluster','')}
Product Family: {st.session_state.s1_classification.get('product_family','')}
Pipeline Route: {st.session_state.s1_classification.get('pipeline_route','')}
Trend Alignment: {', '.join(st.session_state.s1_classification.get('trend_alignment',[]))}

Stage 02 — Market Intelligence ({weights['market']}% weight): {market_score}/10
- Market: {s2d.get('market',{}).get('market_name','')}
- Size: {_mval(s2d.get('market',{}).get('market_size_current') or s2d.get('market',{}).get('market_size_2024',''))}
- CAGR: {s2d.get('market',{}).get('cagr','')}
- Maturity: {s2d.get('market',{}).get('market_maturity','')}
- Primary sectors: {', '.join(s2d.get('sectors',{}).get('primary_sectors',[]))}
- Competition: {s2d.get('comp',{}).get('competitive_intensity','')}
- White space: {s2d.get('comp',{}).get('white_space','')}

Stage 03 — Patent Intelligence ({weights['patent']}% weight): {patent_score}/10
- Activity: {s3d.get('landscape',{}).get('activity_level','')}
- Trend: {s3d.get('landscape',{}).get('filing_trend','')}
- Novelty signal: {s3d.get('ansoff_data',{}).get('novelty_signal','')}
- IP risk: {s3d.get('ansoff_data',{}).get('ip_risk','')}
- Schaeffler IP gap: {s3d.get('ansoff_data',{}).get('schaeffler_position',{}).get('gap','')}

Stage 04 — Technical Feasibility ({weights['feasibility']}% weight): {feasibility_score}/10
- TRL: {s4d.get('trl',{}).get('trl_level','')} — {s4d.get('trl',{}).get('trl_label','')}
- Existence: {s4d.get('existence',{}).get('existence_verdict','')}
- Entry readiness: {s4d.get('trl',{}).get('schaeffler_entry_readiness','')}
- Time to readiness: {s4d.get('existence',{}).get('time_to_readiness','')}
{('⚠️ SENSING PHASE (FUTURE OPTIONS TRACK): This idea was classified as a Future Option at Stage 1 — technology novelty in the borderline 5.0–6.5 band or no innovation cluster assigned yet. Frame the recommendation and next steps around what evidence is needed to graduate this idea to the full pipeline. Use "SENSING PHASE — CONTINUE MONITORING" as the recommendation where the data supports further maturation rather than immediate investment or rejection.' if st.session_state.s1_classification.get('route') == 'FUTURE_OPTIONS' else '')}

Stage 05 — P³ Perspective ({weights.get('org',15)}% weight): {org_score}/10
- P³ Portfolio: {s5d.get('p_portfolio',5)}/10
- P³ People: {s5d.get('p_people',5)}/10
- P³ Process: {s5d.get('p_process',5)}/10
- Critical competency gap: {org_d.get('p3_people',{}).get('competency_gap','')}
- Build strategy: {org_d.get('build_or_partner',{}).get('recommendation','')}
- Time to TRL6 with partner: {org_d.get('build_or_partner',{}).get('time_to_trl6_partner','')}
- Innovation Model: {st.session_state.s1_classification.get('innovation_model','Integrated') or 'Integrated'}
- Pipeline Route: {st.session_state.s1_classification.get('pipeline_route','')}
"""

        # ── Call 1: structured fields ─────────────────────────
        system_structured = """You are a senior Schaeffler innovation strategist.
Return ONLY valid JSON with exactly these fields — no markdown, no extra text, no trailing commas:
{
  "headline": "one direct sentence summarising the overall verdict on this idea",
  "recommendation": "PROCEED or PROCEED WITH CONDITIONS or SENSING PHASE — CONTINUE MONITORING or DEFER or REJECT",
  "recommendation_rationale": "2-3 sentences explaining why this recommendation",
  "strongest_signals": ["specific positive signal from the data 1", "specific positive signal 2", "specific positive signal 3"],
  "key_concerns": ["specific concern from the data 1", "specific concern 2", "specific concern 3"],
  "conditions": ["specific condition to meet before proceeding 1", "condition 2"],
  "strategic_fit": "2-3 sentences on how this fits Schaeffler strategy — reference electrification transition, Vitesco merger, E-Mobility growth, and the specific product family and innovation cluster",
  "risks": ["specific risk 1 with mitigation approach", "specific risk 2 with mitigation approach", "specific risk 3 with mitigation approach"],
  "next_steps": ["concrete Schaeffler-specific action step 1", "concrete action step 2", "concrete action step 3", "concrete action step 4"]
}"""
        raw1 = call_claude(system_structured, synthesis_context, max_tokens=2000)
        raw1_clean = raw1.strip().replace("```json","").replace("```","").strip()
        fb = raw1_clean.find("{"); lb = raw1_clean.rfind("}") + 1
        if fb >= 0: raw1_clean = raw1_clean[fb:lb]
        try:
            synthesis_structured = json.loads(raw1_clean)
        except:
            synthesis_structured = {
                "headline": f"IPI score of {ipi}/10 — {'strong' if ipi>=7 else 'moderate' if ipi>=4 else 'weak'} opportunity in {s2d.get('market',{}).get('market_name','this market')}.",
                "recommendation": "PROCEED WITH CONDITIONS" if ipi >= 5 else "DEFER",
                "recommendation_rationale": f"The idea scores {ipi}/10 across market, patent, and feasibility dimensions. Market opportunity is {'strong' if market_score>=7 else 'moderate'}, IP position is {'favourable' if patent_score>=7 else 'mixed'}, and technical readiness is {'high' if feasibility_score>=7 else 'still developing'}.",
                "strongest_signals": [
                    f"Market Intelligence score of {market_score}/10 — {s2d.get('market',{}).get('market_name','')} shows growth",
                    f"Primary sector fit: {', '.join(s2d.get('sectors',{}).get('primary_sectors',[])[:2])}",
                    f"Patent novelty signal: {s3d.get('ansoff_data',{}).get('novelty_signal','Moderate')}"
                ],
                "key_concerns": [
                    f"Technical feasibility at TRL {s4d.get('trl',{}).get('trl_level',3)} — {s4d.get('trl',{}).get('schaeffler_entry_readiness','')}",
                    f"IP risk level: {s3d.get('ansoff_data',{}).get('ip_risk','Medium')}",
                    f"Competitive intensity: {s2d.get('comp',{}).get('competitive_intensity','')}"
                ],
                "conditions": [
                    "Confirm IP freedom-to-operate before committing R&D budget",
                    "Validate technical feasibility with internal engineering team"
                ],
                "strategic_fit": f"This idea aligns with Schaeffler's {quadrant} innovation quadrant and targets {', '.join(s2d.get('sectors',{}).get('primary_sectors',[])[:2])} — sectors central to Schaeffler's post-Vitesco portfolio. The electrification transition creates urgency for exactly this type of innovation investment.",
                "risks": [
                    f"IP risk: {s3d.get('ansoff_data',{}).get('ip_risk','Medium')} — conduct freedom-to-operate analysis before R&D commitment",
                    f"Technical maturity: TRL {s4d.get('trl',{}).get('trl_level',3)} — further development required before production readiness",
                    "Market timing risk — validate demand with target customers before scaling investment"
                ],
                "next_steps": [
                    "Commission internal engineering feasibility review within 30 days",
                    "Conduct freedom-to-operate IP analysis with Schaeffler patent team",
                    f"Identify pilot customer in {s2d.get('sectors',{}).get('primary_sectors',['target sector'])[0]} for co-development conversation",
                    "Present to Innovation steering committee with this assessment as supporting material"
                ]
            }

        # ── Call 2: narrative separately as plain text ────────
        system_narrative = f"""You are a senior Schaeffler innovation strategist. Write a 5-paragraph narrative synthesis for this innovation assessment. Write in flowing prose — no bullet points, no headers. Be specific about the market opportunity, IP landscape, technical maturity, and strategic recommendation. Reference Schaeffler's context (electrification, Vitesco merger, E-Mobility growth, OEM relationships)."""
        raw2 = call_claude(system_narrative, synthesis_context + f"\n\nRecommendation: {synthesis_structured.get('recommendation','')}\nIPI: {ipi}/10", max_tokens=1400)
        narrative_text = raw2.strip().replace("```","").strip()

        synthesis = {**synthesis_structured, "narrative": narrative_text}

        progress.progress(80)
        status.markdown("✓ Complete.")
        time.sleep(0.5)

        st.session_state.s6_data = {
            "ipi": ipi,
            "weights": weights,
            "synthesis": synthesis,
            "scores": {
                "market": market_score,
                "patent": patent_score,
                "feasibility": feasibility_score,
                "p3": org_score
            }
        }
        st.session_state.s6_step = "done"
        st.rerun()

    # ── Results ───────────────────────────────────────────────
    elif st.session_state.s6_step == "done":
        d         = st.session_state.s6_data
        ipi       = d.get("ipi", 0)
        weights   = d.get("weights", {"market":25,"patent":25,"feasibility":25,"org":25})
        synthesis = d.get("synthesis", {})
        scores    = d.get("scores", {"market":5,"patent":5,"feasibility":5,"p3":5})

        rec = synthesis.get("recommendation","PROCEED WITH CONDITIONS")
        rec_colours = {
            "PROCEED":                             "#22c55e",
            "PROCEED WITH CONDITIONS":             "#f59e0b",
            "SENSING PHASE — CONTINUE MONITORING": "#818cf8",
            "DEFER":                               "#f97316",
            "REJECT":                              "#ef4444"
        }
        rec_col = rec_colours.get(rec, "#f59e0b")
        ipi_col = "#22c55e" if ipi>=7 else "#f59e0b" if ipi>=4 else "#ef4444"

        # ── Auto-save to Google Sheets (once per session) ─────
        if not st.session_state.get("_s6_saved_to_sheets"):
            s2d_sv = st.session_state.get("s2_data", {})
            s3d_sv = st.session_state.get("s3_data", {})
            s4d_sv = st.session_state.get("s4_data", {})
            s5d_sv = st.session_state.get("s5_data", {})
            qa_pairs = []
            for q, a in zip(st.session_state.get("s1_questions",[]), st.session_state.get("s1_answers",[])):
                qa_pairs.append(f"Q: {q} / A: {a}")
            row = {
                "Date":                  datetime.now().strftime("%Y-%m-%d %H:%M"),
                "Submitter Name":        st.session_state.get("user_name",""),
                "Position":              st.session_state.get("user_position",""),
                "Department":            st.session_state.get("user_dept",""),
                "Full Idea Description": st.session_state.get("s1_idea",""),
                "Clarifying Q&A":        " | ".join(qa_pairs),
                "Quadrant":              quadrant,
                "Innovation Cluster":    s1c.get("innovation_cluster",""),
                "Product Family":        s1c.get("product_family",""),
                "Market Score":          str(scores.get("market","")),
                "Patent Score":          str(scores.get("patent","")),
                "Feasibility Score":     str(scores.get("feasibility","")),
                "P³ Score":   str(scores.get("p3","")),
                "IPI Score":             str(ipi),
                "Recommendation":        rec,
                "Key Concerns":          " | ".join(synthesis.get("key_concerns",[])[:3]),
                "Next Steps":            " | ".join(synthesis.get("next_steps",[])[:4]),
                "Market Name":           s2d_sv.get("market",{}).get("market_name",""),
                "Market Size 2024":      _mval(s2d_sv.get("market",{}).get("market_size_current") or s2d_sv.get("market",{}).get("market_size_2024","")),
                "CAGR":                  _mval(s2d_sv.get("market",{}).get("cagr","")),
                "TRL Level":             str(s4d_sv.get("trl",{}).get("trl_level","")),
                "Build Strategy":        s5d_sv.get("org_data",{}).get("build_or_partner",{}).get("recommendation",""),
            }
            saved, save_err = save_idea_to_sheets(row)
            if saved:
                st.session_state["_s6_saved_to_sheets"] = True
                st.success("✅ Idea automatically saved to the Innovation Ideas Log.")
            else:
                st.warning(f"⚠️ Could not save to Ideas Log — {save_err}. Your analysis is complete; the log entry can be added manually.")

        # ── Sensing-phase banner (Future Options ideas only) ──────
        if s1c.get("route") == "FUTURE_OPTIONS":
            st.markdown(f"""
<div style="background:#0f0f1e;border-radius:8px;padding:14px 20px;margin-bottom:16px;border:1px solid #818cf844;border-left:4px solid #818cf8;">
  <div style="color:#818cf8;font-size:10px;font-weight:700;letter-spacing:1.5px;margin-bottom:6px;">⬡ FUTURE OPTIONS TRACK — SENSING PHASE ASSESSMENT</div>
  <div style="color:#c7d2fe;font-size:13px;line-height:1.6;">
    This idea entered the pipeline as a <b>Future Option</b> — technology novelty score <b>{s1c.get('tech_score_final', '')}/10</b> ({s1c.get('technology_level','')}) places it in the sensing band.
    The scores below reflect current maturity. The recommendation and next steps are framed around
    <b>what evidence is needed to graduate this idea to the full Innovation Pipeline</b>, not a standard invest/reject decision.
  </div>
  <div style="color:#818cf8;font-size:11px;margin-top:8px;">Assigned cluster: <b>{s1c.get('innovation_cluster','Not yet assigned') or 'Not yet assigned'}</b> · Quadrant: <b>{quadrant}</b></div>
</div>
""", unsafe_allow_html=True)

        # ── IPI banner ────────────────────────────────────────
        _inno_model   = s1c.get("innovation_model", "") or ""
        _project_type = s1c.get("project_type", "") or ""
        _pipeline_rt  = s1c.get("pipeline_route", "") or _inno_model
        _model_col    = "#22c55e" if "Integrated" in _inno_model else "#60a5fa" if _inno_model else "#94a3b8"
        _model_pill   = ""
        if _inno_model or _project_type:
            _pill_parts = [p for p in [_inno_model, _project_type] if p]
            _model_pill = f'<div style="display:inline-flex;gap:6px;margin-top:8px;flex-wrap:wrap;">' + \
                "".join(f'<span style="background:{_model_col}22;color:{_model_col};font-size:10px;font-weight:700;padding:3px 10px;border-radius:10px;letter-spacing:0.5px;border:1px solid {_model_col}44;">{p}</span>' for p in _pill_parts) + \
                f'</div>'

        st.markdown(f"""
<div style="background:#0f1e35;border-radius:8px;padding:20px 24px;margin-bottom:20px;border:1px solid #2a4a70;">
  <div style="display:flex;justify-content:space-between;align-items:flex-start;">
    <div>
      <div style="color:#94a3b8;font-size:9px;letter-spacing:2px;font-weight:700;font-family:Arial,sans-serif;margin-bottom:4px;">INNOVATION POTENTIAL INDEX</div>
      <div style="color:{ipi_col};font-size:48px;font-weight:700;line-height:1;">{ipi}<span style="font-size:20px;color:#94a3b8;"> / 10</span></div>
      <div style="color:{WHITE};font-size:13px;margin-top:6px;opacity:0.8;">{synthesis.get('headline','')}</div>
      {_model_pill}
    </div>
    <div style="text-align:right;">
      <div style="color:{WHITE};font-size:11px;opacity:0.5;margin-bottom:4px;">RECOMMENDATION</div>
      <div style="background:{rec_col}22;color:{rec_col};font-size:16px;font-weight:700;padding:8px 16px;border-radius:6px;border:1px solid {rec_col}44;">{rec}</div>
      <div style="color:#94a3b8;font-size:10px;margin-top:6px;">{quadrant} · {s1c.get('innovation_cluster','')}</div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)

        # ── Radar chart ───────────────────────────────────────
        st.markdown("#### 📡 Innovation Potential Radar")

        categories   = ["Market Intelligence", "Patent Intelligence", "Technical Feasibility", "P³ Score"]
        values       = [scores["market"], scores["patent"], scores["feasibility"], scores.get("p3", 5)]
        values_close = values + [values[0]]
        cats_close   = categories + [categories[0]]

        fig_radar = go.Figure()
        fig_radar.add_trace(go.Scatterpolar(
            r=values_close, theta=cats_close,
            fill="toself",
            fillcolor=f"rgba(96,165,250,0.15)",
            line=dict(color=BLUE, width=2),
            marker=dict(size=8, color=BLUE),
            name="Score"
        ))
        # Add benchmark line at 7
        fig_radar.add_trace(go.Scatterpolar(
            r=[7,7,7,7,7], theta=cats_close,
            line=dict(color="#4a6fa5", width=1, dash="dot"),
            marker=dict(size=0),
            fill=None,
            name="Target (7/10)",
            showlegend=True
        ))
        fig_radar.update_layout(
            polar=dict(
                bgcolor=BG,
                radialaxis=dict(
                    visible=True, range=[0,10],
                    tickfont=dict(color=WHITE, size=10),
                    gridcolor="#2a4a70", linecolor="#2a4a70",
                    tickvals=[2,4,6,8,10]
                ),
                angularaxis=dict(
                    tickfont=dict(color=WHITE, size=12),
                    gridcolor="#2a4a70", linecolor="#2a4a70"
                )
            ),
            paper_bgcolor=BG,
            plot_bgcolor=BG,
            legend=dict(font=dict(color=WHITE), bgcolor=BG),
            height=400,
            margin=dict(l=60,r=60,t=40,b=40)
        )
        st.plotly_chart(fig_radar, use_container_width=True)

        # Score + weight breakdown
        col1, col2, col3, col4 = st.columns(4)
        col1.metric("Market Intelligence",   f"{scores['market']} / 10",   f"{weights['market']}% weight")
        col2.metric("Patent Intelligence",   f"{scores['patent']} / 10",   f"{weights['patent']}% weight")
        col3.metric("Technical Feasibility", f"{scores['feasibility']} / 10", f"{weights['feasibility']}% weight")
        col4.metric("P³ Score",         f"{scores.get('p3',5)} / 10", f"{weights.get('org',15)}% weight")
        st.markdown("---")

        # ── Recommendation ────────────────────────────────────
        st.markdown("#### 🎯 Recommendation")
        st.markdown(f"""
<div style="background:{rec_col}11;border:1px solid {rec_col}44;border-radius:8px;padding:16px 20px;margin-bottom:16px;">
  <div style="color:{rec_col};font-size:16px;font-weight:700;margin-bottom:6px;">{rec}</div>
  <div style="color:{WHITE};font-size:13px;">{synthesis.get('recommendation_rationale','')}</div>
</div>
""", unsafe_allow_html=True)

        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown("**Strongest signals**")
            for s in synthesis.get("strongest_signals",[])[:3]:
                st.markdown(f"✅ {s}")
        with col_b:
            st.markdown("**Key concerns**")
            for c in synthesis.get("key_concerns",[])[:3]:
                st.markdown(f"⚠️ {c}")

        if synthesis.get("conditions"):
            st.markdown("**Conditions / requirements:**")
            for cond in synthesis.get("conditions",[])[:3]:
                st.markdown(f"→ {cond}")
        st.markdown("---")

        # ── Strategic fit ─────────────────────────────────────
        st.markdown("#### 🏭 Strategic Fit for Schaeffler")
        st.markdown(synthesis.get("strategic_fit",""))
        st.markdown("---")

        # ── Next steps ────────────────────────────────────────
        st.markdown("#### 📋 Recommended Next Steps")
        for i, step in enumerate(synthesis.get("next_steps",[])[:5], 1):
            st.markdown(f"**{i}.** {step}")
        st.markdown("---")

        # ── Full narrative ────────────────────────────────────
        with st.expander("#### 📖 Full Narrative Synthesis", expanded=False):
            st.markdown(synthesis.get("narrative",""))

        # ── Visual Innovation Brief ───────────────────────────
        st.markdown("---")
        st.markdown("#### 📊 Development Roadmap")
        st.caption("Indicative development timeline based on current TRL and P³ build/partner recommendation.")

        import re as _re_bp
        _s4v = st.session_state.get("s4_data", {})
        _s5v = st.session_state.get("s5_data", {})

        # ── Gantt: Development Roadmap ────────────────────────
        def _parse_months(s, default):
            nums = _re_bp.findall(r'\d+', str(s))
            return int(nums[0]) if nums else default

        _trl_val = _s4v.get("trl", {}).get("trl_level", 3)
        try: _trl_val = int(_trl_val)
        except: _trl_val = 3

        _t_partner  = _parse_months(_s5v.get("org_data",{}).get("build_or_partner",{}).get("time_to_trl6_partner","18"), 18)

        _roadmap = [
            ("Sensing & Validation",          0,             3,             "#818cf8"),
            ("Proof of Concept (TRL 4)",       3,             6,             "#60a5fa"),
            ("Prototype (TRL 5–6)",            6,             _t_partner,    "#34d399"),
            ("Pilot / Field Trial (TRL 7–8)",  _t_partner,    _t_partner+10, "#f59e0b"),
            ("Commercialisation (TRL 9)",      _t_partner+10, _t_partner+16, "#22c55e"),
        ]
        fig_road = go.Figure()
        for task, start, end, colour in _roadmap:
            fig_road.add_trace(go.Bar(
                x=[max(1, end-start)], y=[task], base=[start], orientation="h",
                marker_color=colour, marker_line_width=0,
                text=f"M{start}–M{end}", textposition="inside",
                textfont=dict(color=WHITE, size=10), showlegend=False,
                hovertemplate=f"<b>{task}</b><br>Month {start} → Month {end}<extra></extra>"
            ))
        _trl_month = {1:0,2:1,3:2,4:4,5:7,6:10,7:14,8:20,9:28}.get(_trl_val, 3)
        fig_road.add_vline(x=_trl_month, line_color="#ef4444", line_width=2, line_dash="dot",
                           annotation_text=f"Current TRL {_trl_val}",
                           annotation_font_color="#ef4444", annotation_position="top right")
        fig_road.update_layout(
            title=dict(text="Indicative Development Roadmap (months from today)", font=dict(size=13,color=WHITE), x=0.5),
            barmode="overlay", plot_bgcolor=BG, paper_bgcolor=BG, height=300,
            xaxis=dict(title="Months", tickfont=dict(color=WHITE), gridcolor="#2a4a70",
                       title_font=dict(color=DIM), range=[0, _t_partner+18]),
            yaxis=dict(tickfont=dict(color=WHITE, size=11)),
            margin=dict(l=10,r=40,t=50,b=40), font=dict(color=WHITE)
        )
        st.plotly_chart(fig_road, use_container_width=True)

        # ── Risk Register ─────────────────────────────────────
        _risks_vis = synthesis.get("risks", [])
        if _risks_vis:
            st.markdown("**⚠️ Risk Register**")
            _sev_map = {"High":"#ef4444","Medium":"#f59e0b","Low":"#22c55e"}
            for r in _risks_vis[:4]:
                _r_text = r.get("risk", str(r)) if isinstance(r, dict) else str(r)
                _r_mit  = r.get("mitigation","") if isinstance(r, dict) else ""
                _sev = "High" if any(w in _r_text.lower() for w in ["critical","major","significant","high"]) else \
                       "Low"  if any(w in _r_text.lower() for w in ["minor","low","small"]) else "Medium"
                _rc2 = _sev_map[_sev]
                st.markdown(f"""
<div style="background:#1a2d45;border-radius:6px;padding:10px 14px;margin:4px 0;border-left:3px solid {_rc2};">
  <div style="display:flex;justify-content:space-between;align-items:center;">
    <span style="color:#e2e8f0;font-size:12px;font-weight:600;">{_r_text}</span>
    <span style="color:{_rc2};font-size:10px;background:{_rc2}22;padding:2px 8px;border-radius:8px;">{_sev}</span>
  </div>
  {f'<div style="color:#94a3b8;font-size:11px;margin-top:4px;">→ {_r_mit}</div>' if _r_mit else ''}
</div>""", unsafe_allow_html=True)

        # ── 90-Day Action Brief (plain text — no JSON) ────────
        st.markdown("---")
        st.markdown("#### 📍 90-Day Action Brief")
        if st.session_state.get("s6_action_brief") is None:
            if st.button("Generate 90-Day Action Brief →", type="secondary", key="s6_action_brief_btn"):
                with st.spinner("Writing action brief..."):
                    _brief_sys = "You are a senior Schaeffler innovation strategist. Write a concise, direct 90-day action brief for this idea. 3–4 short paragraphs. Plain prose — no bullet points, no JSON. Reference Schaeffler's specific context, the innovation cluster, and the build/partner recommendation. Be concrete about who does what and by when."
                    _brief_ctx = (
                        f"Idea: {idea}\nQuadrant: {quadrant}\nIPI: {ipi}/10\nRecommendation: {rec}\n"
                        f"Cluster: {s1c.get('innovation_cluster','')}\n"
                        f"Build strategy: {_s5v.get('org_data',{}).get('build_or_partner',{}).get('recommendation','')}\n"
                        f"Key concern: {synthesis.get('key_concerns',[''])[0]}\n"
                        f"Strategic fit: {synthesis.get('strategic_fit','')}"
                    )
                    _brief = call_claude(_brief_sys, _brief_ctx, max_tokens=600)
                    st.session_state.s6_action_brief = _brief.strip().replace("```","")
                    st.rerun()
        else:
            st.markdown(f"""
<div style="background:#007A3D11;border:1px solid #007A3D44;border-radius:8px;padding:16px 20px;">
  <div style="color:#e2e8f0;font-size:13px;line-height:1.75;">{st.session_state.s6_action_brief}</div>
</div>""", unsafe_allow_html=True)
            if st.button("↺ Regenerate", key="s6_action_brief_redo"):
                st.session_state.s6_action_brief = None
                st.rerun()

        # ── Master report download ────────────────────────────
        st.markdown("---")
        _s2d = st.session_state.s2_data
        _s3d = st.session_state.s3_data
        _s4d = st.session_state.s4_data
        _s5d = st.session_state.s5_data
        _one_click_dl(
            T("dl_master"),
            lambda: generate_master_report(idea, quadrant, s1c, _s2d, _s3d, _s4d, _s5d, d),
            f"Schaeffler_Innovation_Assessment_{datetime.now().strftime('%Y%m%d')}.docx"
        )
        st.caption(T("dl_caption"))

        # ── Chat ──────────────────────────────────────────────
        st.markdown("---")
        st.subheader(T("s6_chat_header"))
        for msg in st.session_state.s6_chat:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        user_q = st.chat_input(T("s6_chat_ph"))
        if user_q:
            st.session_state.s6_chat.append({"role":"user","content":user_q})
            with st.chat_message("user"):
                st.markdown(user_q)
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    ctx = f"""You are a senior Schaeffler innovation strategist discussing the final assessment.
Idea: {idea} | Quadrant: {quadrant}
IPI Score: {ipi}/10 | Recommendation: {rec}
Market: {scores['market']}/10 | Patent: {scores['patent']}/10 | Feasibility: {scores['feasibility']}/10 | Org: {scores.get('p3',5)}/10
Headline: {synthesis.get('headline','')}
Strongest signals: {synthesis.get('strongest_signals',[])}
Key concerns: {synthesis.get('key_concerns',[])}
Strategic fit: {synthesis.get('strategic_fit','')}
Be direct and specific. Reference Schaeffler's context where relevant. 3-4 sentences."""
                    history = [{"role":m["role"],"content":m["content"]} for m in st.session_state.s6_chat]
                    reply = call_claude_chat(ctx, history)
                    st.markdown(reply)
                    st.session_state.s6_chat.append({"role":"assistant","content":reply})

        # ── Re-run with different weights ─────────────────────
        st.markdown("---")
        if st.button(T("s6_rerun"), key="s6_rerun"):
            st.session_state.s6_step = "intro"
            st.session_state.s6_data = {}
            st.session_state.s6_chat = []
            st.session_state.s6_report_buf = None
            st.rerun()
